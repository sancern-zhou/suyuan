#!/usr/bin/env python3
"""
臭氧垂直报告批量处理脚本

功能：
- 自动批量处理臭氧垂直分析报告
- 依次执行5个分析和替换步骤
- 进度跟踪和错误恢复
- 生成处理报告

使用方法：
1. 配置 reports_dir 指向报告文件夹
2. 运行脚本：python batch_ozone_report_processor.py
3. 查看处理报告：batch_ozone_report_processor.json

作者：Claude Code
日期：2026-04-14
"""

import os
import sys
import json
import asyncio
import httpx
import base64
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
from docx import Document
import structlog

# 配置日志
logger = structlog.get_logger()

# ==================== 配置区 ====================

# 报告文件夹路径（请修改为实际路径）
REPORTS_DIR = "/home/xckj/suyuan/ozone_reports"

# 输出文件夹路径（处理后的报告）
OUTPUT_DIR = "/home/xckj/suyuan/ozone_reports_processed"

# 进度文件路径
PROGRESS_FILE = "/home/xckj/suyuan/batch_ozone_report_processor_progress.json"

# 处理报告路径
REPORT_FILE = "/home/xckj/suyuan/batch_ozone_report_processor_report.json"

# LLM API配置（从环境变量自动获取）
LLM_API_KEY = os.getenv("QWEN_VL_API_KEY", "")
LLM_BASE_URL = os.getenv("QWEN_VL_BASE_URL", "https://dashscope.aliyuncs.com/compatible-mode/v1")
LLM_MODEL = os.getenv("QWEN_VL_MODEL", "qwen-vl-max-latest")

# 并发处理数（同时处理多少份报告）
# 注意：通义千问API有并发限制，建议设置为1（顺序执行）
# 如果API并发限制较高，可适当增加（不建议超过3）
CONCURRENT_TASKS = 1

# API限流保护（每个请求之间的间隔，秒）
API_REQUEST_INTERVAL = 2  # 每次API请求间隔2秒

# API重试配置
MAX_RETRIES = 3  # 最大重试次数
RETRY_DELAY = 5  # 重试延迟（秒）

# ==================== 工具函数 ====================

class BatchProcessor:
    """批量处理器"""

    def __init__(self):
        self.reports_dir = Path(REPORTS_DIR)
        self.output_dir = Path(OUTPUT_DIR)
        self.progress_file = Path(PROGRESS_FILE)
        self.report_file = Path(REPORT_FILE)

        # 创建输出目录
        self.output_dir.mkdir(parents=True, exist_ok=True)

        # 加载进度
        self.progress = self._load_progress()

        # 统计信息
        self.stats = {
            "total": 0,
            "success": 0,
            "failed": 0,
            "skipped": 0,
            "start_time": None,
            "end_time": None
        }

    def _load_progress(self) -> Dict[str, Any]:
        """加载进度文件"""
        if self.progress_file.exists():
            try:
                with open(self.progress_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning("load_progress_failed", error=str(e))
        return {
            "completed": [],
            "failed": [],
            "current_file": None
        }

    def _save_progress(self):
        """保存进度文件"""
        try:
            with open(self.progress_file, 'w', encoding='utf-8') as f:
                json.dump(self.progress, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error("save_progress_failed", error=str(e))

    def _save_report(self):
        """保存处理报告"""
        try:
            report = {
                "statistics": self.stats,
                "progress": self.progress,
                "generated_at": datetime.now().isoformat()
            }
            with open(self.report_file, 'w', encoding='utf-8') as f:
                json.dump(report, f, ensure_ascii=False, indent=2)
        except Exception as e:
            logger.error("save_report_failed", error=str(e))

    def get_reports(self) -> List[Path]:
        """获取所有待处理的报告文件"""
        if not self.reports_dir.exists():
            logger.error("reports_dir_not_exists", dir=str(self.reports_dir))
            return []

        # 查找所有Word文档
        reports = list(self.reports_dir.glob("*.docx"))
        logger.info("found_reports", count=len(reports))

        # 过滤已完成的报告
        pending = [r for r in reports if r.name not in self.progress["completed"]]
        logger.info("pending_reports", total=len(reports), pending=len(pending))

        return pending

    async def process_report(self, report_path: Path) -> bool:
        """处理单份报告

        Args:
            report_path: 报告文件路径

        Returns:
            是否处理成功
        """
        report_name = report_path.name
        logger.info("processing_report", name=report_name)

        try:
            # 1. 解包文档
            unpack_dir = await self._unpack_document(report_path)
            if not unpack_dir:
                return False

            # 2. 执行5个步骤
            steps_result = await self._execute_steps(report_path, unpack_dir)

            # 3. 保存到输出目录
            output_path = self.output_dir / report_name
            if report_path != output_path:
                import shutil
                shutil.copy2(report_path, output_path)
                logger.info("report_saved", path=str(output_path))

            # 记录完成
            self.progress["completed"].append(report_name)
            if report_name in self.progress["failed"]:
                self.progress["failed"].remove(report_name)
            self._save_progress()

            logger.info("report_processed_success", name=report_name)
            return True

        except Exception as e:
            logger.error("report_processing_failed", name=report_name, error=str(e))
            self.progress["failed"].append(report_name)
            self._save_progress()
            return False

    async def _unpack_document(self, doc_path: Path) -> Optional[Path]:
        """解包Word文档"""
        import zipfile

        try:
            # 创建解包目录
            unpack_dir = doc_path.parent / f"unpacked_{doc_path.stem}"
            unpack_dir.mkdir(parents=True, exist_ok=True)

            # 解包
            with zipfile.ZipFile(doc_path, 'r') as zip_ref:
                zip_ref.extractall(unpack_dir)

            logger.info("document_unpacked", path=str(doc_path), unpack_dir=str(unpack_dir))
            return unpack_dir

        except Exception as e:
            logger.error("unpack_failed", path=str(doc_path), error=str(e))
            return None

    async def _execute_steps(self, doc_path: Path, unpack_dir: Path) -> bool:
        """执行5个处理步骤"""

        # 打开文档
        doc = Document(str(doc_path))

        # 提取图片路径
        image_files = self._extract_images(unpack_dir)

        # 执行步骤
        steps = [
            self._step1_data_analysis,
            self._step2_no2_o3_analysis,
            self._step3_spatial_distribution,
            self._step4_vertical_distribution,
            self._step5_summary
        ]

        for i, step_func in enumerate(steps, 1):
            try:
                logger.info("executing_step", step=i)

                # 调用步骤函数
                result = await step_func(doc, image_files, unpack_dir)

                if result:
                    logger.info("step_completed", step=i)
                else:
                    logger.warning("step_failed", step=i)
                    # 继续执行下一步骤

            except Exception as e:
                logger.error("step_error", step=i, error=str(e))
                # 继续执行下一步骤

        # 保存文档
        doc.save(str(doc_path))
        return True

    def _extract_images(self, unpack_dir: Path) -> List[Path]:
        """提取文档中的图片"""
        media_dir = unpack_dir / "word" / "media"
        if media_dir.exists():
            images = sorted(media_dir.glob("*"))
            return images
        return []

    async def _step1_data_analysis(self, doc: Document, images: List[Path], unpack_dir: Path) -> bool:
        """第1步：数据特征分析"""
        try:
            # 读取表格数据
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)

            # 构造提示词
            prompt = """请分析以下表格数据，给出数据特征分析。

要求：
- 不需要任何结论和建议
- 找出关键信息
- 表达流畅易读
- 内容丰富

表格数据：
""" + str(tables)

            # 调用LLM
            analysis = await self._call_llm(prompt)

            # 替换文本
            self._replace_text(doc, "数据特征分析：", analysis)

            return True

        except Exception as e:
            logger.error("step1_failed", error=str(e))
            return False

    async def _step2_no2_o3_analysis(self, doc: Document, images: List[Path], unpack_dir: Path) -> bool:
        """第2步：NO2、O3污染特征分析"""
        try:
            if len(images) < 2:
                logger.warning("not_enough_images", count=len(images))
                return False

            # 分析图片1和2
            prompt = """请分析以下两张图片，给出NO2、O3污染特征分析结果。

要求：
- 不需要任何结论和建议
- 找出关键信息
- 内容丰富、表达流畅易读

第1张图是NO2相关图表，第2张图是O3相关图表。
"""

            # 并发分析图片
            analysis_results = await self._analyze_images_concurrent(images[:2], prompt)

            # 合并分析结果
            combined_analysis = self._combine_analyses(analysis_results)

            # 替换文本
            self._replace_text(doc, "NO2、O3污染特征分析：", combined_analysis)

            return True

        except Exception as e:
            logger.error("step2_failed", error=str(e))
            return False

    async def _step3_spatial_distribution(self, doc: Document, images: List[Path], unpack_dir: Path) -> bool:
        """第3步：O3空间分布特征"""
        try:
            if len(images) < 3:
                logger.warning("not_enough_images", count=len(images))
                return False

            # 分析图片3
            prompt = """请分析第3张图片，给出O3空间分布特征分析结果。

要求：
- 不需要任何建议
- 找出关键信息
- 内容丰富、表达流畅易读
"""

            analysis = await self._analyze_single_image(images[2], prompt)

            # 替换文本
            self._replace_text(doc, "空间分布特征：", analysis)

            return True

        except Exception as e:
            logger.error("step3_failed", error=str(e))
            return False

    async def _step4_vertical_distribution(self, doc: Document, images: List[Path], unpack_dir: Path) -> bool:
        """第4步：臭氧垂直分布分析"""
        try:
            if len(images) < 5:
                logger.warning("not_enough_images", count=len(images))
                return False

            # 分析图片4和5（臭氧雷达图）
            prompt = """请分析以下两张图片，给出臭氧垂直分布分析结果。

说明：
- 第4张图是臭氧雷达316nm消光系数图
- 第5张图是臭氧浓度雷达监测图

要求：
- 不需要任何建议
- 找出关键信息
- 内容丰富、表达流畅易读
"""

            # 并发分析图片
            analysis_results = await self._analyze_images_concurrent(images[3:5], prompt)

            # 合并分析结果
            combined_analysis = self._combine_analyses(analysis_results)

            # 替换文本
            self._replace_text(doc, "臭氧垂直分布分析：", combined_analysis)

            return True

        except Exception as e:
            logger.error("step4_failed", error=str(e))
            return False

    async def _step5_summary(self, doc: Document, images: List[Path], unpack_dir: Path) -> bool:
        """第5步：生成总结"""
        try:
            # 读取文档已有内容
            doc_text = "\n".join([p.text for p in doc.paragraphs])

            prompt = f"""请根据以下文档内容，生成一段总结，综合以上分析内容，给出当天的臭氧污染情况和分析结论。

要求：
- 不需要任何建议
- 找出关键信息
- 内容丰富、表达流畅易读

文档内容：
{doc_text}
"""

            # 调用LLM
            summary = await self._call_llm(prompt)

            # 替换文本
            self._replace_text(doc, "小结：", summary)

            return True

        except Exception as e:
            logger.error("step5_failed", error=str(e))
            return False

    async def _analyze_single_image(self, image_path: Path, prompt: str) -> str:
        """分析单张图片（带重试和限流）"""
        for retry in range(MAX_RETRIES):
            try:
                # 添加请求间隔（限流）
                if retry > 0:
                    await asyncio.sleep(RETRY_DELAY)
                else:
                    await asyncio.sleep(API_REQUEST_INTERVAL)

                # 读取图片
                with open(image_path, 'rb') as f:
                    image_bytes = f.read()
                    base64_data = base64.b64encode(image_bytes).decode('utf-8')

                file_ext = image_path.suffix[1:]

                # 构造data URL
                data_url = f"data:image/{file_ext};base64,{base64_data}"

                # 调用Vision API
                async with httpx.AsyncClient(timeout=120.0) as client:
                    response = await client.post(
                        f"{LLM_BASE_URL}/chat/completions",
                        headers={
                            "Authorization": f"Bearer {LLM_API_KEY}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": LLM_MODEL,
                            "messages": [
                                {
                                    "role": "user",
                                    "content": [
                                        {
                                            "type": "image_url",
                                            "image_url": {"url": data_url}
                                        },
                                        {
                                            "type": "text",
                                            "text": prompt
                                        }
                                    ]
                                }
                            ],
                            "max_tokens": 2000,
                            "temperature": 0.3
                        }
                    )

                    # 检查HTTP状态码
                    if response.status_code == 429:
                        logger.warning("api_rate_limit_image", retry=retry+1, image=image_path.name)
                        await asyncio.sleep(RETRY_DELAY * (retry + 1))  # 指数退避
                        continue

                    response.raise_for_status()
                    result = response.json()
                    analysis = result["choices"][0]["message"]["content"]

                    return analysis

            except httpx.HTTPStatusError as e:
                if e.response.status_code == 429:
                    logger.warning("api_rate_limit_image_http", retry=retry+1, image=image_path.name)
                    await asyncio.sleep(RETRY_DELAY * (retry + 1))
                    continue
                else:
                    logger.error("image_http_error", image=str(image_path), status=e.response.status_code)
                    if retry == MAX_RETRIES - 1:
                        return f"图片分析失败：HTTP {e.response.status_code}"
            except Exception as e:
                logger.error("image_analysis_failed", image=str(image_path), retry=retry+1, error=str(e))
                if retry == MAX_RETRIES - 1:
                    return f"图片分析失败：{str(e)}"

        return f"图片分析失败：超过最大重试次数"

    async def _analyze_images_concurrent(self, images: List[Path], prompt: str) -> List[str]:
        """并发分析多张图片"""
        tasks = [self._analyze_single_image(img, prompt) for img in images]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 处理异常
        final_results = []
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error("image_analysis_error", index=i, error=str(result))
                final_results.append(f"图片{i+1}分析失败：{str(result)}")
            else:
                final_results.append(result)

        return final_results

    def _combine_analyses(self, analyses: List[str]) -> str:
        """合并多个分析结果"""
        return "\n\n".join([f"分析结果{i+1}：\n{analysis}" for i, analysis in enumerate(analyses)])

    async def _call_llm(self, prompt: str) -> str:
        """调用LLM API（带重试和限流）"""
        for retry in range(MAX_RETRIES):
            try:
                # 添加请求间隔（限流）
                if retry > 0:
                    await asyncio.sleep(RETRY_DELAY)
                else:
                    await asyncio.sleep(API_REQUEST_INTERVAL)

                async with httpx.AsyncClient(timeout=120.0) as client:
                    response = await client.post(
                        f"{LLM_BASE_URL}/chat/completions",
                        headers={
                            "Authorization": f"Bearer {LLM_API_KEY}",
                            "Content-Type": "application/json"
                        },
                        json={
                            "model": LLM_MODEL,
                            "messages": [
                                {
                                    "role": "user",
                                    "content": prompt
                                }
                            ],
                            "max_tokens": 2000,
                            "temperature": 0.3
                        }
                    )

                    # 检查HTTP状态码
                    if response.status_code == 429:
                        logger.warning("api_rate_limit", retry=retry+1, max_retries=MAX_RETRIES)
                        await asyncio.sleep(RETRY_DELAY * (retry + 1))  # 指数退避
                        continue

                    response.raise_for_status()
                    result = response.json()
                    content = result["choices"][0]["message"]["content"]

                    return content

            except httpx.HTTPStatusError as e:
                if e.response.status_code == 429:
                    logger.warning("api_rate_limit_http", retry=retry+1)
                    await asyncio.sleep(RETRY_DELAY * (retry + 1))
                    continue
                else:
                    logger.error("llm_http_error", status=e.response.status_code, error=str(e))
                    if retry == MAX_RETRIES - 1:
                        return f"分析失败：HTTP {e.response.status_code}"
            except Exception as e:
                logger.error("llm_call_failed", retry=retry+1, error=str(e))
                if retry == MAX_RETRIES - 1:
                    return f"分析失败：{str(e)}"

        return f"分析失败：超过最大重试次数"

    def _replace_text(self, doc: Document, find_text: str, replace_text: str):
        """在文档中查找并替换文本"""
        replacements = 0

        # 遍历段落
        for paragraph in doc.paragraphs:
            if find_text in paragraph.text:
                # 替换文本
                new_text = paragraph.text.replace(find_text, f"{find_text}\n\n{replace_text}")
                self._replace_paragraph_text(paragraph, new_text)
                replacements += 1
                break  # 只替换第一个匹配项

        # 遍历表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if find_text in paragraph.text:
                            new_text = paragraph.text.replace(find_text, f"{find_text}\n\n{replace_text}")
                            self._replace_paragraph_text(paragraph, new_text)
                            replacements += 1
                            break
                    if replacements > 0:
                        break
                if replacements > 0:
                    break
            if replacements > 0:
                break

        logger.info("text_replaced", find_text=find_text, replacements=replacements)

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """替换段落文本（保留格式）"""
        # 清空段落内容
        for run in paragraph.runs:
            run.text = ""

        # 添加新文本到第一个run（如果存在）
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    async def run(self):
        """运行批量处理"""
        logger.info("batch_processor_starting")

        # 获取待处理报告
        reports = self.get_reports()

        if not reports:
            logger.warning("no_reports_to_process")
            return

        self.stats["total"] = len(reports)
        self.stats["start_time"] = datetime.now().isoformat()

        # 并发处理报告
        semaphore = asyncio.Semaphore(CONCURRENT_TASKS)

        async def process_with_semaphore(report):
            async with semaphore:
                return await self.process_report(report)

        tasks = [process_with_semaphore(report) for report in reports]
        results = await asyncio.gather(*tasks, return_exceptions=True)

        # 统计结果
        for i, result in enumerate(results):
            if isinstance(result, Exception):
                logger.error("processing_error", report=reports[i].name, error=str(result))
                self.stats["failed"] += 1
            elif result:
                self.stats["success"] += 1
            else:
                self.stats["failed"] += 1

        self.stats["end_time"] = datetime.now().isoformat()

        # 保存报告
        self._save_report()

        # 打印总结
        self._print_summary()

    def _print_summary(self):
        """打印处理总结"""
        print("\n" + "="*60)
        print("批量处理完成")
        print("="*60)
        print(f"总报告数：{self.stats['total']}")
        print(f"成功：{self.stats['success']}")
        print(f"失败：{self.stats['failed']}")
        print(f"跳过：{self.stats['skipped']}")
        print(f"\n开始时间：{self.stats['start_time']}")
        print(f"结束时间：{self.stats['end_time']}")
        print(f"\n处理报告：{self.report_file}")
        print(f"进度文件：{self.progress_file}")
        print("="*60)


# ==================== 主程序 ====================

async def main():
    """主函数"""
    processor = BatchProcessor()
    await processor.run()


if __name__ == "__main__":
    # 检查配置
    if not LLM_API_KEY:
        print("错误：未配置 LLM_API_KEY 环境变量")
        print("请设置 QWEN_VL_API_KEY 环境变量")
        sys.exit(1)

    # 检查报告目录
    if not Path(REPORTS_DIR).exists():
        print(f"错误：报告目录不存在：{REPORTS_DIR}")
        print("请修改脚本中的 REPORTS_DIR 配置")
        sys.exit(1)

    # 运行处理器
    asyncio.run(main())
