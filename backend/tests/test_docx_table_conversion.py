"""
测试 docx 表格转换功能

验证点：
1. 表格正确转换为 Markdown 格式
2. 段落和表格的顺序保持一致
3. 标题、加粗、斜体等格式正确转换
4. 合并单元格正确处理
"""
import sys
import os
import io

# 添加项目根目录到路径
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from app.routers.utils_docx import (
    convert_docx_to_markdown,
    _convert_table_to_markdown,
    _paragraph_to_markdown
)


def test_simple_table_conversion():
    """测试简单表格转换"""
    try:
        from docx import Document
    except ImportError:
        print("⚠️  跳过测试: 需要安装 python-docx")
        return

    # 创建测试文档
    doc = Document()
    doc.add_heading('测试报告', 1)
    doc.add_paragraph('这是一段普通文本。')

    # 添加简单表格
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = '城市'
    table.rows[0].cells[1].text = 'PM2.5'
    table.rows[0].cells[2].text = 'AQI'

    table.rows[1].cells[0].text = '广州'
    table.rows[1].cells[1].text = '23'
    table.rows[1].cells[2].text = '92.3'

    table.rows[2].cells[0].text = '深圳'
    table.rows[2].cells[1].text = '18'
    table.rows[2].cells[2].text = '96.1'

    doc.add_paragraph('表格后的段落。')

    # 保存为字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # 转换
    markdown = convert_docx_to_markdown(buffer.read())

    print("\n" + "="*60)
    print("测试1: 简单表格转换")
    print("="*60)
    print(markdown)
    print("="*60)

    # 验证
    assert '# 测试报告' in markdown, "❌ 标题未正确转换"
    assert '这是一段普通文本' in markdown, "❌ 段落未正确转换"
    assert '| 城市 | PM2.5 | AQI |' in markdown, "❌ 表头未正确转换"
    assert '| --- | --- | --- |' in markdown, "❌ 表格分隔线缺失"
    assert '| 广州 | 23 | 92.3 |' in markdown, "❌ 表格数据行未正确转换"
    assert '表格后的段落' in markdown, "❌ 表格后段落缺失"

    # 验证顺序
    title_pos = markdown.index('# 测试报告')
    para1_pos = markdown.index('这是一段普通文本')
    table_pos = markdown.index('| 城市')
    para2_pos = markdown.index('表格后的段落')

    assert title_pos < para1_pos < table_pos < para2_pos, "❌ 文档元素顺序错误"

    print("✅ 测试1通过: 简单表格转换正确")


def test_format_preservation():
    """测试格式保留（加粗、斜体）"""
    try:
        from docx import Document
    except ImportError:
        print("⚠️  跳过测试: 需要安装 python-docx")
        return

    doc = Document()

    # 添加带格式的段落
    para = doc.add_paragraph()
    para.add_run('这是普通文本，')
    para.add_run('这是加粗文本').bold = True
    para.add_run('，')
    para.add_run('这是斜体文本').italic = True
    para.add_run('，')
    run = para.add_run('这是加粗斜体文本')
    run.bold = True
    run.italic = True

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    markdown = convert_docx_to_markdown(buffer.read())

    print("\n" + "="*60)
    print("测试2: 格式保留")
    print("="*60)
    print(markdown)
    print("="*60)

    assert '**这是加粗文本**' in markdown, "❌ 加粗格式未正确转换"
    assert '*这是斜体文本*' in markdown, "❌ 斜体格式未正确转换"
    assert '***这是加粗斜体文本***' in markdown, "❌ 加粗斜体格式未正确转换"

    print("✅ 测试2通过: 格式保留正确")


def test_merged_cells():
    """测试合并单元格处理"""
    try:
        from docx import Document
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except ImportError:
        print("⚠️  跳过测试: 需要安装 python-docx")
        return

    doc = Document()
    table = doc.add_table(rows=3, cols=3)

    # 填充表格
    table.rows[0].cells[0].text = '指标'
    table.rows[0].cells[1].text = '值1'
    table.rows[0].cells[2].text = '值2'

    # 横向合并第二行的前两列（手动设置）
    # 注意：python-docx 的 merge_cells 方法会自动处理
    cell_a = table.rows[1].cells[0]
    cell_b = table.rows[1].cells[1]
    cell_a.text = '合并单元格'

    # 使用 python-docx 的 merge 功能
    # 但由于我们的转换逻辑基于 cell._element 的 ID，
    # 合并后的单元格会被正确处理
    cell_a.merge(cell_b)

    table.rows[1].cells[2].text = '普通单元格'

    table.rows[2].cells[0].text = '数据1'
    table.rows[2].cells[1].text = '数据2'
    table.rows[2].cells[2].text = '数据3'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    markdown = convert_docx_to_markdown(buffer.read())

    print("\n" + "="*60)
    print("测试3: 合并单元格处理")
    print("="*60)
    print(markdown)
    print("="*60)

    # 验证表格存在
    assert '| 指标 | 值1 | 值2 |' in markdown, "❌ 表头未正确转换"
    assert '合并单元格' in markdown, "❌ 合并单元格内容缺失"

    print("✅ 测试3通过: 合并单元格处理正确（内容保留）")


def test_multiple_tables():
    """测试多个表格的文档"""
    try:
        from docx import Document
    except ImportError:
        print("⚠️  跳过测试: 需要安装 python-docx")
        return

    doc = Document()
    doc.add_heading('第一章', 1)
    doc.add_paragraph('第一段文字')

    # 第一个表格
    table1 = doc.add_table(rows=2, cols=2)
    table1.rows[0].cells[0].text = '表1列1'
    table1.rows[0].cells[1].text = '表1列2'
    table1.rows[1].cells[0].text = 'A'
    table1.rows[1].cells[1].text = 'B'

    doc.add_heading('第二章', 2)
    doc.add_paragraph('第二段文字')

    # 第二个表格
    table2 = doc.add_table(rows=2, cols=3)
    table2.rows[0].cells[0].text = '表2列1'
    table2.rows[0].cells[1].text = '表2列2'
    table2.rows[0].cells[2].text = '表2列3'
    table2.rows[1].cells[0].text = 'X'
    table2.rows[1].cells[1].text = 'Y'
    table2.rows[1].cells[2].text = 'Z'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    markdown = convert_docx_to_markdown(buffer.read())

    print("\n" + "="*60)
    print("测试4: 多个表格")
    print("="*60)
    print(markdown)
    print("="*60)

    assert '# 第一章' in markdown, "❌ 第一章标题缺失"
    assert '## 第二章' in markdown, "❌ 第二章标题缺失"
    assert '| 表1列1 | 表1列2 |' in markdown, "❌ 第一个表格缺失"
    assert '| 表2列1 | 表2列2 | 表2列3 |' in markdown, "❌ 第二个表格缺失"

    # 验证顺序
    ch1_pos = markdown.index('# 第一章')
    t1_pos = markdown.index('| 表1列1')
    ch2_pos = markdown.index('## 第二章')
    t2_pos = markdown.index('| 表2列1')

    assert ch1_pos < t1_pos < ch2_pos < t2_pos, "❌ 文档元素顺序错误"

    print("✅ 测试4通过: 多个表格顺序正确")


def test_empty_cells():
    """测试空单元格处理"""
    try:
        from docx import Document
    except ImportError:
        print("⚠️  跳过测试: 需要安装 python-docx")
        return

    doc = Document()
    table = doc.add_table(rows=2, cols=3)

    table.rows[0].cells[0].text = '列1'
    table.rows[0].cells[1].text = '列2'
    table.rows[0].cells[2].text = '列3'

    # 第二行留空单元格
    table.rows[1].cells[0].text = 'A'
    # table.rows[1].cells[1] 留空
    table.rows[1].cells[2].text = 'C'

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    markdown = convert_docx_to_markdown(buffer.read())

    print("\n" + "="*60)
    print("测试5: 空单元格处理")
    print("="*60)
    print(markdown)
    print("="*60)

    # 空单元格应该被填充为一个空格
    lines = markdown.split('\n')
    data_row = next(line for line in lines if line.startswith('| A'))

    assert data_row.count('|') == 4, "❌ 表格列数不正确（应该是4个竖线）"

    print("✅ 测试5通过: 空单元格正确填充")


if __name__ == '__main__':
    print("\n" + "="*60)
    print("DOCX 表格转换测试套件")
    print("="*60)

    try:
        test_simple_table_conversion()
        test_format_preservation()
        test_merged_cells()
        test_multiple_tables()
        test_empty_cells()

        print("\n" + "="*60)
        print("🎉 所有测试通过！")
        print("="*60)

    except AssertionError as e:
        print(f"\n❌ 测试失败: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
    except Exception as e:
        print(f"\n❌ 测试异常: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
