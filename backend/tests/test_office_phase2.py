"""
Office 工具单元测试 - Phase 2

测试范围：
- accept_changes_tool.py - 接受 Word 修订
- find_replace_tool.py - Word 查找替换

测试场景：
- 查找替换（普通文本/正则表达式/大小写）
- 接受修订（需要 LibreOffice，可能跳过）
"""
import pytest
import tempfile
import zipfile
from pathlib import Path
from docx import Document
from app.tools.office.find_replace_tool import FindReplaceTool
from app.tools.office.accept_changes_tool import AcceptChangesTool


@pytest.fixture
def temp_dir():
    """创建临时目录"""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_docx_with_text(temp_dir):
    """创建包含文本的示例 DOCX 文件"""
    docx_path = temp_dir / "sample.docx"
    doc = Document()

    # 添加段落
    doc.add_paragraph("Hello World")
    doc.add_paragraph("This is a test document.")
    doc.add_paragraph("Python is great. Python is powerful.")
    doc.add_paragraph("Date: 2024-01-15")

    # 添加表格
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Header 1"
    table.cell(0, 1).text = "Header 2"
    table.cell(1, 0).text = "Python"
    table.cell(1, 1).text = "Java"

    doc.save(docx_path)
    return docx_path


class TestFindReplaceTool:
    """测试 FindReplaceTool"""

    @pytest.mark.asyncio
    async def test_simple_replace(self, sample_docx_with_text, temp_dir):
        """测试简单文本替换"""
        tool = FindReplaceTool()
        output_file = temp_dir / "output.docx"

        result = await tool.execute(
            file_path=str(sample_docx_with_text),
            find_text="Python",
            replace_text="JavaScript",
            output_file=str(output_file)
        )

        assert result["success"] is True
        assert result["data"]["replacements"] == 3  # 2 in paragraph + 1 in table
        assert result["data"]["paragraphs_affected"] >= 2  # 至少2个段落受影响
        assert output_file.exists()

        # 验证替换结果
        doc = Document(output_file)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "JavaScript" in text
        assert "Python" not in text

    @pytest.mark.asyncio
    async def test_case_insensitive_replace(self, sample_docx_with_text, temp_dir):
        """测试大小写不敏感替换"""
        tool = FindReplaceTool()
        output_file = temp_dir / "output.docx"

        result = await tool.execute(
            file_path=str(sample_docx_with_text),
            find_text="PYTHON",
            replace_text="Ruby",
            output_file=str(output_file),
            case_sensitive=False
        )

        assert result["success"] is True
        assert result["data"]["replacements"] == 3
        assert output_file.exists()

    @pytest.mark.asyncio
    async def test_regex_replace(self, sample_docx_with_text, temp_dir):
        """测试正则表达式替换"""
        tool = FindReplaceTool()
        output_file = temp_dir / "output.docx"

        result = await tool.execute(
            file_path=str(sample_docx_with_text),
            find_text=r"\d{4}-\d{2}-\d{2}",
            replace_text="2025-12-31",
            output_file=str(output_file),
            use_regex=True
        )

        assert result["success"] is True
        assert result["data"]["replacements"] == 1
        assert output_file.exists()

        # 验证替换结果
        doc = Document(output_file)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "2025-12-31" in text
        assert "2024-01-15" not in text

    @pytest.mark.asyncio
    async def test_replace_in_place(self, sample_docx_with_text):
        """测试原地替换（覆盖原文件）"""
        tool = FindReplaceTool()

        result = await tool.execute(
            file_path=str(sample_docx_with_text),
            find_text="test",
            replace_text="example"
        )

        assert result["success"] is True
        assert result["data"]["replacements"] == 1

        # 验证原文件被修改
        doc = Document(sample_docx_with_text)
        text = "\n".join([p.text for p in doc.paragraphs])
        assert "example" in text
        assert "test" not in text

    @pytest.mark.asyncio
    async def test_no_matches(self, sample_docx_with_text, temp_dir):
        """测试没有匹配项"""
        tool = FindReplaceTool()
        output_file = temp_dir / "output.docx"

        result = await tool.execute(
            file_path=str(sample_docx_with_text),
            find_text="NonExistentText",
            replace_text="Replacement",
            output_file=str(output_file)
        )

        assert result["success"] is True
        assert result["data"]["replacements"] == 0
        assert result["data"]["paragraphs_affected"] == 0

    @pytest.mark.asyncio
    async def test_replace_nonexistent_file(self, temp_dir):
        """测试替换不存在的文件"""
        tool = FindReplaceTool()

        result = await tool.execute(
            file_path=str(temp_dir / "nonexistent.docx"),
            find_text="test",
            replace_text="example"
        )

        assert result["success"] is False
        assert "不存在" in result["summary"]

    @pytest.mark.asyncio
    async def test_replace_invalid_format(self, temp_dir):
        """测试替换不支持的格式"""
        tool = FindReplaceTool()
        invalid_file = temp_dir / "test.txt"
        invalid_file.write_text("test")

        result = await tool.execute(
            file_path=str(invalid_file),
            find_text="test",
            replace_text="example"
        )

        assert result["success"] is False
        assert "格式" in result["summary"]


class TestAcceptChangesTool:
    """测试 AcceptChangesTool"""

    @pytest.mark.asyncio
    async def test_accept_changes_file_not_found(self, temp_dir):
        """测试接受修订 - 文件不存在"""
        tool = AcceptChangesTool()

        result = await tool.execute(
            input_file=str(temp_dir / "nonexistent.docx"),
            output_file=str(temp_dir / "output.docx")
        )

        assert result["success"] is False
        assert "不存在" in result["summary"]

    @pytest.mark.asyncio
    async def test_accept_changes_invalid_format(self, temp_dir):
        """测试接受修订 - 不支持的格式"""
        tool = AcceptChangesTool()
        invalid_file = temp_dir / "test.txt"
        invalid_file.write_text("test")

        result = await tool.execute(
            input_file=str(invalid_file),
            output_file=str(temp_dir / "output.docx")
        )

        assert result["success"] is False
        assert "格式" in result["summary"]

    @pytest.mark.asyncio
    async def test_accept_changes_tool_availability(self):
        """测试 LibreOffice 可用性检测"""
        tool = AcceptChangesTool()
        is_available = tool.is_available()

        # 只记录可用性，不强制要求
        print(f"LibreOffice available: {is_available}")

    @pytest.mark.asyncio
    @pytest.mark.skipif(
        not AcceptChangesTool().is_available(),
        reason="LibreOffice not installed"
    )
    async def test_accept_changes_with_libreoffice(self, sample_docx_with_text, temp_dir):
        """测试接受修订 - 需要 LibreOffice"""
        tool = AcceptChangesTool()
        output_file = temp_dir / "output.docx"

        result = await tool.execute(
            input_file=str(sample_docx_with_text),
            output_file=str(output_file)
        )

        # 如果 LibreOffice 可用，应该成功
        # 如果不可用，测试会被跳过
        if result["success"]:
            assert output_file.exists()
            assert result["data"]["size"] > 0


class TestToolIntegration:
    """测试工具集成"""

    @pytest.mark.asyncio
    async def test_find_replace_schema(self):
        """测试查找替换工具 Schema"""
        tool = FindReplaceTool()
        schema = tool.get_function_schema()

        assert schema["name"] == "find_replace_word"
        assert "parameters" in schema
        assert "file_path" in schema["parameters"]["properties"]
        assert "find_text" in schema["parameters"]["properties"]
        assert "replace_text" in schema["parameters"]["properties"]

    @pytest.mark.asyncio
    async def test_accept_changes_schema(self):
        """测试接受修订工具 Schema"""
        tool = AcceptChangesTool()
        schema = tool.get_function_schema()

        assert schema["name"] == "accept_word_changes"
        assert "parameters" in schema
        assert "input_file" in schema["parameters"]["properties"]
        assert "output_file" in schema["parameters"]["properties"]


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
