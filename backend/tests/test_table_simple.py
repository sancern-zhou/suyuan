"""
简化版测试 - 验证表格转换的核心功能
避免emoji字符，解决Windows cmd编码问题
"""
import sys
import os

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

def main():
    print("=" * 60)
    print("DOCX Table Conversion Test")
    print("=" * 60)

    # 测试1: 检查依赖
    try:
        from docx import Document
        print("[OK] python-docx imported successfully")
    except ImportError:
        print("[SKIP] python-docx not installed")
        print("Please run: pip install python-docx")
        return 1

    # 测试2: 导入转换函数
    try:
        from app.routers.utils_docx import convert_docx_to_markdown
        print("[OK] convert_docx_to_markdown imported")
    except ImportError as e:
        print(f"[ERROR] Cannot import conversion function: {e}")
        return 1

    # 测试3: 创建测试docx
    import io
    doc = Document()
    doc.add_heading('Test Report', 1)
    doc.add_paragraph('This is a paragraph.')

    # 添加表格
    table = doc.add_table(rows=3, cols=3)
    table.rows[0].cells[0].text = 'City'
    table.rows[0].cells[1].text = 'PM2.5'
    table.rows[0].cells[2].text = 'AQI'

    table.rows[1].cells[0].text = 'Guangzhou'
    table.rows[1].cells[1].text = '23'
    table.rows[1].cells[2].text = '92.3'

    table.rows[2].cells[0].text = 'Shenzhen'
    table.rows[2].cells[1].text = '18'
    table.rows[2].cells[2].text = '96.1'

    doc.add_paragraph('Paragraph after table.')

    # 保存为字节
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    docx_bytes = buffer.read()

    print(f"[OK] Test docx created, size: {len(docx_bytes)} bytes")

    # 测试4: 转换
    try:
        markdown = convert_docx_to_markdown(docx_bytes)
        print("[OK] Conversion completed")
    except Exception as e:
        print(f"[ERROR] Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        return 1

    # 测试5: 验证结果
    print("\n" + "=" * 60)
    print("Markdown Output:")
    print("=" * 60)
    print(markdown)
    print("=" * 60)

    # 验证关键元素
    checks = [
        ('# Test Report' in markdown, 'Heading conversion'),
        ('This is a paragraph' in markdown, 'Paragraph conversion'),
        ('| City | PM2.5 | AQI |' in markdown, 'Table header'),
        ('| --- | --- | --- |' in markdown, 'Table separator'),
        ('| Guangzhou | 23 | 92.3 |' in markdown, 'Table data row'),
        ('Paragraph after table' in markdown, 'Paragraph after table'),
    ]

    all_passed = True
    for passed, description in checks:
        status = "[PASS]" if passed else "[FAIL]"
        print(f"{status} {description}")
        if not passed:
            all_passed = False

    # 验证顺序
    try:
        title_pos = markdown.index('# Test Report')
        para1_pos = markdown.index('This is a paragraph')
        table_pos = markdown.index('| City')
        para2_pos = markdown.index('Paragraph after table')

        order_correct = title_pos < para1_pos < table_pos < para2_pos
        status = "[PASS]" if order_correct else "[FAIL]"
        print(f"{status} Element order preservation")
        if not order_correct:
            all_passed = False
    except ValueError as e:
        print(f"[FAIL] Element order check: {e}")
        all_passed = False

    print("\n" + "=" * 60)
    if all_passed:
        print("All tests PASSED!")
        print("=" * 60)
        return 0
    else:
        print("Some tests FAILED!")
        print("=" * 60)
        return 1


if __name__ == '__main__':
    sys.exit(main())
