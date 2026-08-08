from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from app.services.report.government_docx_style import finalize_government_docx
from app.services.quarto_report_renderer import QuartoReportRenderer


def _font_attr(run, attr: str) -> str | None:
    r_fonts = run._element.get_or_add_rPr().rFonts
    return r_fonts.get(qn(f"w:{attr}"))


def _add_bottom_border_to_title_style(doc: Document) -> None:
    p_pr = doc.styles["Title"]._element.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:color"), "4F81BD")


def _paragraph_has_border(paragraph) -> bool:
    return paragraph._p.get_or_add_pPr().find(qn("w:pBdr")) is not None


def _style_has_border(style) -> bool:
    p_pr = style._element.get_or_add_pPr()
    return p_pr.find(qn("w:pBdr")) is not None


def _save_sample_docx(path: Path) -> None:
    doc = Document()
    doc.add_paragraph("测试报告", style="Title")
    doc.add_heading("章标题", level=1)
    doc.add_heading("一级节标题", level=2)
    doc.add_heading("二级节标题", level=3)
    doc.add_heading("三级节标题", level=4)
    doc.add_heading("四级节标题", level=5)
    body = doc.add_paragraph("这是正文段落，用于检查中文和西文字体。")
    body.add_run(" Latin text")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "表头"
    table.cell(0, 1).text = "Header"
    table.cell(1, 0).text = "内容"
    table.cell(1, 1).text = "Value"
    doc.save(path)


def test_finalize_government_docx_applies_requested_body_toc_and_table_styles(tmp_path):
    docx_path = tmp_path / "report.docx"
    _save_sample_docx(docx_path)

    result = finalize_government_docx(docx_path)

    assert result["toc_inserted"] is True
    doc = Document(docx_path)

    toc_title = next(paragraph for paragraph in doc.paragraphs if paragraph.text.replace(" ", "") == "目录")
    assert toc_title.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert toc_title.runs[0].font.size.pt == 12
    assert _font_attr(toc_title.runs[0], "eastAsia") == "宋体"
    assert _font_attr(toc_title.runs[0], "ascii") == "Times New Roman"

    body = next(paragraph for paragraph in doc.paragraphs if paragraph.text.startswith("这是正文段落"))
    assert body.paragraph_format.line_spacing == 1.0
    assert body.paragraph_format.line_spacing_rule == WD_LINE_SPACING.SINGLE
    assert body.runs[0].font.size.pt == 14
    assert _font_attr(body.runs[0], "eastAsia") == "宋体"
    assert _font_attr(body.runs[0], "ascii") == "Times New Roman"

    table = doc.tables[0]
    assert table.alignment is not None
    header_run = table.cell(0, 0).paragraphs[0].runs[0]
    body_run = table.cell(1, 0).paragraphs[0].runs[0]
    assert header_run.bold is True
    assert header_run.font.size.pt == 12
    assert body_run.font.size.pt == 12
    assert _font_attr(header_run, "eastAsia") == "宋体"
    assert _font_attr(body_run, "ascii") == "Times New Roman"


def test_finalize_government_docx_removes_title_blue_bottom_borders(tmp_path):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    _add_bottom_border_to_title_style(doc)
    doc.add_paragraph("测试报告", style="Title")
    doc.add_heading("章标题", level=1)
    doc.save(docx_path)

    finalize_government_docx(docx_path)

    doc = Document(docx_path)
    assert not _style_has_border(doc.styles["Title"])
    title_paragraph = next(paragraph for paragraph in doc.paragraphs if paragraph.text == "测试报告")
    toc_paragraph = next(paragraph for paragraph in doc.paragraphs if paragraph.text.replace(" ", "") == "目录")
    assert not _paragraph_has_border(title_paragraph)
    assert not _paragraph_has_border(toc_paragraph)


def test_finalize_government_docx_applies_requested_heading_levels(tmp_path):
    docx_path = tmp_path / "report.docx"
    _save_sample_docx(docx_path)

    finalize_government_docx(docx_path)

    doc = Document(docx_path)
    headings = {
        paragraph.style.name: paragraph
        for paragraph in doc.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    }

    expected = {
        "Heading 1": (22, 14, 14, True),
        "Heading 2": (16, 7, 7, False),
        "Heading 3": (15, 7, 7, False),
        "Heading 4": (14, 7, 7, False),
        "Heading 5": (12, 3.5, 3.5, False),
    }

    for style_name, (size_pt, before_pt, after_pt, page_break) in expected.items():
        paragraph = headings[style_name]
        assert paragraph.paragraph_format.first_line_indent.pt == 0
        assert paragraph.paragraph_format.space_before.pt == before_pt
        assert paragraph.paragraph_format.space_after.pt == after_pt
        assert paragraph.paragraph_format.page_break_before is page_break
        assert paragraph.runs[0].font.size.pt == size_pt
        assert _font_attr(paragraph.runs[0], "eastAsia") == "Times New Roman"
        assert _font_attr(paragraph.runs[0], "ascii") == "Times New Roman"
        assert _font_attr(paragraph.runs[-1], "eastAsia") == "黑体"


def test_finalize_government_docx_strips_bare_numeric_heading_prefixes(tmp_path):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("测试报告", style="Title")
    doc.add_heading("1 项目名称", level=2)
    doc.add_heading("0 业务架构概述", level=2)
    doc.save(docx_path)

    finalize_government_docx(docx_path)

    doc = Document(docx_path)
    headings = [
        paragraph.text.strip()
        for paragraph in doc.paragraphs
        if paragraph.style and paragraph.style.name == "Heading 2"
    ]
    assert headings == ["（一）项目名称", "（二）业务架构概述"]


def test_prepare_docx_qmd_disables_quarto_docx_toc_and_numbering(tmp_path):
    report_dir = tmp_path / "report"
    report_dir.mkdir()
    qmd_path = report_dir / "report.qmd"
    qmd_path.write_text(
        """---
title: 测试报告
toc: true
number-sections: true
format:
  html:
    toc: true
  docx:
    toc: true
    number-sections: true
---

# 第一章
""",
        encoding="utf-8",
    )
    renderer = QuartoReportRenderer(report_dir.parent)

    prepared = renderer._prepare_docx_qmd(report_dir, qmd_path)

    prepared_text = prepared.read_text(encoding="utf-8")
    assert "toc: false" in prepared_text
    assert "number-sections: false" in prepared_text
    assert "docx:\n    toc: false\n    number-sections: false" in prepared_text
    assert prepared != qmd_path
