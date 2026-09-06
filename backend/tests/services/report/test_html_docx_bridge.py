import shutil
import subprocess

from docx import Document
from docx.oxml.ns import qn
import pytest

from app.services.report.html_docx_bridge import RenderedHtmlReport
from app.services.report.government_docx_style import finalize_government_docx


def convert(tmp_path, content):
    source = tmp_path / "report.html"
    source.write_text(content, encoding="utf-8")
    report = RenderedHtmlReport(source)
    prepared = tmp_path / "prepared.html"
    report.prepare(prepared)
    output = tmp_path / "report.docx"
    if not shutil.which("quarto"):
        pytest.skip("Quarto is required for the actual HTML/DOCX reader contract")
    subprocess.run(
        ["quarto", "pandoc", str(prepared), "-f", "html", "-t", "docx", "-o", str(output)],
        check=True,
        capture_output=True,
    )
    finalize_government_docx(output)
    report.apply_styles(output)
    return report, output


def test_html_tables_css_cascade_merged_cells_and_content_order(tmp_path):
    report, output = convert(
        tmp_path,
        """<html><head><style>
    th { background: #1f4e79; color: white }
    .risk { background-color: #ffc7ce; color: #9c0006; font-weight:700 }
    td.risk { color: rgb(0,128,0) !important }
    </style></head><body><main id="quarto-document-content">
    <header id="title-block-header"><h1>Report title</h1></header>
    <h1>Section</h1><p>Before</p>
    <div><TABLE><THEAD><TR><TH colspan="2">Header</TH></TR></THEAD>
    <tbody><tr><td rowspan="2" class="risk" style="color:blue">High</td><td>A</td></tr>
    <tr><td><span class="risk">Medium</span></td></tr></tbody></TABLE></div>
    <p>After <span class="risk">Warning</span></p>
    </main></body></html>""",
    )
    doc = Document(output)
    assert len(doc.tables) == 1
    table = doc.tables[0]
    assert len(table.rows) == 3
    assert table.cell(0, 0)._tc is table.cell(0, 1)._tc
    assert table.cell(1, 0)._tc is table.cell(2, 0)._tc
    assert table.cell(0, 0)._tc.xpath("./w:tcPr/w:shd/@w:fill") == ["1F4E79"]
    assert table.cell(1, 0)._tc.xpath(".//w:rPr/w:color/@w:val") == ["008000"]
    assert doc.element.body.xpath('.//w:r[w:t="Medium"]/w:rPr/w:shd/@w:fill') == ["FFC7CE"]
    assert doc.element.body.xpath('.//w:r[w:t="Warning"]/w:rPr/w:color/@w:val') == ["9C0006"]
    assert report.validate(output)["cells"] == 4


def test_nested_table_and_rowspan_zero(tmp_path):
    report, output = convert(
        tmp_path,
        """<body><table><tbody>
    <tr><td rowspan="0">Shared</td><td><table><tr><td>Inner</td></tr></table></td></tr>
    <tr><td>Last</td></tr></tbody></table></body>""",
    )
    assert report.validate(output)["tables"] == 2


def test_no_false_positive_for_examples_comments_or_hidden_content(tmp_path):
    source = tmp_path / "report.html"
    source.write_text(
        """<body><!-- <table><tr><td>comment</td></tr></table> -->
    <pre><code>&lt;table&gt;&lt;/table&gt;</code></pre>
    <script>const example = '<table></table>'</script>
    <template><table></table></template>
    <div style="display:none"><table><tr><td>Hidden</td></tr></table></div>
    <p>Text</p></body>""",
        encoding="utf-8",
    )
    assert not RenderedHtmlReport(source).tables


def test_validation_rejects_flattened_tables_and_removed_colors(tmp_path):
    report, output = convert(
        tmp_path, '<body><table><tr><td style="background:red">A</td></tr></table></body>'
    )
    doc = Document(output)
    shading = doc.tables[0].cell(0, 0)._tc.find(".//" + qn("w:shd"))
    shading.set(qn("w:fill"), "FFFFFF")
    doc.save(output)
    with pytest.raises(ValueError, match="background"):
        report.validate(output)
    with pytest.raises(ValueError, match="background"):
        RenderedHtmlReport(tmp_path / "report.html").validate(output)
    doc = Document()
    doc.add_paragraph("A")
    doc.save(output)
    with pytest.raises(ValueError, match="count mismatch"):
        report.validate(output)


def test_local_stylesheet_and_independent_color_validation(tmp_path):
    (tmp_path / "style.css").write_text(
        ".risk { color: #9c0006; background: #ffc7ce }", encoding="utf-8"
    )
    report, output = convert(
        tmp_path,
        '<head><link rel="stylesheet" href="style.css"></head><body><table><tr><td><span class="risk">High</span></td></tr></table></body>',
    )
    assert RenderedHtmlReport(tmp_path / "report.html").validate(output)["tables"] == 1
    doc = Document(output)
    for color in doc.element.body.xpath(".//w:color"):
        color.getparent().remove(color)
    doc.save(output)
    with pytest.raises(ValueError, match="text color"):
        RenderedHtmlReport(tmp_path / "report.html").validate(output)


def test_quarto_resolves_raw_markdown_included_and_generated_tables(tmp_path):
    from app.services.quarto_report_renderer import QuartoReportRenderer

    report_dir = tmp_path / "report"
    report_dir.mkdir()
    (report_dir / "generated.lua").write_text(
        """function Para(p)
  if pandoc.utils.stringify(p) == "GENERATED" then
    return pandoc.RawBlock("html", "<table><tr><td>Generated</td></tr></table>")
  end
end""",
        encoding="utf-8",
    )
    (report_dir / "included.qmd").write_text(
        "<div><TABLE><tr><th>Included</th></tr><tr><td>Value</td></tr></TABLE></div>",
        encoding="utf-8",
    )
    (report_dir / "report.qmd").write_text(
        """---
title: Report
format: html
filters:
  - generated.lua
---

# Content

```html
<table><tr><td>Example</td></tr></table>
```

`<table>inline example</table>`

<!-- <table><tr><td>Comment</td></tr></table> -->

{{< include included.qmd >}}

| Markdown | Column |
|----------|--------|
| A        | B      |

```{=html}
<table><tr><td>Raw block</td></tr></table>
```

GENERATED
""",
        encoding="utf-8",
    )
    renderer = QuartoReportRenderer(tmp_path)
    result = renderer.render_docx("report")
    doc = Document(result)
    assert len(doc.tables) == 4
    assert [t.cell(0, 0).text for t in doc.tables] == [
        "Included",
        "Markdown",
        "Raw block",
        "Generated",
    ]


def test_css_custom_properties_and_fallback(tmp_path):
    report, output = convert(
        tmp_path,
        """<style>
    :root { --risk: #ffc7ce; --ink: #9c0006 }
    .risk { background:var(--risk); color:var(--missing, var(--ink)) }
    </style><body><table><tr><td class="risk">High</td></tr></table></body>""",
    )
    assert report.validate(output)["tables"] == 1
    tc = Document(output).tables[0].cell(0, 0)._tc
    assert tc.xpath("./w:tcPr/w:shd/@w:fill") == ["FFC7CE"]
    assert tc.xpath(".//w:rPr/w:color/@w:val") == ["9C0006"]


@pytest.mark.asyncio
async def test_report_package_validator_rejects_structural_loss(tmp_path, monkeypatch):
    from app.services.quarto_report_renderer import QuartoReportRenderer
    from app.tools.report.report_package import tool

    report_dir = tmp_path / "report"
    report_dir.mkdir()
    (report_dir / "report.qmd").write_text("<table><tr><td>A</td></tr></table>", encoding="utf-8")
    (report_dir / "report.html").write_text(
        "<body><table><tr><td>A</td></tr></table></body>", encoding="utf-8"
    )
    doc = Document()
    doc.add_paragraph("A")
    doc.save(report_dir / "report.docx")
    monkeypatch.setattr(tool, "quarto_report_renderer", QuartoReportRenderer(tmp_path))
    result = await tool.ValidateReportPackageTool().execute("report", require_docx=True)
    assert not result["success"]
    assert any("table count mismatch" in error for error in result["data"]["errors"])


def test_failed_bridge_keeps_existing_docx(tmp_path, monkeypatch):
    from app.services.quarto_report_renderer import QuartoReportRenderer, ReportRenderError

    source = tmp_path / "report.html"
    source.write_text("<body><table><tr><td>A</td></tr></table></body>", encoding="utf-8")
    qmd = tmp_path / "report.qmd"
    qmd.write_text("# Test", encoding="utf-8")
    existing = tmp_path / "report.docx"
    existing.write_bytes(b"existing artifact")
    renderer = QuartoReportRenderer(tmp_path)

    def fake_pandoc(cwd, args):
        doc = Document()
        doc.add_paragraph("A")
        doc.save(cwd / args[args.index("--output") + 1])

    monkeypatch.setattr(renderer, "_run_quarto", fake_pandoc)
    with pytest.raises(ReportRenderError, match="table count mismatch"):
        renderer._render_docx_from_resolved_html(tmp_path, qmd, RenderedHtmlReport(source))
    assert existing.read_bytes() == b"existing artifact"
    assert not list(tmp_path.glob("report_docx_*"))
