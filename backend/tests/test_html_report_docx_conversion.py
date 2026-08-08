from pathlib import Path

from docx import Document

from app.services.report.government_docx_style import convert_html_report_to_government_docx


def test_convert_quarto_html_table_with_thead_tbody(tmp_path: Path):
    html_path = tmp_path / "report.html"
    docx_path = tmp_path / "report.docx"
    html_path.write_text(
        """
        <!doctype html>
        <html>
          <head><title>测试报告</title></head>
          <body>
            <main class="container">
              <section class="section">
                <h2>一、审核范围</h2>
                <table class="caption-top table">
                  <thead>
                    <tr class="header"><th>项目</th><th>说明</th></tr>
                  </thead>
                  <tbody>
                    <tr class="odd"><td>工单总数</td><td>173 条</td></tr>
                    <tr class="even"><td>审核方法</td><td>规则筛查</td></tr>
                  </tbody>
                </table>
              </section>
            </main>
          </body>
        </html>
        """,
        encoding="utf-8",
    )

    result = convert_html_report_to_government_docx(html_path, docx_path)

    assert result["tables"] == 1
    doc = Document(str(docx_path))
    assert len(doc.tables) == 1
    assert doc.tables[0].cell(0, 0).text == "项目"
    assert doc.tables[0].cell(1, 0).text == "工单总数"
