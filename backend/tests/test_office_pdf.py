"""
测试 PDF 转换功能

运行此脚本检查 PDF 转换是否正常工作
"""
import asyncio
import sys
from pathlib import Path
from datetime import datetime

# 添加父目录到路径
sys.path.insert(0, str(Path(__file__).parent.parent))

from app.tools.office.soffice import run_soffice, get_soffice_path


async def test_libreoffice():
    """测试 LibreOffice 是否可用"""
    print("=" * 50)
    print("测试 LibreOffice 是否安装")
    print("=" * 50)

    # 首先显示查找的路径
    soffice_path = get_soffice_path()
    print(f"LibreOffice 路径: {soffice_path}")

    # 检查文件是否存在
    from pathlib import Path
    if not Path(soffice_path).exists() and soffice_path == "soffice":
        print("\n正在搜索 LibreOffice 安装位置...")
        import os

        # Windows 搜索
        if os.name == "nt":
            program_files = [
                os.environ.get("ProgramFiles", r"C:\Program Files"),
                os.environ.get("ProgramFiles(x86)", r"C:\Program Files (x86)")
            ]

            found_locations = []
            for base in program_files:
                try:
                    libreoffice_dir = Path(base) / "LibreOffice"
                    if libreoffice_dir.exists():
                        # 搜索 soffice.exe
                        for exe_file in libreoffice_dir.rglob("soffice.exe"):
                            if "program" in str(exe_file.parent).lower():
                                found_locations.append(str(exe_file))
                except Exception:
                    pass

            if found_locations:
                print("\n找到以下 LibreOffice 安装:")
                for loc in found_locations:
                    print(f"  - {loc}")
                print(f"\n将使用: {found_locations[0]}")

                # 更新全局路径缓存
                import app.tools.office.soffice as soffice_module
                soffice_module._soffice_path = found_locations[0]
                soffice_path = found_locations[0]
            else:
                print("\n未找到 LibreOffice 安装!")
                print("\n请安装 LibreOffice:")
                print("  1. 下载: https://www.libreoffice.org/download/download/")
                print("  2. 安装到默认路径 (推荐)")
                print("  3. 重新运行此测试")

    try:
        result = run_soffice(["--version"])
        if result.returncode == 0:
            # LibreOffice 版本信息可能在 stdout 或 stderr 中
            version = (result.stdout or "").strip() or (result.stderr or "").strip()
            if version:
                print(f"\nLibreOffice 已安装: {version}")
                return True
            else:
                print("\nLibreOffice 运行成功，但没有版本信息")
                return True
        else:
            print(f"LibreOffice 返回错误 (代码 {result.returncode})")
            stderr = (result.stderr or "").strip()
            stdout = (result.stdout or "").strip()
            if stderr:
                print(f"错误信息: {stderr}")
            if stdout:
                print(f"输出信息: {stdout}")
            return False
    except FileNotFoundError:
        print(f"\n错误: 找不到 LibreOffice 可执行文件")
        return False
    except Exception as e:
        print(f"LibreOffice 不可用: {e}")
        import traceback
        traceback.print_exc()
        return False


async def test_pdf_converter():
    """测试 PDF 转换器"""
    print("\n" + "=" * 50)
    print("测试 PDF 转换器")
    print("=" * 50)

    try:
        from app.services.pdf_converter import pdf_converter
        print(f"PDF 转换器已加载")
        print(f"输出目录: {pdf_converter.output_dir}")
        print(f"输出目录存在: {pdf_converter.output_dir.exists()}")
        return True
    except ImportError as e:
        print(f"PDF 转换器导入失败: {e}")
        print("请确保已安装 pypdf: pip install pypdf")
        return False
    except Exception as e:
        print(f"PDF 转换器初始化失败: {e}")
        return False


async def test_docx_conversion():
    """测试实际的 DOCX 转 PDF"""
    print("\n" + "=" * 50)
    print("测试 DOCX 转 PDF")
    print("=" * 50)

    # 查找或创建测试文件
    test_file = None
    possible_paths = [
        Path(__file__).parent.parent.parent / "backend_data" / "test.docx",
        Path(__file__).parent / "fixtures" / "sample.docx",
    ]

    for path in possible_paths:
        if path.exists():
            test_file = path
            break

    # 如果没有找到，尝试创建一个简单的测试文档
    if not test_file:
        try:
            from docx import Document

            # 创建测试文档
            test_file = Path(__file__).parent / "fixtures" / "sample.docx"
            test_file.parent.mkdir(exist_ok=True)

            doc = Document()
            doc.add_heading('Office PDF 预览测试文档', 0)
            doc.add_paragraph('这是一个测试文档，用于验证 PDF 转换功能。')
            doc.add_paragraph('LibreOffice 已成功安装并可以运行！')
            doc.add_paragraph('测试时间：' + str(datetime.now()))

            doc.save(str(test_file))
            print(f"已创建测试文档: {test_file}")

        except ImportError:
            print("未找到测试 DOCX 文件，且 python-docx 未安装")
            print("请安装 python-docx: pip install python-docx")
            print("或手动将一个 .docx 文件放在以下位置之一：")
            for path in possible_paths:
                print(f"  - {path}")
            return False
        except Exception as e:
            print(f"创建测试文档失败: {e}")
            return False

    print(f"使用测试文件: {test_file}")

    try:
        from app.services.pdf_converter import pdf_converter
        result = await pdf_converter.convert_to_pdf(str(test_file))

        print(f"转换成功!")
        print(f"  PDF ID: {result['pdf_id']}")
        print(f"  PDF URL: {result['pdf_url']}")
        print(f"  页数: {result['pages']}")
        print(f"  大小: {result['size']} 字节")

        # 验证文件存在
        pdf_path = pdf_converter.get_pdf_path(result['pdf_id'])
        if pdf_path.exists():
            print(f"  PDF 文件已创建: {pdf_path}")
        else:
            print(f"  警告: PDF 文件不存在")
            return False

        return True

    except Exception as e:
        print(f"转换失败: {e}")
        import traceback
        traceback.print_exc()
        return False


async def main():
    """主测试函数"""
    print("\nOffice PDF 预览功能测试\n")

    # 1. 测试 LibreOffice
    libreoffice_ok = await test_libreoffice()

    # 2. 测试 PDF 转换器
    converter_ok = await test_pdf_converter()

    # 3. 测试实际转换
    conversion_ok = False
    if libreoffice_ok and converter_ok:
        conversion_ok = await test_docx_conversion()

    # 总结
    print("\n" + "=" * 50)
    print("测试结果总结")
    print("=" * 50)
    print(f"LibreOffice: {'✓' if libreoffice_ok else '✗'}")
    print(f"PDF 转换器: {'✓' if converter_ok else '✗'}")
    print(f"DOCX 转 PDF: {'✓' if conversion_ok else '✗'}")

    if libreoffice_ok and converter_ok and conversion_ok:
        print("\n所有测试通过! PDF 预览功能应该正常工作。")
        print("\n如果前端仍然没有显示预览，请检查:")
        print("1. 后端是否已重启")
        print("2. 浏览器控制台是否有错误")
        print("3. 网络请求是否成功 (检查 /api/office/pdf/xxx)")
    else:
        print("\n存在问题，请根据上述提示修复。")

    return libreoffice_ok and converter_ok and conversion_ok


if __name__ == "__main__":
    success = asyncio.run(main())
    sys.exit(0 if success else 1)
