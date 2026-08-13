from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

from app.services.report.government_docx_style import finalize_government_docx


def test_finalize_government_docx_uses_fourth_size_spaced_toc_title(tmp_path):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("测试报告")
    doc.add_heading("一、正文", level=1)
    doc.add_paragraph("内容")
    doc.save(docx_path)

    finalize_government_docx(docx_path)

    updated = Document(docx_path)
    toc_title = next(paragraph for paragraph in updated.paragraphs if paragraph.text.strip() == "目  录")
    assert toc_title.runs[0].font.size.pt == 14


def test_finalize_government_docx_formats_image_caption_with_number(tmp_path):
    image_path = tmp_path / "chart.png"
    Image.new("RGB", (80, 40), "white").save(image_path)

    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("测试报告")
    doc.add_heading("一、正文", level=1)
    doc.add_paragraph().add_run().add_picture(str(image_path))
    doc.add_paragraph("AQI日变化趋势")
    doc.save(docx_path)

    finalize_government_docx(docx_path)

    updated = Document(docx_path)
    caption = next(paragraph for paragraph in updated.paragraphs if paragraph.text.startswith("图1 "))
    assert caption.text == "图1 AQI日变化趋势"
    assert caption.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert all(run.italic for run in caption.runs if run.text)


def test_finalize_government_docx_applies_heading_numbering(tmp_path):
    docx_path = tmp_path / "report.docx"
    doc = Document()
    doc.add_paragraph("测试报告")
    doc.add_heading("总体情况", level=1)
    doc.add_paragraph("内容")
    doc.add_heading("处理效率", level=2)
    doc.add_heading("重点问题", level=3)
    doc.save(docx_path)

    result = finalize_government_docx(docx_path, add_toc=False)

    updated = Document(docx_path)
    headings = [
        paragraph.text
        for paragraph in updated.paragraphs
        if paragraph.style and paragraph.style.name.startswith("Heading")
    ]
    assert result["heading_numbers"] == 3
    assert headings == ["一、总体情况", "（一）处理效率", "1. 重点问题"]
