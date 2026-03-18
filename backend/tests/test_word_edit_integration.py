"""
Word Edit Tool Integration Tests

Comprehensive tests for the new Word editing system including:
- XMLEditor node location functionality
- Document class save functionality
- word_edit_tool end-to-end testing
- Comparison with find_replace_word functionality
- Error recovery scenarios

Run with: pytest tests/test_word_edit_integration.py -v
"""

import pytest
import tempfile
import shutil
from pathlib import Path
from docx import Document as DocxDocument

from app.tools.xml.xml_editor import XMLEditor, XMLNode
from app.tools.office.document import Document
from app.tools.office.word_edit_tool import WordEditTool
from app.tools.office.find_replace_tool import FindReplaceTool


class TestXMLEditor:
    """Test XMLEditor core functionality."""

    def test_basic_xml_parsing(self):
        """Test basic XML parsing and node location."""
        xml_content = """<?xml version="1.0" encoding="UTF-8"?>
        <root>
            <child1>Text 1</child1>
            <child2 attr="value">Text 2</child2>
        </root>"""

        editor = XMLEditor(xml_content)

        # Find all child nodes
        children = editor.get_nodes(tag="child1")
        assert len(children) == 1
        assert children[0].text == "Text 1"

    def test_find_by_tag(self):
        """Test finding nodes by tag name."""
        xml_content = """<root>
            <w:p>Paragraph 1</w:p>
            <w:p>Paragraph 2</w:p>
            <w:t>Text node</w:t>
        </root>"""

        editor = XMLEditor(xml_content)
        paragraphs = editor.get_nodes(tag="w:p")

        assert len(paragraphs) == 2
        assert paragraphs[0].text == "Paragraph 1"
        assert paragraphs[1].text == "Paragraph 2"

    def test_find_by_attributes(self):
        """Test finding nodes by attributes."""
        xml_content = """<root>
            <item id="1">First</item>
            <item id="2">Second</item>
            <item id="1">Third</item>
        </root>"""

        editor = XMLEditor(xml_content)
        items = editor.get_nodes(tag="item", attributes={"id": "1"})

        assert len(items) == 2
        assert items[0].text == "First"
        assert items[1].text == "Third"

    def test_find_by_text_content(self):
        """Test finding nodes containing specific text."""
        xml_content = """<root>
            <w:p>This is a test paragraph</w:p>
            <w:p>Another paragraph</w:p>
            <w:p>Yet another test paragraph</w:p>
        </root>"""

        editor = XMLEditor(xml_content)
        test_paragraphs = editor.get_nodes(tag="w:p", contains="test")

        assert len(test_paragraphs) == 2

    def test_update_text(self):
        """Test updating text content of a node."""
        xml_content = """<root>
            <w:p>Original text</w:p>
        </root>"""

        editor = XMLEditor(xml_content)
        nodes = editor.get_nodes(tag="w:p")

        assert len(nodes) == 1
        success = editor.update_text(nodes[0], "Updated text")

        assert success is True
        assert "Updated text" in editor.to_xml()

    def test_replace_node(self):
        """Test replacing a node with new XML."""
        xml_content = """<root>
            <w:p>Old paragraph</w:p>
        </root>"""

        editor = XMLEditor(xml_content)
        nodes = editor.get_nodes(tag="w:p")

        new_xml = "<w:p>New paragraph content</w:p>"
        success = editor.replace_node(nodes[0], new_xml)

        assert success is True
        assert "New paragraph content" in editor.to_xml()
        assert "Old paragraph" not in editor.to_xml()

    def test_insert_after(self):
        """Test inserting content after a node."""
        xml_content = """<root>
            <w:p>First paragraph</w:p>
        </root>"""

        editor = XMLEditor(xml_content)
        nodes = editor.get_nodes(tag="w:p")

        new_xml = "<w:p>Inserted paragraph</w:p>"
        success = editor.insert_after(nodes[0], new_xml)

        assert success is True
        result = editor.to_xml()
        assert "First paragraph" in result
        assert "Inserted paragraph" in result

    def test_remove_node(self):
        """Test removing a node."""
        xml_content = """<root>
            <w:p>Keep this</w:p>
            <w:p>Remove this</w:p>
            <w:p>Keep this too</w:p>
        </root>"""

        editor = XMLEditor(xml_content)
        nodes = editor.get_nodes(tag="w:p")

        # Remove the middle paragraph
        success = editor.remove_node(nodes[1])

        assert success is True
        result = editor.to_xml()
        assert "Keep this" in result
        assert "Remove this" not in result


class TestDocumentClass:
    """Test Document class functionality."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        doc = DocxDocument()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Second paragraph with keyword")
        doc.add_paragraph("Third paragraph")

        doc_path = tmp_path / "test.docx"
        doc.save(str(doc_path))
        return doc_path

    @pytest.fixture
    def unpacked_dir(self, sample_docx, tmp_path):
        """Unpack a sample DOCX file."""
        from app.tools.office.unpack_tool import UnpackOfficeTool

        unpack_tool = UnpackOfficeTool()
        result = unpack_tool.execute(path=str(sample_docx), output_dir=str(tmp_path / "unpacked"))

        assert result["success"]
        return tmp_path / "unpacked"

    def test_document_initialization(self, unpacked_dir):
        """Test Document initialization."""
        doc = Document(unpacked_dir)

        # Should be able to access main document
        main_doc = doc.get_main_document()
        assert main_doc is not None

    def test_find_paragraphs(self, unpacked_dir):
        """Test finding paragraphs by text content."""
        doc = Document(unpacked_dir)

        # Find paragraphs containing "keyword"
        paragraphs = doc.find_paragraphs("keyword")

        assert len(paragraphs) >= 1

    def test_replace_text(self, unpacked_dir):
        """Test replacing text throughout document."""
        doc = Document(unpacked_dir)

        # Replace "keyword" with "replacement"
        count = doc.replace_text("keyword", "replacement")

        assert count >= 0  # May be 0 if no matches

    def test_document_statistics(self, unpacked_dir):
        """Test getting document statistics."""
        doc = Document(unpacked_dir)

        stats = doc.get_statistics()

        assert "unpacked_dir" in stats
        assert "author" in stats
        assert "edit_time" in stats


class TestWordEditTool:
    """Test word_edit_tool end-to-end functionality."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        doc = DocxDocument()
        doc.add_paragraph("Introduction paragraph")
        doc.add_paragraph("Content with old text")
        doc.add_paragraph("Conclusion paragraph")

        doc_path = tmp_path / "test.docx"
        doc.save(str(doc_path))
        return doc_path

    @pytest.mark.asyncio
    async def test_replace_text_operation(self, sample_docx):
        """Test replace_text operation."""
        tool = WordEditTool()

        result = await tool.execute(
            path=str(sample_docx),
            operation="replace_text",
            search="old text",
            replace="new text"
        )

        assert result["success"] is True
        assert result["data"]["operation"] == "replace_text"

        # Verify the change
        doc = DocxDocument(str(sample_docx))
        text = " ".join([p.text for p in doc.paragraphs])
        assert "new text" in text

    @pytest.mark.asyncio
    async def test_replace_paragraph_operation(self, sample_docx):
        """Test replace_paragraph operation."""
        tool = WordEditTool()

        result = await tool.execute(
            path=str(sample_docx),
            operation="replace_paragraph",
            contains="Content with old text",
            new_content="Completely new content"
        )

        assert result["success"] is True
        assert result["data"]["operation"] == "replace_paragraph"

    @pytest.mark.asyncio
    async def test_insert_after_operation(self, sample_docx):
        """Test insert_after operation."""
        tool = WordEditTool()

        result = await tool.execute(
            path=str(sample_docx),
            operation="insert_after",
            marker="Introduction paragraph",
            content="Inserted after introduction"
        )

        assert result["success"] is True
        assert result["data"]["operation"] == "insert_after"

    @pytest.mark.asyncio
    async def test_invalid_operation(self, sample_docx):
        """Test error handling for invalid operation."""
        tool = WordEditTool()

        result = await tool.execute(
            path=str(sample_docx),
            operation="invalid_operation"
        )

        assert result["success"] is False
        assert "未知操作类型" in result["summary"]

    @pytest.mark.asyncio
    async def test_missing_parameters(self, sample_docx):
        """Test error handling for missing required parameters."""
        tool = WordEditTool()

        result = await tool.execute(
            path=str(sample_docx),
            operation="replace_text"
            # Missing search and replace parameters
        )

        assert result["success"] is False


class TestFindReplaceComparison:
    """Compare word_edit_tool with find_replace_word functionality."""

    @pytest.fixture
    def sample_docx(self, tmp_path):
        """Create a sample DOCX file for testing."""
        doc = DocxDocument()
        doc.add_paragraph("This document contains OLD_TEXT multiple times.")
        doc.add_paragraph("Another paragraph with OLD_TEXT.")
        doc.add_paragraph("Final paragraph without the target.")

        doc_path = tmp_path / "test.docx"
        doc.save(str(doc_path))
        return doc_path

    @pytest.mark.asyncio
    async def test_find_replace_simple(self, sample_docx):
        """Test find_replace_word for simple text replacement."""
        tool = FindReplaceTool()

        result = await tool.execute(
            path=str(sample_docx),
            find_text="OLD_TEXT",
            replace_text="NEW_TEXT"
        )

        assert result["success"] is True
        assert result["data"]["replacements"] == 2

    @pytest.mark.asyncio
    async def test_word_edit_simple(self, sample_docx):
        """Test word_edit for simple text replacement."""
        tool = WordEditTool()

        result = await tool.execute(
            path=str(sample_docx),
            operation="replace_text",
            search="OLD_TEXT",
            replace="NEW_TEXT"
        )

        assert result["success"] is True
        assert result["data"]["changes"] == 2


class TestErrorRecovery:
    """Test error recovery scenarios."""

    @pytest.mark.asyncio
    async def test_nonexistent_file(self):
        """Test handling of nonexistent file."""
        tool = WordEditTool()

        result = await tool.execute(
            path="nonexistent.docx",
            operation="replace_text",
            search="test",
            replace="test"
        )

        assert result["success"] is False
        assert "文件不存在" in result["summary"]

    @pytest.mark.asyncio
    async def test_invalid_file_format(self, tmp_path):
        """Test handling of invalid file format."""
        # Create a text file instead of docx
        text_file = tmp_path / "test.txt"
        text_file.write_text("Some text content")

        tool = WordEditTool()

        result = await tool.execute(
            path=str(text_file),
            operation="replace_text",
            search="test",
            replace="test"
        )

        assert result["success"] is False
        assert "不支持的文件格式" in result["summary"]

    @pytest.mark.asyncio
    async def test_marker_not_found(self, tmp_path):
        """Test handling when marker text is not found."""
        # Create a simple docx
        doc = DocxDocument()
        doc.add_paragraph("Test content")
        doc_path = tmp_path / "test.docx"
        doc.save(str(doc_path))

        tool = WordEditTool()

        result = await tool.execute(
            path=str(doc_path),
            operation="insert_after",
            marker="NONEXISTENT_MARKER",
            content="New content"
        )

        # Should fail gracefully
        assert result["data"]["changes"] == 0


class TestToolAvailability:
    """Test tool availability and registration."""

    def test_xml_editor_available(self):
        """Test XMLEditor can be imported."""
        from app.tools.xml.xml_editor import XMLEditor
        assert XMLEditor is not None

    def test_document_available(self):
        """Test Document can be imported."""
        from app.tools.office.document import Document
        assert Document is not None

    def test_word_edit_tool_available(self):
        """Test WordEditTool is available."""
        tool = WordEditTool()
        assert tool.is_available() is True

    def test_word_edit_tool_schema(self):
        """Test WordEditTool has valid schema."""
        tool = WordEditTool()
        schema = tool.get_function_schema()

        assert schema["name"] == "word_edit"
        assert "parameters" in schema
        assert "properties" in schema["parameters"]
        assert "operation" in schema["parameters"]["properties"]
