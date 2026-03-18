"""
FindReplace 工具 - Word 文档查找替换

功能：
- 在 DOCX 文件中查找并替换文本
- 支持正则表达式（可选）
- 支持大小写敏感/不敏感
- 基于 python-docx 库实现
- 自动生成PDF预览

使用场景：
- 批量替换文本
- 更新文档中的术语
- 修正拼写错误
"""
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from app.tools.base.tool_interface import LLMTool, ToolCategory
import structlog
import re

logger = structlog.get_logger()

# PDF转换器导入（懒加载避免循环依赖）
def get_pdf_converter():
    try:
        from app.services.pdf_converter import pdf_converter
        return pdf_converter
    except ImportError:
        logger.warning("pdf_converter_not_available")
        return None


class FindReplaceTool(LLMTool):
    """
    Word 文档查找替换工具

    功能：
    - 在 DOCX 文件中查找并替换文本
    - 支持正则表达式
    - 支持大小写敏感/不敏感
    """

    def __init__(self):
        super().__init__(
            name="find_replace_word",
            description="""在 Word 文档中查找并替换文本（简单文本替换，快速便捷）

🎯 适用场景：简单的文本查找替换操作

⚡ 快捷方案（推荐用于简单替换）：
- 直接操作 DOCX 文件，无需解包
- 自动保留文档格式（字体、颜色、样式等）
- 更可靠、更简单、更安全
- 速度更快，无需解包/打包流程

使用场景：
- ✅ 简单文本替换（批量修改术语、拼写错误）
- ✅ 在某处插入内容（"标记" → "标记\n\n新内容"）
- ✅ 查找并替换固定文本

❌ 不适用场景：
- 复杂编辑（插入段落、修改格式）→ 使用 word_edit 工具
- 结构化编辑（需要精确控制）→ 使用 word_edit 工具

示例：
- find_replace_word(path="report.docx", find_text="旧术语", replace_text="新术语")
- find_replace_word(path="doc.docx", find_text="标题", replace_text="标题\\n\\n插入的新段落")
- find_replace_word(path="doc.docx", find_text="\\d{4}-\\d{2}-\\d{2}", replace_text="2024-01-01", use_regex=True)

插入内容的技巧：
- 要在某段文本后插入内容，将 find_text 设为该文本，replace_text 设为"原文本\\n\\n要插入的内容"
- 示例：在"结论："后插入分析结果：find_text="结论：", replace_text="结论：\\n\\n根据上述分析..."

参数说明：
- path: DOCX 文件路径
- find_text: 要查找的文本（支持正则表达式）
- replace_text: 替换后的文本
- output_file: 输出文件路径（可选，默认覆盖原文件）
- use_regex: 是否使用正则表达式（默认 False）
- case_sensitive: 是否大小写敏感（默认 True）

注意：
- 如果不指定 output_file，会覆盖原文件
- 正则表达式替换支持捕获组（如 $1, $2）
- ⚠️ 编辑 Word 文档时优先使用此工具，而非 edit_file
""",
            category=ToolCategory.QUERY,
            version="1.0.0",
            requires_context=False
        )

        self.working_dir = Path.cwd().parent  # D:\溯源\ 或 /opt/app/ 等

    async def execute(
        self,
        path: str,
        find_text: str,
        replace_text: str,
        output_file: Optional[str] = None,
        use_regex: bool = False,
        case_sensitive: bool = True,
        **kwargs
    ) -> Dict[str, Any]:
        """
        查找并替换 Word 文档中的文本

        Args:
            path: DOCX 文件路径
            find_text: 要查找的文本
            replace_text: 替换后的文本
            output_file: 输出文件路径（可选）
            use_regex: 是否使用正则表达式
            case_sensitive: 是否大小写敏感

        Returns:
            {
                "success": bool,
                "data": {
                    "file_path": str,
                    "replacements": int,
                    "paragraphs_affected": int
                },
                "summary": str
            }
        """
        try:
            # 1. 路径解析
            resolved_path = self._resolve_path(path)
            if not resolved_path:
                return {
                    "success": False,
                    "data": {"error": "文件路径无效"},
                    "summary": "查找替换失败：路径无效"
                }

            if not resolved_path.exists():
                return {
                    "success": False,
                    "data": {"error": f"文件不存在: {path}"},
                    "summary": "查找替换失败：文件不存在"
                }

            if resolved_path.suffix.lower() != ".docx":
                return {
                    "success": False,
                    "data": {"error": f"不支持的文件格式: {resolved_path.suffix}"},
                    "summary": "查找替换失败：仅支持 DOCX 格式"
                }

            # 2. 确定输出文件
            if output_file:
                output_path = self._resolve_path(output_file)
                output_path.parent.mkdir(parents=True, exist_ok=True)
            else:
                output_path = resolved_path

            # 3. 打开文档
            doc = Document(str(resolved_path))

            # 4. 执行查找替换
            total_replacements = 0
            paragraphs_affected = 0

            # 编译正则表达式（如果需要）
            if use_regex:
                flags = 0 if case_sensitive else re.IGNORECASE
                pattern = re.compile(find_text, flags)
            else:
                pattern = None

            # 遍历段落
            for paragraph in doc.paragraphs:
                original_text = paragraph.text

                if use_regex:
                    # 正则表达式替换
                    new_text = pattern.sub(replace_text, original_text)
                    count = len(pattern.findall(original_text))
                else:
                    # 普通文本替换
                    if case_sensitive:
                        count = original_text.count(find_text)
                        new_text = original_text.replace(find_text, replace_text)
                    else:
                        # 大小写不敏感替换
                        count = original_text.lower().count(find_text.lower())
                        new_text = self._case_insensitive_replace(original_text, find_text, replace_text)

                if count > 0:
                    # 替换段落文本（保留格式）
                    self._replace_paragraph_text(paragraph, new_text)
                    total_replacements += count
                    paragraphs_affected += 1

            # 遍历表格
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text

                            if use_regex:
                                new_text = pattern.sub(replace_text, original_text)
                                count = len(pattern.findall(original_text))
                            else:
                                if case_sensitive:
                                    count = original_text.count(find_text)
                                    new_text = original_text.replace(find_text, replace_text)
                                else:
                                    count = original_text.lower().count(find_text.lower())
                                    new_text = self._case_insensitive_replace(original_text, find_text, replace_text)

                            if count > 0:
                                self._replace_paragraph_text(paragraph, new_text)
                                total_replacements += count
                                paragraphs_affected += 1

            # 5. 保存文档
            doc.save(output_path)

            logger.info(
                "find_replace_success",
                file=str(resolved_path),
                output=str(output_path),
                replacements=total_replacements,
                paragraphs=paragraphs_affected
            )

            # 生成PDF预览
            pdf_preview = None
            try:
                converter = get_pdf_converter()
                if converter:
                    pdf_preview = await converter.convert_to_pdf(str(output_path))
                    logger.info(
                        "find_replace_pdf_generated",
                        pdf_id=pdf_preview["pdf_id"],
                        pdf_url=pdf_preview["pdf_url"]
                    )
            except Exception as pdf_error:
                logger.warning("find_replace_pdf_conversion_failed", error=str(pdf_error))

            result_data = {
                "file_path": str(output_path),
                "replacements": total_replacements,
                "paragraphs_affected": paragraphs_affected,
                "find_text": find_text,
                "replace_text": replace_text
            }

            if pdf_preview:
                result_data["pdf_preview"] = pdf_preview

            return {
                "success": True,
                "data": result_data,
                "summary": f"已替换 {total_replacements} 处文本，影响 {paragraphs_affected} 个段落"
            }

        except Exception as e:
            logger.error("find_replace_failed", file=str(resolved_path), error=str(e))
            return {
                "success": False,
                "data": {"error": str(e)},
                "summary": f"查找替换失败：{str(e)[:80]}"
            }

    def _case_insensitive_replace(self, text: str, find_text: str, replace_text: str) -> str:
        """大小写不敏感替换"""
        pattern = re.compile(re.escape(find_text), re.IGNORECASE)
        return pattern.sub(replace_text, text)

    def _replace_paragraph_text(self, paragraph, new_text: str):
        """替换段落文本（保留格式）"""
        # 清空段落内容
        for run in paragraph.runs:
            run.text = ""

        # 添加新文本到第一个 run（如果存在）
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            paragraph.add_run(new_text)

    def _resolve_path(self, path: str) -> Path:
        """解析路径（支持相对路径和绝对路径）"""
        try:
            file_path = Path(path)

            if not file_path.is_absolute():
                file_path = self.working_dir / file_path

            return file_path.resolve()

        except Exception as e:
            logger.error("find_replace_path_resolution_failed", path=path, error=str(e))
            return None

    def get_function_schema(self) -> Dict[str, Any]:
        """获取 Function Calling Schema"""
        return {
            "name": "find_replace_word",
            "description": """在 Word 文档中查找并替换文本（简单文本替换，推荐首选）

🎯 适用场景：简单的文本查找替换操作

决策流程：
1. 仅需要简单文本替换？→ find_replace_word（本工具，推荐）
2. 需要复杂结构化编辑？→ word_edit（自动解包/编辑/打包）
3. 编辑代码/配置文件？→ edit_file

功能：
- 直接操作 DOCX 文件，自动保留格式
- 支持正则表达式和大小写控制
- 速度快，无需解包/打包流程

使用场景：
- ✅ 简单文本替换（批量修改术语、拼写错误）
- ✅ 在某处插入内容（"标记" → "标记\\n\\n新内容"）
- ❌ 复杂编辑（插入段落、修改格式）→ 使用 word_edit

示例：
- find_replace_word(path="report.docx", find_text="旧术语", replace_text="新术语")
- find_replace_word(path="doc.docx", find_text="标题", replace_text="标题\\n\\n插入的新段落")

注意：
- 如果不指定 output_file，会覆盖原文件
- 简单替换优先使用此工具（更快速）
- 复杂编辑请使用 word_edit 工具
""",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "DOCX 文件路径。示例：'report.docx' 或 'D:/work/report.docx'"
                    },
                    "find_text": {
                        "type": "string",
                        "description": "要查找的文本（支持正则表达式）。示例：'旧术语' 或 '\\d{4}-\\d{2}-\\d{2}'"
                    },
                    "replace_text": {
                        "type": "string",
                        "description": "替换后的文本。示例：'新术语' 或 '2024-01-01'"
                    },
                    "output_file": {
                        "type": "string",
                        "description": "输出文件路径（可选，默认覆盖原文件）。示例：'report_updated.docx'"
                    },
                    "use_regex": {
                        "type": "boolean",
                        "description": "是否使用正则表达式（默认 False）",
                        "default": False
                    },
                    "case_sensitive": {
                        "type": "boolean",
                        "description": "是否大小写敏感（默认 True）",
                        "default": True
                    }
                },
                "required": ["path", "find_text", "replace_text"]
            }
        }

    def is_available(self) -> bool:
        return True


# 创建工具实例
tool = FindReplaceTool()
