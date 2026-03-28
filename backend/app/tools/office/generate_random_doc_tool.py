"""
Generate Random DOC Tool - Create random content Word documents

Provides functionality to generate random content Word documents with various formats and content types.
"""

from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import structlog
import random
import uuid
from datetime import datetime

from app.tools.base.tool_interface import LLMTool, ToolCategory

logger = structlog.get_logger()


class GenerateRandomDocTool(LLMTool):
    """
    Tool for generating random content Word documents.
    
    This tool creates new .docx files with random content including:
    - Random paragraphs of text
    - Random reports with structured sections
    - Various content types (news, reports, articles, etc.)
    """

    def __init__(self):
        super().__init__(
            name="generate_random_doc",
            description="""生成随机内容的DOC文件

功能：
- 创建新的Word文档 (.docx)
- 生成随机文本内容（段落、标题、列表等）
- 支持多种内容类型：报告、新闻、文章等
- 可自定义内容长度和格式

参数说明：
- output_path: 输出文件路径（必填）
- content_type: 内容类型（可选，默认：'report'）
  - 'report': 随机报告
  - 'news': 随机新闻
  - 'article': 随机文章
  - 'mixed': 混合内容
- paragraphs: 段落数量（可选，默认：5）
- title: 文档标题（可选，默认：随机标题）

示例：
- generate_random_doc(output_path='/root/随机文件.docx')
- generate_random_doc(output_path='/root/随机文件.docx', content_type='report', paragraphs=10)
- generate_random_doc(output_path='/root/随机文件.docx', title='我的随机报告', paragraphs=8)
""",
            category=ToolCategory.UTILITY,
            version="1.0.0",
            requires_context=False
        )

        self.working_dir = Path.cwd().parent

    async def execute(
        self,
        output_path: str,
        content_type: str = "report",
        paragraphs: int = 5,
        title: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        Execute random document generation.

        Args:
            output_path: Output file path
            content_type: Content type (report/news/article/mixed)
            paragraphs: Number of paragraphs to generate
            title: Document title (optional)

        Returns:
            Success/failure response with file path
        """
        try:
            # Resolve output path
            resolved_path = self._resolve_path(output_path)
            if not resolved_path:
                return {
                    "success": False,
                    "data": {"error": f"无效路径: {output_path}"},
                    "summary": "生成失败：无效路径"
                }

            # Create parent directories if needed
            resolved_path.parent.mkdir(parents=True, exist_ok=True)

            # Generate document
            doc = Document()

            # Add title
            if not title:
                title = self._generate_random_title(content_type)
            
            title_para = doc.add_paragraph()
            title_run = title_para.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(16)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add subtitle with date
            subtitle = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
            subtitle_para = doc.add_paragraph(subtitle)
            subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle_para.runs[0].font.size = Pt(10)

            # Add empty paragraph
            doc.add_paragraph()

            # Generate content based on type
            if content_type == "report":
                self._generate_report_content(doc, paragraphs)
            elif content_type == "news":
                self._generate_news_content(doc, paragraphs)
            elif content_type == "article":
                self._generate_article_content(doc, paragraphs)
            else:
                self._generate_mixed_content(doc, paragraphs)

            # Add footer
            doc.add_paragraph()
            footer = doc.add_paragraph("* 此文档由随机内容生成工具创建 *")
            footer.runs[0].font.size = Pt(9)
            footer.runs[0].font.italic = True

            # Save document
            doc.save(str(resolved_path))

            return {
                "success": True,
                "data": {
                    "file_path": str(resolved_path),
                    "content_type": content_type,
                    "paragraphs": paragraphs,
                    "title": title
                },
                "summary": f"成功生成随机文档: {title} ({paragraphs}段落)"
            }

        except Exception as e:
            logger.error("generate_random_doc_failed", error=str(e))
            return {
                "success": False,
                "data": {"error": str(e)},
                "summary": f"生成失败: {str(e)[:80]}"
            }

    def _generate_random_title(self, content_type: str) -> str:
        """Generate random title based on content type."""
        titles = {
            "report": [
                "月度工作总结报告",
                "项目进度分析报告",
                "市场调研报告",
                "技术评估报告",
                "年度总结报告"
            ],
            "news": [
                "今日重要新闻摘要",
                "国际时事动态",
                "科技前沿报道",
                "财经资讯速递",
                "社会热点追踪"
            ],
            "article": [
                "深度分析文章",
                "专题研究报告",
                "观点与评论",
                "行业观察",
                "技术探讨"
            ],
            "mixed": [
                "综合信息文档",
                "混合内容报告",
                "多主题文档",
                "信息汇总"
            ]
        }
        
        type_titles = titles.get(content_type, titles["mixed"])
        return random.choice(type_titles)

    def _generate_random_sentence(self, length: int = 20) -> str:
        """Generate random Chinese sentence."""
        words = [
            "技术", "发展", "创新", "研究", "分析", "报告", "总结", "评估",
            "项目", "进度", "质量", "效率", "成本", "风险", "机遇", "挑战",
            "市场", "竞争", "需求", "供给", "趋势", "变化", "影响", "结果",
            "数据", "信息", "系统", "平台", "工具", "方法", "流程", "标准",
            "团队", "合作", "沟通", "协调", "管理", "执行", "监督", "反馈",
            "用户", "客户", "服务", "体验", "满意度", "价值", "效益", "成果"
        ]
        
        sentence = "".join(random.choices(words, k=random.randint(5, 15)))
        return sentence + "。"

    def _generate_paragraph(self, sentence_count: int = 3) -> str:
        """Generate random paragraph."""
        sentences = [self._generate_random_sentence() for _ in range(sentence_count)]
        return "".join(sentences)

    def _generate_report_content(self, doc: Document, paragraphs: int):
        """Generate report-style content."""
        sections = [
            "一、工作概述",
            "二、主要成果",
            "三、存在问题",
            "四、改进措施",
            "五、下阶段计划"
        ]
        
        for i, section in enumerate(sections[:paragraphs]):
            # Add section heading
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(section)
            heading_run.bold = True
            heading_run.font.size = Pt(12)
            
            # Add content paragraphs
            for _ in range(random.randint(2, 4)):
                doc.add_paragraph(self._generate_paragraph())
            
            doc.add_paragraph()

    def _generate_news_content(self, doc: Document, paragraphs: int):
        """Generate news-style content."""
        news_items = [
            "[科技]",
            "[财经]",
            "[国际]",
            "[社会]",
            "[文化]"
        ]
        
        for i in range(min(paragraphs, len(news_items))):
            # Add news item heading
            heading_para = doc.add_paragraph()
            heading_run = heading_para.add_run(news_items[i] + " " + self._generate_random_sentence(10))
            heading_run.bold = True
            
            # Add news content
            doc.add_paragraph(self._generate_paragraph(random.randint(3, 5)))
            doc.add_paragraph()

    def _generate_article_content(self, doc: Document, paragraphs: int):
        """Generate article-style content."""
        # Add introduction
        intro_para = doc.add_paragraph()
        intro_run = intro_para.add_run("引言")
        intro_run.bold = True
        intro_run.font.size = Pt(12)
        doc.add_paragraph(self._generate_paragraph(4))
        
        # Add body sections
        for i in range(1, paragraphs):
            section_para = doc.add_paragraph()
            section_run = section_para.add_run(f"第{i}部分")
            section_run.bold = True
            section_run.font.size = Pt(12)
            doc.add_paragraph(self._generate_paragraph(random.randint(3, 5)))
        
        # Add conclusion
        conclusion_para = doc.add_paragraph()
        conclusion_run = conclusion_para.add_run("结论")
        conclusion_run.bold = True
        conclusion_run.font.size = Pt(12)
        doc.add_paragraph(self._generate_paragraph(3))

    def _generate_mixed_content(self, doc: Document, paragraphs: int):
        """Generate mixed content."""
        content_types = ["report", "news", "article"]
        
        for i in range(paragraphs):
            if i % 3 == 0:
                # Add section heading
                heading_para = doc.add_paragraph()
                heading_run = heading_para.add_run(f"部分 {i//3 + 1}")
                heading_run.bold = True
                heading_run.font.size = Pt(12)
            
            doc.add_paragraph(self._generate_paragraph(random.randint(2, 4)))

    def _resolve_path(self, path: str) -> Optional[Path]:
        """Resolve file path."""
        try:
            file_path = Path(path)
            if not file_path.is_absolute():
                file_path = self.working_dir / file_path
            return file_path.resolve()
        except Exception as e:
            logger.error("path_resolution_failed", path=path, error=str(e))
            return None

    def get_function_schema(self) -> Dict[str, Any]:
        """Get Function Calling Schema."""
        return {
            "name": "generate_random_doc",
            "description": "生成随机内容的DOC文件",
            "parameters": {
                "type": "object",
                "properties": {
                    "output_path": {
                        "type": "string",
                        "description": "输出文件路径（必填）。示例：'/root/随机文件.docx'"
                    },
                    "content_type": {
                        "type": "string",
                        "enum": ["report", "news", "article", "mixed"],
                        "description": "内容类型（可选，默认：report）",
                        "default": "report"
                    },
                    "paragraphs": {
                        "type": "integer",
                        "description": "段落数量（可选，默认：5）",
                        "default": 5
                    },
                    "title": {
                        "type": "string",
                        "description": "文档标题（可选，默认：随机标题）"
                    }
                },
                "required": ["output_path"]
            }
        }

    def is_available(self) -> bool:
        return True


# Create tool instance
tool = GenerateRandomDocTool()
