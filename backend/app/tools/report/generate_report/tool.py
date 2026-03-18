"""
生成DOCX报告工具

基于模板和数据生成DOCX格式报告
"""

from typing import Dict, Any, Optional, List
from pathlib import Path
import structlog
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import copy

from app.tools.base.tool_interface import LLMTool, ToolCategory

logger = structlog.get_logger()


class GenerateReportTool(LLMTool):
    """
    生成DOCX报告工具

    功能：
    - 使用python-docx生成报告
    - 支持数据替换
    - 支持表格生成
    - 支持文本替换
    """

    def __init__(self):
        super().__init__(
            name="generate_report",
            description="使用python-docx生成DOCX报告。参数: template_path(str), time_range(dict), data_ids(list), output_path(str), data_matches(dict, 可选)",
            category=ToolCategory.ANALYSIS,
            version="1.0.0",
            requires_context=False
        )

        self.function_schema = {
            "name": "generate_report",
            "description": "基于模板和数据生成DOCX报告文档",
            "parameters": {
                "type": "object",
                "properties": {
                    "template_path": {
                        "type": "string",
                        "description": "DOCX模板文件路径（可选，如果不提供则创建新文档）"
                    },
                    "time_range": {
                        "type": "object",
                        "description": "报告时间范围",
                        "properties": {
                            "start_time": {"type": "string", "description": "开始时间（YYYY-MM-DD HH:MM:SS）"},
                            "end_time": {"type": "string", "description": "结束时间（YYYY-MM-DD HH:MM:SS）"}
                        }
                    },
                    "data_ids": {
                        "type": "array",
                        "description": "数据ID列表（从问数模式获取的data_id）",
                        "items": {"type": "string"}
                    },
                    "output_path": {
                        "type": "string",
                        "description": "输出报告文件路径"
                    },
                    "data_matches": {
                        "type": "object",
                        "description": "数据匹配规则（可选）"
                    },
                    "report_title": {
                        "type": "string",
                        "description": "报告标题（可选）"
                    }
                },
                "required": ["time_range", "output_path"]
            }
        }

    async def execute(
        self,
        time_range: Dict[str, str],
        output_path: str,
        template_path: Optional[str] = None,
        data_ids: Optional[List[str]] = None,
        data_matches: Optional[Dict[str, Any]] = None,
        report_title: Optional[str] = None,
        **kwargs
    ) -> Dict[str, Any]:
        """
        执行报告生成

        Args:
            time_range: 时间范围
            output_path: 输出路径
            template_path: 模板路径（可选）
            data_ids: 数据ID列表
            data_matches: 数据匹配规则
            report_title: 报告标题

        Returns:
            生成结果
        """
        try:
            # 创建或加载文档
            if template_path and Path(template_path).exists():
                doc = Document(str(template_path))
                logger.info("loaded_template", path=template_path)
            else:
                doc = Document()
                # 添加默认标题
                if report_title:
                    title = doc.add_heading(report_title, 0)
                    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                logger.info("created_new_document")

            # 替换时间范围占位符
            self._replace_time_placeholders(doc, time_range)

            # 替换数据占位符
            if data_matches:
                self._replace_data_placeholders(doc, data_matches)

            # 如果有data_ids，添加数据表格
            if data_ids:
                await self._add_data_tables(doc, data_ids, time_range)

            # 添加生成时间
            self._add_generation_info(doc, time_range)

            # 确保输出目录存在
            output_file = Path(output_path)
            output_file.parent.mkdir(parents=True, exist_ok=True)

            # 保存文档
            doc.save(str(output_file))

            result = {
                "success": True,
                "data": {
                    "output_path": output_path,
                    "file_name": output_file.name,
                    "file_size": output_file.stat().st_size if output_file.exists() else 0,
                    "time_range": time_range,
                    "data_ids": data_ids or []
                },
                "summary": f"报告已生成: {output_path}"
            }

            logger.info(
                "report_generated",
                path=output_path,
                size=result["data"]["file_size"]
            )

            return result

        except Exception as e:
            logger.error(
                "generate_report_failed",
                error=str(e),
                exc_info=True
            )
            return {
                "success": False,
                "data": {"error": str(e)},
                "summary": f"生成失败: {str(e)}"
            }

    def _replace_time_placeholders(self, doc: Document, time_range: Dict[str, str]):
        """替换时间占位符"""
        time_replacements = {
            "{start_time}": time_range.get("start_time", ""),
            "{end_time}": time_range.get("end_time", ""),
            "{{start_time}}": time_range.get("start_time", ""),
            "{{end_time}}": time_range.get("end_time", ""),
            "${start_time}": time_range.get("start_time", ""),
            "${end_time}": time_range.get("end_time", ""),
        }

        # 替换段落中的时间占位符
        for para in doc.paragraphs:
            for placeholder, value in time_replacements.items():
                if placeholder in para.text:
                    self._replace_text_in_paragraph(para, placeholder, value)

        # 替换表格中的时间占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in time_replacements.items():
                            if placeholder in para.text:
                                self._replace_text_in_paragraph(para, placeholder, value)

    def _replace_data_placeholders(self, doc: Document, data_matches: Dict[str, Any]):
        """替换数据占位符"""
        # data_matches 格式: {"placeholder_name": "replacement_value"}
        for para in doc.paragraphs:
            for placeholder, value in data_matches.items():
                # 支持多种占位符格式
                placeholders = [
                    f"{{{placeholder}}}",
                    f"{{{{{placeholder}}}}}",
                    f"${{{placeholder}}}",
                    f"#{placeholder}#",
                    f"[{placeholder}]",
                ]
                for ph in placeholders:
                    if ph in para.text:
                        self._replace_text_in_paragraph(para, ph, str(value))

        # 替换表格中的占位符
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for placeholder, value in data_matches.items():
                            placeholders = [
                                f"{{{placeholder}}}",
                                f"{{{{{placeholder}}}}}",
                                f"${{{placeholder}}}",
                            ]
                            for ph in placeholders:
                                if ph in para.text:
                                    self._replace_text_in_paragraph(para, ph, str(value))

    def _replace_text_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """在段落中替换文本（保留格式）"""
        if old_text not in paragraph.text:
            return

        # 清空段落内容
        paragraph.clear()

        # 添加新文本（保留原有格式或使用默认格式）
        run = paragraph.add_run(new_text)

        # 如果原段落有格式，尝试保留
        if paragraph.runs and len(paragraph.runs) > 0:
            original_run = paragraph.runs[0]
            if original_run.bold:
                run.bold = True
            if original_run.italic:
                run.italic = True
            if original_run.font.size:
                run.font.size = original_run.font.size

    async def _add_data_tables(self, doc: Document, data_ids: List[str], time_range: Dict[str, str]):
        """添加数据表格到文档"""
        # 这里需要从DataContextManager获取数据
        # 但由于这是非Context-Aware工具，我们通过简单的数据ID引用
        # 实际数据填充由用户在data_matches中提供

        # 添加数据表格占位符
        for i, data_id in enumerate(data_ids):
            heading = doc.add_heading(f"数据表格 {i+1}", level=2)
            doc.add_paragraph(f"数据来源: {data_id}")
            doc.add_paragraph(f"时间范围: {time_range.get('start_time', '')} 至 {time_range.get('end_time', '')}")

            # 创建表格占位符
            table = doc.add_table(rows=2, cols=3)
            table.style = 'Light Grid Accent 1'

            # 设置表头
            headers = table.rows[0].cells
            headers[0].text = "字段"
            headers[1].text = "数值"
            headers[2].text = "单位"

            # 添加示例数据行
            data_row = table.rows[1].cells
            data_row[0].text = "示例数据"
            data_row[1].text = "-"
            data_row[2].text = "-"

            doc.add_paragraph()  # 空行

    def _add_generation_info(self, doc: Document, time_range: Dict[str, str]):
        """添加生成信息"""
        # 在文档末尾添加生成信息
        doc.add_paragraph()
        info = doc.add_paragraph()
        info.add_run("---\n").bold = True
        info.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        if time_range.get("start_time") and time_range.get("end_time"):
            info.add_run(f"数据时间范围: {time_range['start_time']} 至 {time_range['end_time']}\n")
        info.add_run("生成工具: 大气污染溯源分析系统")
