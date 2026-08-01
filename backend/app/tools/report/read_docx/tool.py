"""
读取DOCX文档工具

使用python-docx直接读取DOCX文件内容，无需解包
"""

from typing import Dict, Any, Optional
from pathlib import Path
from docx import Document
from docx.table import Table
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
import structlog

from app.tools.base.tool_interface import LLMTool, ToolCategory

logger = structlog.get_logger()


class ReadDocxTool(LLMTool):
    """
    读取DOCX文档工具

    功能：
    - 直接读取DOCX文件内容（无需解包）
    - 提取段落和表格
    - 保留文档结构
    """

    def __init__(self):
        super().__init__(
            name="read_docx",
            description="""读取DOCX文档内容（直接读取，无需解包）

功能：
- 提取文档段落和表格
- 保留文档结构（标题层级）
- 自动生成 PDF 预览

⚠️ 格式说明：
- 标题会转换为 Markdown 格式（如 ## 标题），便于理解层级
- 这是工具添加的格式标记，实际文档不包含 # 符号
- 不要把这些 Markdown 标记当作文档原始内容

参数：
- path: DOCX 文件路径
- offset: 非空段落起始偏移量（可选，默认0）
- limit: 本次最多读取的非空段落数（可选）
- max_paragraphs: 最大段落数（可选，默认100，用于控制token消耗）
- include_tables: 是否包含表格（可选，默认true）
""",
            category=ToolCategory.QUERY,
            version="1.0.0",
            requires_context=False
        )

        self.function_schema = {
            "name": "read_docx",
            "description": "读取DOCX文档内容，提取段落和表格结构",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {
                        "type": "string",
                        "description": "DOCX文件路径"
                    },
                    "offset": {
                        "type": "integer",
                        "description": "非空段落起始偏移量（从0开始）",
                        "default": 0
                    },
                    "limit": {
                        "type": "integer",
                        "description": "本次最多读取的非空段落数"
                    },
                    "max_paragraphs": {
                        "type": "integer",
                        "description": "最大段落数（默认100，用于控制token消耗）",
                        "default": 100
                    },
                    "include_tables": {
                        "type": "boolean",
                        "description": "是否包含表格内容（默认true）",
                        "default": True
                    }
                },
                "required": ["path"]
            }
        }

    async def execute(
        self,
        path: str,
        offset: int = 0,
        limit: Optional[int] = None,
        max_paragraphs: int = 100,
        include_tables: bool = True,
        **kwargs
    ) -> Dict[str, Any]:
        """
        执行DOCX文档读取

        Args:
            path: DOCX文件路径
            offset: 非空段落起始偏移量（从0开始）
            limit: 本次最多读取的非空段落数
            max_paragraphs: 最大段落数
            include_tables: 是否包含表格

        Returns:
            读取结果
        """
        try:
            effective_limit = limit if limit is not None else max_paragraphs
            if offset < 0 or effective_limit <= 0:
                return {
                    "success": False,
                    "data": {
                        "error": (
                            "DOCX分页参数无效: offset 必须大于等于0，"
                            "limit 必须大于0"
                        )
                    },
                    "summary": "DOCX分页参数无效",
                }

            file_path = Path(path)

            # 验证文件存在
            if not file_path.exists():
                return {
                    "success": False,
                    "data": {"error": f"文件不存在: {path}"},
                    "summary": "文件不存在"
                }

            # 验证文件格式
            if file_path.suffix.lower() != ".docx":
                return {
                    "success": False,
                    "data": {"error": "只支持DOCX格式文件"},
                    "summary": "文件格式不支持"
                }

            # 尝试使用python-docx读取文档，如果失败则回退到手动解析
            try:
                doc = Document(str(file_path))
            except Exception as docx_error:
                logger.warning("python_docx_failed_fallback_to_manual", error=str(docx_error))
                # 回退方案：手动解包并提取文本
                return await self._read_docx_fallback(
                    file_path,
                    offset,
                    effective_limit,
                    include_tables,
                )

            # 提取内容（保持段落和表格的原始顺序）
            content_parts = []
            paragraph_count = 0  # 统计本页非空段落数
            table_count = 0
            image_count = 0  # 统计图片数量

            total_paragraphs = sum(1 for para in doc.paragraphs if para.text.strip())
            page_end = min(offset + effective_limit, total_paragraphs)
            paragraph_index = 0

            # 使用 iter_inner_content() 按原始顺序遍历段落和表格
            from docx.table import Table
            from docx.text.paragraph import Paragraph

            for item in doc.iter_inner_content():
                if isinstance(item, Paragraph):
                    # 处理段落
                    text = item.text.strip()
                    if text:
                        if offset <= paragraph_index < page_end:
                            # 检查是否是标题
                            if item.style and item.style.name.startswith("Heading"):
                                level = item.style.name.replace("Heading ", "")
                                try:
                                    level_int = int(level)
                                    content_parts.append(f"{'#' * level_int} {text}")
                                except ValueError:
                                    content_parts.append(text)
                            else:
                                content_parts.append(text)
                            paragraph_count += 1
                        paragraph_index += 1

                elif isinstance(item, Table) and include_tables:
                    if not (offset < total_paragraphs and offset <= paragraph_index < page_end):
                        continue
                    # 处理表格
                    if table_count >= 20:  # 限制表格数量
                        break

                    table_text = self._extract_table_text(item)
                    if table_text:
                        content_parts.append(table_text)
                        table_count += 1

            # 统计段落中嵌入的图片
            for para in doc.paragraphs:
                for run in para.runs:
                    if run._element.xpath('.//pic:pic'):
                        image_count += 1

            # 如果有图片，自动解包文档获取图片路径
            image_paths = []
            if image_count > 0:
                import tempfile
                import zipfile
                import shutil
                import uuid

                # 创建临时解包目录
                temp_dir = Path(tempfile.gettempdir()) / f"docx_images_{uuid.uuid4().hex[:8]}"
                try:
                    # 解包DOCX（本质是ZIP）
                    with zipfile.ZipFile(str(file_path), 'r') as zip_ref:
                        zip_ref.extractall(temp_dir)

                    # 查找所有图片文件（word/media目录）
                    media_dir = temp_dir / "word" / "media"
                    if media_dir.exists():
                        for img_file in sorted(media_dir.iterdir()):
                            if img_file.is_file() and img_file.suffix.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg']:
                                # 使用绝对路径
                                image_paths.append(str(img_file.resolve()))

                    logger.info(
                        "docx_images_extracted",
                        temp_dir=str(temp_dir),
                        image_count=len(image_paths)
                    )
                except Exception as extract_error:
                    logger.warning("docx_image_extraction_failed", error=str(extract_error))
                finally:
                    # 注意：不立即删除临时目录，因为 read_file 多模态读取需要访问
                    # 临时文件会在系统清理时删除（如/tmp定时清理）
                    pass

            # 构建结果
            content = "\n\n".join(content_parts)

            # 构建摘要信息
            summary_parts = [f"{paragraph_count}个段落", f"{table_count}个表格"]
            if image_count > 0:
                summary_parts.append(f"{image_count}个图片")

            result_data = {
                "type": "docx",
                "content": content,
                "path": str(file_path),
                "file_name": file_path.name,
                "file_size": file_path.stat().st_size,
                "paragraph_count": paragraph_count,
                "table_count": table_count,
                "image_count": image_count,
                "has_images": image_count > 0,
                "offset": offset,
                "limit": effective_limit,
                "next_offset": page_end if page_end < total_paragraphs else None,
                "is_truncated": page_end < total_paragraphs,
                "total_paragraphs": total_paragraphs,
                "total_tables": len(doc.tables)
            }

            # 如果有图片，添加图片路径列表
            if image_paths:
                result_data["image_paths"] = image_paths
                result_data["image_note"] = f"文档包含 {len(image_paths)} 张图片。"
                result_data["image_suggestion"] = "图片已提取到临时目录，可使用 read_file 并设置 as_multimodal_attachment=true 查看（使用上面 image_paths 中的路径）。"
            elif image_count > 0:
                # 如果检测到图片但提取失败
                result_data["image_note"] = f"文档包含 {image_count} 张图片，但提取失败。"
                result_data["image_suggestion"] = "图片未能自动提取；当前不暴露 Office XML 解包工具。"

            result = {
                "success": True,
                "data": result_data,
                "summary": f"读取成功: {file_path.name} ({', '.join(summary_parts)})",
                "metadata": {
                    "generator": "read_docx",
                    "tool_name": "read_docx"
                }
            }

            logger.info(
                "docx_read_success",
                path=str(file_path),
                paragraphs=paragraph_count,
                tables=table_count,
                images=image_count
            )

            return result

        except Exception as e:
            logger.error(
                "read_docx_failed",
                path=path,
                error=str(e),
                exc_info=True
            )
            return {
                "success": False,
                "data": {"error": str(e)},
                "summary": f"读取失败: {str(e)[:50]}"
            }

    def _extract_table_text(self, table: Table) -> str:
        """从表格提取文本（完整保留内容）"""
        try:
            lines = []
            col_count = 0

            for row_idx, row in enumerate(table.rows):
                cells = []
                for cell in row.cells:
                    # 提取文本并处理换行符
                    text = cell.text.strip()
                    # 替换换行符为空格，保持内容在同一行
                    text = text.replace('\n', ' ').replace('\r', '').strip()
                    # 移除多余空格
                    text = ' '.join(text.split())
                    # ✅ 完整保留内容，不再限制字符长度
                    cells.append(text)

                line = " | ".join(cells)
                lines.append(line)

                # 记录列数（以第一行为准）
                if row_idx == 0:
                    col_count = len(cells)

            if lines and col_count > 0:
                # 添加表头分隔符（与列数一致）
                separator = " | ".join(["---"] * col_count)
                return "\n".join([lines[0], separator] + lines[1:])

            return ""

        except Exception as e:
            logger.warning("extract_table_failed", error=str(e))
            return "[表格解析失败]"

    async def _read_docx_fallback(
        self,
        file_path: Path,
        offset: int,
        limit: int,
        include_tables: bool
    ) -> Dict[str, Any]:
        """
        回退方案：手动解包DOCX并提取文本（绕过python-docx的XML解析）

        用于处理python-docx无法解析的损坏文档
        """
        import zipfile
        import re
        from xml.etree import ElementTree as ET

        temp_dir = None
        try:
            # 创建临时解包目录
            import tempfile
            import uuid
            temp_dir = Path(tempfile.gettempdir()) / f"docx_fallback_{uuid.uuid4().hex[:8]}"
            temp_dir.mkdir(exist_ok=True)

            # 解包DOCX
            with zipfile.ZipFile(str(file_path), 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # 读取主文档
            doc_xml_path = temp_dir / "word" / "document.xml"
            if not doc_xml_path.exists():
                return {
                    "success": False,
                    "data": {"error": "无法找到word/document.xml，文档可能已损坏"},
                    "summary": "文档结构异常"
                }

            # 读取XML内容
            with open(doc_xml_path, 'r', encoding='utf-8', errors='ignore') as f:
                xml_content = f.read()

            # 方法1：尝试标准XML解析（忽略命名空间）
            try:
                # 移除命名空间以简化解析
                xml_content_clean = re.sub(r'xmlns[^"]*"[^"]*"', '', xml_content)
                root = ET.fromstring(xml_content_clean)

                # 提取所有文本
                paragraphs = []
                for elem in root.iter():
                    if elem.text and elem.text.strip():
                        text = elem.text.strip()
                        if text:
                            paragraphs.append(text)

                page = paragraphs[offset:offset + limit]
                next_offset = offset + len(page)
                is_truncated = next_offset < len(paragraphs)
                content = "\n\n".join(page)

                # 提取图片
                image_paths = self._extract_images_fallback(temp_dir)

                result_data = {
                    "type": "docx",
                    "content": content,
                    "path": str(file_path),
                    "file_name": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "paragraph_count": len(page),
                    "table_count": 0,
                    "image_count": len(image_paths),
                    "has_images": len(image_paths) > 0,
                    "offset": offset,
                    "limit": limit,
                    "next_offset": next_offset if is_truncated else None,
                    "is_truncated": is_truncated,
                    "total_paragraphs": len(paragraphs),
                    "extraction_method": "fallback_xml",
                    "warning": "文档XML格式不规范，使用回退方案提取"
                }

                if image_paths:
                    result_data["image_paths"] = image_paths
                    result_data["image_note"] = f"文档包含 {len(image_paths)} 张图片。"
                    result_data["image_suggestion"] = "图片已提取到临时目录，可使用 read_file 并设置 as_multimodal_attachment=true 查看。"

                return {
                    "success": True,
                    "data": result_data,
                    "summary": f"读取成功（回退模式）: {file_path.name} ({len(page)}个段落)",
                    "metadata": {
                        "generator": "read_docx",
                        "tool_name": "read_docx",
                        "extraction_method": "fallback"
                    }
                }

            except ET.ParseError:
                # 方法2：XML解析失败，使用正则表达式提取文本
                logger.warning("xml_parse_failed_using_regex")

                # 提取所有 <w:t> 标签中的文本
                text_pattern = re.compile(r'<w:t[^>]*>([^<]+)</w:t>', re.DOTALL)
                texts = text_pattern.findall(xml_content)

                # 清理文本
                paragraphs = []
                for text in texts:
                    text = text.strip()
                    if text and len(text) > 1:  # 过滤单个字符
                        paragraphs.append(text)

                page = paragraphs[offset:offset + limit]
                next_offset = offset + len(page)
                is_truncated = next_offset < len(paragraphs)
                content = "\n\n".join(page)

                # 提取图片
                image_paths = self._extract_images_fallback(temp_dir)

                result_data = {
                    "type": "docx",
                    "content": content,
                    "path": str(file_path),
                    "file_name": file_path.name,
                    "file_size": file_path.stat().st_size,
                    "paragraph_count": len(page),
                    "table_count": 0,
                    "image_count": len(image_paths),
                    "has_images": len(image_paths) > 0,
                    "offset": offset,
                    "limit": limit,
                    "next_offset": next_offset if is_truncated else None,
                    "is_truncated": is_truncated,
                    "total_paragraphs": len(paragraphs),
                    "extraction_method": "fallback_regex",
                    "warning": "文档严重损坏，使用正则表达式提取（可能丢失格式）"
                }

                if image_paths:
                    result_data["image_paths"] = image_paths
                    result_data["image_note"] = f"文档包含 {len(image_paths)} 张图片。"

                return {
                    "success": True,
                    "data": result_data,
                    "summary": f"读取成功（正则模式）: {file_path.name} ({len(page)}个段落)",
                    "metadata": {
                        "generator": "read_docx",
                        "tool_name": "read_docx",
                        "extraction_method": "regex"
                    }
                }

        except Exception as e:
            logger.error("read_docx_fallback_failed", error=str(e), exc_info=True)
            return {
                "success": False,
                "data": {"error": f"回退方案也失败: {str(e)}"},
                "summary": f"文档读取失败: {str(e)[:50]}"
            }
        finally:
            # 清理临时目录
            if temp_dir and temp_dir.exists():
                try:
                    import shutil
                    shutil.rmtree(temp_dir, ignore_errors=True)
                except:
                    pass

    def _extract_images_fallback(self, temp_dir: Path) -> list:
        """从解包的DOCX中提取图片路径"""
        image_paths = []
        try:
            media_dir = temp_dir / "word" / "media"
            if media_dir.exists():
                for img_file in sorted(media_dir.iterdir()):
                    if img_file.is_file() and img_file.suffix.lower() in ['.png', '.jpg', '.jpeg', '.gif', '.bmp', '.svg']:
                        # 复制到持久位置
                        import tempfile
                        import uuid
                        import shutil
                        persist_dir = Path(tempfile.gettempdir()) / "docx_images_persistent"
                        persist_dir.mkdir(exist_ok=True)
                        dest_path = persist_dir / f"{uuid.uuid4().hex[:8]}_{img_file.name}"
                        shutil.copy2(img_file, dest_path)
                        image_paths.append(str(dest_path.resolve()))
        except Exception as e:
            logger.warning("extract_images_fallback_failed", error=str(e))
        return image_paths
