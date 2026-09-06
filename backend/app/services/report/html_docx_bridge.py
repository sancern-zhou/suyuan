"""Preserve rendered report tables and author CSS across Pandoc's HTML reader.

Quarto owns QMD parsing/execution. Pandoc owns HTML table parsing (including
merged and nested cells). This module only resolves CSS and verifies the result.
"""

from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import re
from urllib.parse import unquote, urlparse
from zipfile import BadZipFile

import cssselect2
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import html
import tinycss2
from tinycss2.color3 import parse_color


def _text(value: str) -> str:
    return re.sub(r"\s+", "", value)


def _color(value: str | None) -> str | None:
    if not value:
        return None
    parsed = parse_color(value)
    if parsed is None or isinstance(parsed, str) or parsed.alpha == 0:
        return None
    # Word colors have no alpha channel; composite translucent colors on white.
    return "".join(f"{round((c * parsed.alpha + 1 - parsed.alpha) * 255):02X}" for c in parsed[:3])


def _declarations(content):
    for item in tinycss2.parse_declaration_list(content, skip_comments=True, skip_whitespace=True):
        if item.type != "declaration":
            continue
        name = item.name if item.name.startswith("--") else item.lower_name
        value = tinycss2.serialize(item.value).strip()
        if name == "background":
            name = "background-color"
        if name.startswith("--") or name in {
            "color",
            "background-color",
            "font-weight",
            "font-style",
            "display",
        }:
            yield name, value, item.important


def _resolve_variables(value, properties, seen=frozenset()):
    parts = []
    for token in tinycss2.parse_component_value_list(value):
        if token.type == "function" and token.lower_name == "var":
            arguments = token.arguments
            variable = next((t.value for t in arguments if t.type == "ident"), "")
            comma = next(
                (i for i, t in enumerate(arguments) if t.type == "literal" and t.value == ","), None
            )
            fallback = tinycss2.serialize(arguments[comma + 1 :]) if comma is not None else ""
            resolved = properties.get(variable, fallback) if variable not in seen else fallback
            parts.append(
                _resolve_variables(resolved, properties, seen | {variable})
                if resolved and variable not in seen
                else fallback
            )
        else:
            parts.append(tinycss2.serialize([token]))
    return "".join(parts).strip()


def _rules(content):
    for rule in tinycss2.parse_stylesheet(content, skip_comments=True, skip_whitespace=True):
        if rule.type == "qualified-rule":
            yield rule
        elif rule.type == "at-rule" and rule.lower_at_keyword == "media" and rule.content:
            media = tinycss2.serialize(rule.prelude).strip().lower()
            if media in {"print", "all"}:
                yield from _rules(tinycss2.serialize(rule.content))


@dataclass
class HtmlCell:
    row: int
    column: int
    rowspan: int
    colspan: int
    element: object
    fill: str | None


class RenderedHtmlReport:
    def __init__(self, path: Path):
        self.path = path
        self.tree = html.document_fromstring(path.read_text(encoding="utf-8"))
        self.styles = self._compute_styles()
        candidates = self.tree.xpath('//main[@id="quarto-document-content"]')
        self.body = candidates[0] if candidates else self.tree.find("body")
        if self.body is None:
            raise ValueError("Rendered report has no HTML body")
        # These are page controls or examples, never report content.
        for node in list(
            self.body.xpath(".//nav|.//script|.//style|.//template|.//pre//table|.//code//table")
        ):
            if node.getparent() is not None:
                node.drop_tree()
        for node in list(self.body.iterdescendants()):
            if (
                self.styles.get(node, {}).get("display") == "none" or node.get("hidden") is not None
            ) and node.getparent() is not None:
                node.drop_tree()
        self.tables = list(self.body.iter("table"))
        self.cells = [self._cells(table) for table in self.tables]
        self.run_styles: dict[str, dict[str, str]] = {}
        self._mark_text_styles()

    def _compute_styles(self):
        matcher = cssselect2.Matcher()
        # Only author CSS is exported. Quarto's framework CSS is screen layout.
        sheets = self.tree.xpath('//style|//link[@rel="stylesheet"]')
        for sheet in sheets:
            if sheet.get("id") in {"quarto-bootstrap", "quarto-text-highlighting-styles"}:
                continue
            if sheet.get("media", "all").lower() not in {"all", "print", "screen"}:
                continue
            content = sheet.text or ""
            if sheet.tag == "link":
                url = urlparse(sheet.get("href", ""))
                if url.scheme or url.netloc:
                    raise ValueError("DOCX author stylesheets must be local report assets")
                css_path = (self.path.parent / unquote(url.path)).resolve()
                if not css_path.is_relative_to(self.path.parent.resolve()):
                    raise ValueError("DOCX stylesheet is outside the report package")
                content = css_path.read_text(encoding="utf-8")
            for rule in _rules(content):
                try:
                    selectors = cssselect2.compile_selector_list(rule.prelude)
                except cssselect2.SelectorError:
                    continue
                for selector in selectors:
                    if selector.pseudo_element is None:
                        matcher.add_selector(selector, list(_declarations(rule.content)))
        results = {}
        root = cssselect2.ElementWrapper.from_html_root(self.tree)
        for wrapper in root.iter_subtree():
            element = wrapper.etree_element
            inherited = results.get(element.getparent(), {})
            result = {
                k: v
                for k, v in inherited.items()
                if k.startswith("--") or k in {"color", "font-weight", "font-style"}
            }
            if element.tag in {"strong", "b"}:
                result["font-weight"] = "bold"
            if element.tag in {"em", "i"}:
                result["font-style"] = "italic"
            winners = {}
            for specificity, order, _, declarations in matcher.match(wrapper):
                for name, value, important in declarations:
                    rank = (important, 0, specificity, order)
                    if name not in winners or rank >= winners[name][0]:
                        winners[name] = (rank, value)
            for name, value, important in _declarations(element.get("style", "")):
                rank = (important, 1, (0, 0, 0), 0)
                if name not in winners or rank >= winners[name][0]:
                    winners[name] = (rank, value)
            for name, (_, value) in winners.items():
                result[name] = inherited.get(name, "") if value == "inherit" else value
            for name in ("color", "background-color", "font-weight", "font-style", "display"):
                if name in result:
                    result[name] = _resolve_variables(result[name], result)
            if "background-color" in result:
                tokens = tinycss2.parse_component_value_list(result["background-color"])
                result["background-color"] = next(
                    (tinycss2.serialize([t]) for t in tokens if _color(tinycss2.serialize([t]))),
                    "transparent",
                )
            if element.get("bgcolor") and "background-color" not in winners:
                result["background-color"] = element.get("bgcolor")
            if element.tag == "font" and element.get("color") and "color" not in winners:
                result["color"] = element.get("color")
            results[element] = result
        return results

    def _cells(self, table):
        rows = [r for r in table.iter("tr") if next(r.iterancestors("table"), None) is table]
        # HTML layout puts the footer last regardless of its source position.
        rows.sort(
            key=lambda r: (
                0
                if next(r.iterancestors("thead"), None) is not None
                else (2 if next(r.iterancestors("tfoot"), None) is not None else 1)
            )
        )
        occupied = set()
        cells = []
        for row_index, row in enumerate(rows):
            column = 0
            for cell in row:
                if cell.tag not in {"td", "th"}:
                    continue
                while (row_index, column) in occupied:
                    column += 1
                rowspan = int(cell.get("rowspan", "1"))
                colspan = int(cell.get("colspan", "1"))
                if rowspan == 0:
                    rowspan = sum(r.getparent() is row.getparent() for r in rows[row_index:])
                    cell.set("rowspan", str(rowspan))
                if rowspan < 1 or colspan < 1 or rowspan > 1000 or colspan > 1000:
                    raise ValueError("Invalid HTML table span")
                fill = None
                for ancestor in [cell, *cell.iterancestors()]:
                    fill = _color(self.styles.get(ancestor, {}).get("background-color"))
                    if fill or ancestor is table:
                        break
                cells.append(HtmlCell(row_index, column, rowspan, colspan, cell, fill))
                occupied.update(
                    (r, c)
                    for r in range(row_index, row_index + rowspan)
                    for c in range(column, column + colspan)
                )
                column += colspan
        return cells

    def _mark_text_styles(self):
        # Character styles survive the HTML reader and DOCX writer. Resolve them
        # after the government formatter, which otherwise resets explicit colors.
        style_ids = {}
        for element in list(self.body.iter()):
            if not isinstance(element.tag, str):
                continue
            # Headings and captions are rebuilt by the government formatter.
            if any(
                node.tag in {"h1", "h2", "h3", "h4", "h5", "h6", "figcaption"}
                for node in [element, *element.iterancestors()]
            ):
                continue
            properties = self.styles.get(element, {})
            values = {}
            for name in ("color", "background-color"):
                color = _color(properties.get(name))
                if color:
                    values[name] = color
            for name in ("font-weight", "font-style"):
                if properties.get(name):
                    values[name] = properties[name]
            if not values:
                continue
            if not (
                (element.text and element.text.strip())
                or any(child.tail and child.tail.strip() for child in element)
            ):
                continue
            key = tuple(sorted(values.items()))
            style_id = style_ids.setdefault(key, f"ReportHtml{len(style_ids) + 1}")
            self.run_styles[style_id] = values
            if element.text and element.text.strip():
                span = html.Element("span", {"data-custom-style": style_id})
                span.text, element.text = element.text, None
                element.insert(0, span)
            for child in list(element):
                if child.tail and child.tail.strip():
                    span = html.Element("span", {"data-custom-style": style_id})
                    span.text, child.tail = child.tail, None
                    child.addnext(span)

    def prepare(self, output: Path):
        for title in self.body.xpath('.//header[@id="title-block-header"]'):
            headings = title.xpath(".//h1")
            for heading in headings:
                heading.tag = "p"
                parent = html.Element("div", {"data-custom-style": "Title"})
                heading.addprevious(parent)
                parent.append(heading)
        output.write_text(html.tostring(self.body, encoding="unicode"), encoding="utf-8")

    def _word_cells(self, doc):
        tables = doc.element.body.xpath(".//w:tbl")
        if len(tables) != len(self.tables):
            raise ValueError(
                f"DOCX table count mismatch: HTML={len(self.tables)}, DOCX={len(tables)}"
            )
        for table, expected in zip(tables, self.cells):
            rows = table.findall(qn("w:tr"))
            expected_rows = max((c.row + c.rowspan for c in expected), default=0)
            if len(rows) != expected_rows:
                raise ValueError("DOCX table row count mismatch")
            for cell in expected:
                if cell.row >= len(rows):
                    raise ValueError("DOCX table row missing")
                column = 0
                target = None
                for tc in rows[cell.row].findall(qn("w:tc")):
                    if column == cell.column:
                        target = tc
                        break
                    spans = tc.xpath("./w:tcPr/w:gridSpan/@w:val")
                    column += int(spans[0]) if spans else 1
                if target is None:
                    raise ValueError("DOCX table cell missing")
                spans = target.xpath("./w:tcPr/w:gridSpan/@w:val")
                if (int(spans[0]) if spans else 1) != cell.colspan:
                    raise ValueError("DOCX table column span mismatch")
                if cell.rowspan > 1 and target.xpath("./w:tcPr/w:vMerge/@w:val") != ["restart"]:
                    raise ValueError("DOCX table row span missing")
                actual = "".join(target.xpath(".//w:t/text()"))
                wanted = "".join(cell.element.itertext())
                if _text(actual) != _text(wanted):
                    raise ValueError(
                        f"DOCX table cell content mismatch at row {cell.row + 1}, column {cell.column + 1}"
                    )
                yield target, cell

    def apply_styles(self, path: Path):
        doc = Document(path)
        for tc, cell in self._word_cells(doc):
            if cell.fill:
                _property(tc.get_or_add_tcPr(), "shd", fill=cell.fill, val="clear")
        for run in doc.element.body.xpath(".//w:r"):
            style = run.xpath("./w:rPr/w:rStyle/@w:val")
            values = self.run_styles.get(style[0] if style else "", {})
            if not values:
                continue
            pr = run.get_or_add_rPr()
            if "color" in values:
                _property(pr, "color", val=values["color"])
            if "background-color" in values:
                _property(pr, "shd", fill=values["background-color"], val="clear")
            if "font-weight" in values:
                weight = values["font-weight"]
                _property(
                    pr,
                    "b",
                    val="1"
                    if weight in {"bold", "bolder"} or (weight.isdigit() and int(weight) >= 600)
                    else "0",
                )
            if "font-style" in values:
                _property(
                    pr, "i", val="1" if values["font-style"] in {"italic", "oblique"} else "0"
                )
        doc.save(path)
        self.validate(path)

    def validate(self, path: Path):
        try:
            doc = Document(path)
        except (BadZipFile, PackageNotFoundError) as exc:
            raise ValueError("DOCX package is unreadable") from exc
        for tc, cell in self._word_cells(doc):
            if cell.fill and tc.xpath("./w:tcPr/w:shd/@w:fill") != [cell.fill]:
                raise ValueError("DOCX table background color missing")
        for style_id, values in self.run_styles.items():
            runs = doc.element.body.xpath(f'.//w:r[w:rPr/w:rStyle[@w:val="{style_id}"]]')
            if not runs and ("color" in values or "background-color" in values):
                raise ValueError(f"DOCX styled content missing: {style_id}")
            for run in runs:
                if "color" in values and run.xpath("./w:rPr/w:color/@w:val") != [values["color"]]:
                    raise ValueError("DOCX text color missing")
                if "background-color" in values and run.xpath("./w:rPr/w:shd/@w:fill") != [
                    values["background-color"]
                ]:
                    raise ValueError("DOCX text background missing")
        return {"tables": len(self.tables), "cells": sum(map(len, self.cells))}


def _property(parent, tag, **attributes):
    element = parent.find(qn(f"w:{tag}"))
    if element is None:
        element = OxmlElement(f"w:{tag}")
        parent.append(element)
    for key, value in attributes.items():
        element.set(qn(f"w:{key}"), value)
