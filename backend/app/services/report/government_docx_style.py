"""
Default government-report DOCX styling helpers.

This module keeps report formatting deterministic for both direct
python-docx generation and Quarto/Pandoc DOCX export. User-specific
formatting can still override these defaults in generated code or qmd YAML.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable
from urllib.parse import unquote, urlparse

from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.text.paragraph import Paragraph


DEFAULT_REFERENCE_DOCX = (
    Path(__file__).resolve().parents[3]
    / "backend_data_registry"
    / "report_templates"
    / "government_report_reference.docx"
)


BODY_FONT = "仿宋_GB2312"
BODY_FONT_FALLBACK = "仿宋"
HEADING_FONT = "黑体"
TITLE_FONT = "黑体"
TITLE_FONT_FALLBACK = "黑体"
SECOND_LEVEL_FONT = "楷体_GB2312"
SECOND_LEVEL_FONT_FALLBACK = "楷体"

BODY_SIZE_PT = 12
TITLE_SIZE_PT = 22
SUBTITLE_SIZE_PT = 16
COVER_META_SIZE_PT = 16
LINE_SPACING_PT = 22
COVER_TITLE_LINE_SPACING_PT = 32
COVER_TITLE_SPACE_BEFORE_PT = 128
COVER_SUBTITLE_SPACE_AFTER_PT = 18
COVER_META_SPACE_BEFORE_PT = 255

# 政府公文标题字号规范
HEADING_LEVEL_1_SIZE_PT = 22  # 二号
HEADING_LEVEL_2_SIZE_PT = 16  # 三号
HEADING_LEVEL_3_SIZE_PT = 16  # 三号
HEADING_LEVEL_4_SIZE_PT = 14  # 四号


SKIP_HTML_TAGS = {"script", "style", "meta", "link", "noscript"}

# 标题序号计数器（全局）
_heading_counters = {
    1: 0,
    2: 0,
    3: 0,
    4: 0
}

# 上一个标题的级别（用于重置下级计数器）
_last_heading_level = 0


def _reset_heading_counters(level: int) -> None:
    """重置指定级别及以下的所有计数器"""
    global _last_heading_level
    _last_heading_level = level
    for l in range(level + 1, 5):
        _heading_counters[l] = 0


def _get_heading_number(level: int) -> str:
    """
    生成政府公文格式的标题序号

    Args:
        level: 标题级别（1-4）

    Returns:
        标题序号字符串
    """
    global _heading_counters, _last_heading_level

    # 重置下级计数器
    _reset_heading_counters(level)

    # 增加当前级别计数
    _heading_counters[level] += 1

    # 根据级别生成序号
    if level == 1:
        # 一级：一、二、三、
        chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十",
                       "十一", "十二", "十三", "十四", "十五", "十六", "十七", "十八", "十九", "二十"]
        num = _heading_counters[level]
        if num <= 20:
            return chinese_nums[num - 1] + "、"
        else:
            # 超过20使用数字
            return f"{num}、"
    elif level == 2:
        # 二级：（一）（二）（三）
        chinese_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        num = _heading_counters[level]
        if num <= 10:
            return f"（{chinese_nums[num - 1]}）"
        else:
            return f"（{num}）"
    elif level == 3:
        # 三级：1. 2. 3.
        return f"{_heading_counters[level]}."
    elif level == 4:
        # 四级：（1）（2）（3）
        return f"（{_heading_counters[level]}）"
    return ""


def set_run_font(
    run,
    font_name: str,
    size_pt: float | None = None,
    bold: bool | None = None,
    color: str | None = None,
) -> None:
    """Set Latin and East Asian font metadata on a run."""
    run.font.name = font_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    else:
        _remove_run_color(run)


def _remove_run_color(run) -> None:
    """Use Word's automatic font color by removing explicit color nodes."""
    r_pr = run._element.get_or_add_rPr()
    for color_element in list(r_pr.findall(qn("w:color"))):
        r_pr.remove(color_element)


def set_paragraph_format(
    paragraph,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_chars: float = 2,
    line_spacing_pt: float = LINE_SPACING_PT,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    """Apply common government-report paragraph spacing."""
    fmt = paragraph.paragraph_format
    fmt.alignment = alignment
    fmt.line_spacing = Pt(line_spacing_pt)
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    if first_line_indent_chars:
        # Approximate "2 characters" using the current body font size.
        fmt.first_line_indent = Pt(BODY_SIZE_PT * first_line_indent_chars)
    else:
        fmt.first_line_indent = Pt(0)


def paragraph_has_drawing(paragraph) -> bool:
    """Return True when a paragraph contains an inline picture/drawing."""
    return bool(paragraph._p.xpath(".//w:drawing"))


def set_image_paragraph_format(paragraph) -> None:
    """Keep inline pictures in their own visible Word paragraph."""
    fmt = paragraph.paragraph_format
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = 1.0
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE
    fmt.space_before = Pt(6)
    fmt.space_after = Pt(6)


def normalize_docx_image_paragraphs(docx_path: str | Path) -> dict:
    """Apply safe paragraph formatting to every inline picture in a DOCX."""
    path = Path(docx_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        return {"path": str(path), "image_paragraphs": 0, "updated": False}

    doc = Document(str(path))
    image_paragraphs = 0
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            set_image_paragraph_format(paragraph)
            image_paragraphs += 1

    if image_paragraphs:
        doc.save(str(path))

    return {
        "path": str(path),
        "image_paragraphs": image_paragraphs,
        "updated": bool(image_paragraphs),
    }


def _insert_paragraph_before(paragraph):
    new_p = OxmlElement("w:p")
    paragraph._p.addprevious(new_p)
    return Paragraph(new_p, paragraph._parent)


def _insert_page_break(paragraph) -> None:
    run = paragraph.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _add_field_run(paragraph, instr: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" {instr.strip()} "
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr_text, fld_sep, fld_end])


def _add_toc_field(paragraph) -> None:
    _add_field_run(paragraph, r'TOC \o "1-3" \h \z \u')


def _add_page_field(paragraph) -> None:
    _add_field_run(paragraph, "PAGE")


def _enable_update_fields_on_open(doc: DocumentObject) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _set_section_page_number_start(section, start: int = 1) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), str(start))


def _unlink_header_footer(section) -> None:
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def _clear_paragraph(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag == qn("w:pPr"):
            continue
        paragraph._p.remove(child)


def _clear_footer(footer) -> None:
    for paragraph in footer.paragraphs:
        _clear_paragraph(paragraph)


def _add_centered_page_number(section) -> None:
    footer = section.footer
    _clear_footer(footer)
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_chars=0,
        line_spacing_pt=20,
        space_before_pt=0,
        space_after_pt=0,
    )
    _add_page_field(paragraph)
    for run in paragraph.runs:
        set_run_font(run, BODY_FONT, 10.5)


def _is_cover_or_toc_heading(text: str) -> bool:
    normalized = sanitize_report_text(text)
    return normalized in {"目录"} or normalized.endswith("报告")


def _heading_prefix(level: int, counters: dict[int, int]) -> str:
    if level == 1:
        counters[1] += 1
        counters[2] = 0
        counters[3] = 0
        cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        number = cn_nums[counters[1] - 1] if counters[1] <= len(cn_nums) else str(counters[1])
        return f"{number}、"
    if level == 2:
        counters[2] += 1
        counters[3] = 0
        cn_nums = ["一", "二", "三", "四", "五", "六", "七", "八", "九", "十"]
        number = cn_nums[counters[2] - 1] if counters[2] <= len(cn_nums) else str(counters[2])
        return f"（{number}）"
    counters[3] += 1
    return f"{counters[3]}. "


def _strip_existing_heading_prefix(text: str) -> str:
    text = sanitize_report_text(text)
    patterns = [
        r"^[一二三四五六七八九十百]+、\s*",
        r"^（[一二三四五六七八九十百]+）\s*",
        r"^\d+[\.、]\s*",
        r"^\d+(?:\.\d+)+\s*",
    ]
    for pattern in patterns:
        text = re.sub(pattern, "", text)
    return text


def _heading_level(paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style else ""
    match = re.match(r"Heading\s+([1-6])", style_name)
    if not match:
        return None
    return int(match.group(1))


def _apply_heading_numbering(doc: DocumentObject) -> int:
    counters = {1: 0, 2: 0, 3: 0}
    updated = 0
    for paragraph in doc.paragraphs:
        level = _heading_level(paragraph)
        if not level or level > 3:
            continue
        raw_text = paragraph.text.strip()
        if not raw_text or _is_cover_or_toc_heading(raw_text):
            continue
        clean_text = _strip_existing_heading_prefix(raw_text)
        prefix = _heading_prefix(level, counters)
        target = prefix + clean_text
        if raw_text == target:
            continue
        _clear_paragraph(paragraph)
        run = paragraph.add_run(target)
        font = HEADING_FONT if level == 1 else SECOND_LEVEL_FONT if level == 2 else BODY_FONT
        set_run_font(run, font, BODY_SIZE_PT, bold=(level == 3))
        updated += 1
    return updated


def _first_report_heading(doc: DocumentObject):
    for paragraph in doc.paragraphs:
        level = _heading_level(paragraph)
        text = paragraph.text.strip()
        if level and text and text != "目录":
            return paragraph
    return None


def _toc_already_inserted(doc: DocumentObject) -> bool:
    return any(paragraph.text.strip() == "目录" for paragraph in doc.paragraphs[:8])


def _insert_toc_page(doc: DocumentObject) -> tuple[bool, Paragraph | None]:
    """插入目录页，返回(是否插入, 目录后的分页段落)"""
    if not doc.paragraphs or _toc_already_inserted(doc):
        return False, None

    body_start = _first_report_heading(doc)
    if body_start is None:
        body_start = doc.paragraphs[0]

    cover_break = _insert_paragraph_before(body_start)
    _insert_page_break(cover_break)

    toc_title = _insert_paragraph_before(body_start)
    toc_title.style = doc.styles["Title"]
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_run = toc_title.add_run("目录")
    set_run_font(toc_run, HEADING_FONT, BODY_SIZE_PT, bold=True)

    toc_field = _insert_paragraph_before(body_start)
    toc_field.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_paragraph_format(
        toc_field,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_chars=0,
        line_spacing_pt=22,
        space_before_pt=0,
        space_after_pt=0,
    )
    _add_toc_field(toc_field)

    body_break = _insert_paragraph_before(body_start)
    _insert_page_break(body_break)
    return True, body_break


def _ensure_body_section_starts_after_toc(doc: DocumentObject, toc_break_paragraph: Paragraph | None) -> bool:
    """在目录后的分页段落设置节属性，而不是在第一个标题上"""
    if toc_break_paragraph is None:
        return False

    p_pr = toc_break_paragraph._p.get_or_add_pPr()
    sect_pr = p_pr.find(qn("w:sectPr"))
    if sect_pr is None:
        sect_pr = OxmlElement("w:sectPr")
        p_pr.append(sect_pr)
    type_el = sect_pr.find(qn("w:type"))
    if type_el is None:
        type_el = OxmlElement("w:type")
        sect_pr.append(type_el)
    type_el.set(qn("w:val"), "nextPage")
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")
    return True


def _format_cover_paragraph(
    paragraph: Paragraph,
    *,
    font_name: str,
    size_pt: float,
    bold: bool = False,
    line_spacing_pt: float = COVER_TITLE_LINE_SPACING_PT,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt = paragraph.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fmt.first_line_indent = Pt(0)
    fmt.line_spacing = Pt(line_spacing_pt)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    for run in paragraph.runs:
        set_run_font(run, font_name, size_pt, bold=bold)
        _remove_run_color(run)


def _format_cover_page(doc: DocumentObject) -> bool:
    """Normalize Pandoc-generated title metadata into a formal cover page."""
    cover_paragraphs: list[Paragraph] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            if cover_paragraphs:
                break
            continue
        style_name = paragraph.style.name if paragraph.style else ""
        if text == "目录" or style_name.startswith("Heading"):
            break
        cover_paragraphs.append(paragraph)

    if not cover_paragraphs:
        return False

    _format_cover_paragraph(
        cover_paragraphs[0],
        font_name=TITLE_FONT,
        size_pt=TITLE_SIZE_PT,
        line_spacing_pt=COVER_TITLE_LINE_SPACING_PT,
        space_before_pt=COVER_TITLE_SPACE_BEFORE_PT,
        space_after_pt=0,
    )

    meta_start_index = 1
    if len(cover_paragraphs) >= 2 and cover_paragraphs[1].style and cover_paragraphs[1].style.name == "Subtitle":
        _format_cover_paragraph(
            cover_paragraphs[1],
            font_name=HEADING_FONT,
            size_pt=SUBTITLE_SIZE_PT,
            line_spacing_pt=COVER_TITLE_LINE_SPACING_PT,
            space_before_pt=0,
            space_after_pt=COVER_SUBTITLE_SPACE_AFTER_PT,
        )
        meta_start_index = 2

    meta_paragraphs = [p for p in cover_paragraphs[meta_start_index:] if p.text.strip()]
    for idx, paragraph in enumerate(meta_paragraphs):
        _format_cover_paragraph(
            paragraph,
            font_name=BODY_FONT,
            size_pt=COVER_META_SIZE_PT,
            line_spacing_pt=28,
            space_before_pt=COVER_META_SPACE_BEFORE_PT if idx == 0 else 0,
            space_after_pt=0,
        )

    return True


def finalize_government_docx(docx_path: str | Path, *, add_toc: bool = True) -> dict:
    """Normalize Quarto DOCX output to the project's government-report profile."""
    path = Path(docx_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        return {"path": str(path), "updated": False, "reason": "not_docx"}

    doc = Document(str(path))
    apply_government_report_style(doc)
    invalid_date_removed = 0
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "Invalid Date":
            _clear_paragraph(paragraph)
            invalid_date_removed += 1

    if add_toc:
        toc_inserted, toc_break = _insert_toc_page(doc)
        _ensure_body_section_starts_after_toc(doc, toc_break)
    else:
        toc_inserted = False

    cover_formatted = _format_cover_page(doc)

    heading_numbers = 0
    image_paragraphs = 0
    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            set_image_paragraph_format(paragraph)
            image_paragraphs += 1
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # 根据标题级别设置字体大小
            level = _heading_level(paragraph)
            size_map = {
                1: HEADING_LEVEL_1_SIZE_PT,
                2: HEADING_LEVEL_2_SIZE_PT,
                3: HEADING_LEVEL_3_SIZE_PT,
                4: HEADING_LEVEL_4_SIZE_PT,
            }
            size_pt = size_map.get(level, BODY_SIZE_PT)

            for run in paragraph.runs:
                set_run_font(run, HEADING_FONT, size_pt)
                _remove_run_color(run)

    for table in doc.tables:
        format_government_table(table)

    for idx, section in enumerate(doc.sections):
        _unlink_header_footer(section)
        if idx == 0:
            _clear_footer(section.footer)
        else:
            _set_section_page_number_start(section, 1)
            _add_centered_page_number(section)

    if len(doc.sections) == 1:
        _add_centered_page_number(doc.sections[0])

    _enable_update_fields_on_open(doc)
    doc.save(str(path))
    return {
        "path": str(path),
        "updated": True,
        "toc_inserted": toc_inserted,
        "heading_numbers": heading_numbers,
        "image_paragraphs": image_paragraphs,
        "sections": len(doc.sections),
        "invalid_date_removed": invalid_date_removed,
        "cover_formatted": cover_formatted,
    }


def _set_style_font(style, font_name: str, size_pt: float, bold: bool = False) -> None:
    style.font.name = font_name
    style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    r_pr = style._element.get_or_add_rPr()
    for color_element in list(r_pr.findall(qn("w:color"))):
        r_pr.remove(color_element)


def _set_style_paragraph(
    style,
    *,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent_chars: float = 2,
    line_spacing_pt: float = LINE_SPACING_PT,
    space_before_pt: float = 0,
    space_after_pt: float = 0,
) -> None:
    fmt = style.paragraph_format
    fmt.alignment = alignment
    fmt.line_spacing = Pt(line_spacing_pt)
    fmt.space_before = Pt(space_before_pt)
    fmt.space_after = Pt(space_after_pt)
    if first_line_indent_chars:
        fmt.first_line_indent = Pt(BODY_SIZE_PT * first_line_indent_chars)
    else:
        fmt.first_line_indent = Pt(0)


def _set_table_cell_text(cell, text: str, *, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_chars=0,
        line_spacing_pt=20,
        space_before_pt=0,
        space_after_pt=0,
    )
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, 10.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _set_table_borders(table) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "8")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")


def apply_government_report_style(doc: DocumentObject) -> DocumentObject:
    """
    Apply a default government-report style profile to an existing document.

    Defaults:
    - A4-ish government margins: top 3.7 cm, bottom 3.5 cm, left/right 2.8 cm.
    - Cover title: centered, 22 pt 黑体, fixed 32 pt line spacing.
    - Body: 仿宋 12 pt (小四), fixed 22 pt line spacing, first-line indent 2 chars.
    - Heading 1/2/3: 黑体 12 pt, left aligned, automatic color, no first-line indent.
    """
    for section in doc.sections:
        section.start_type = WD_SECTION_START.NEW_PAGE
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.8)

    styles = doc.styles

    normal = styles["Normal"]
    _set_style_font(normal, BODY_FONT, BODY_SIZE_PT)
    _set_style_paragraph(normal)

    title = styles["Title"]
    _set_style_font(title, TITLE_FONT, TITLE_SIZE_PT)
    _set_style_paragraph(
        title,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_chars=0,
        line_spacing_pt=32,
        space_before_pt=0,
        space_after_pt=18,
    )

    heading_1 = styles["Heading 1"]
    _set_style_font(heading_1, HEADING_FONT, BODY_SIZE_PT)
    _set_style_paragraph(
        heading_1,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_chars=0,
        line_spacing_pt=LINE_SPACING_PT,
        space_before_pt=12,
        space_after_pt=6,
    )

    heading_2 = styles["Heading 2"]
    _set_style_font(heading_2, HEADING_FONT, BODY_SIZE_PT)
    _set_style_paragraph(
        heading_2,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_chars=0,
        line_spacing_pt=LINE_SPACING_PT,
        space_before_pt=6,
        space_after_pt=6,
    )

    heading_3 = styles["Heading 3"]
    _set_style_font(heading_3, HEADING_FONT, BODY_SIZE_PT)
    _set_style_paragraph(
        heading_3,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
        first_line_indent_chars=0,
        line_spacing_pt=LINE_SPACING_PT,
        space_before_pt=6,
        space_after_pt=6,
    )

    for paragraph in doc.paragraphs:
        if paragraph_has_drawing(paragraph):
            set_image_paragraph_format(paragraph)
            continue
        if paragraph.style and paragraph.style.name.startswith("Heading"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = Pt(0)

            # 根据标题级别设置字体大小
            level = _heading_level(paragraph)
            size_map = {
                1: HEADING_LEVEL_1_SIZE_PT,
                2: HEADING_LEVEL_2_SIZE_PT,
                3: HEADING_LEVEL_3_SIZE_PT,
                4: HEADING_LEVEL_4_SIZE_PT,
            }
            size_pt = size_map.get(level, BODY_SIZE_PT)

            for run in paragraph.runs:
                set_run_font(run, HEADING_FONT, size_pt)
                _remove_run_color(run)
            continue
        if paragraph.style and paragraph.style.name == "Title":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                _remove_run_color(run)
            continue
        set_paragraph_format(paragraph)
        for run in paragraph.runs:
            set_run_font(run, BODY_FONT, BODY_SIZE_PT)

    for table in doc.tables:
        format_government_table(table)

    return doc


def add_government_title(doc: DocumentObject, text: str):
    paragraph = doc.add_paragraph(style="Title")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, TITLE_FONT, TITLE_SIZE_PT)
    return paragraph


def add_government_paragraph(doc: DocumentObject, text: str = ""):
    paragraph = doc.add_paragraph(text)
    set_paragraph_format(paragraph)
    for run in paragraph.runs:
        set_run_font(run, BODY_FONT, BODY_SIZE_PT)
    return paragraph


def add_government_heading(
    doc: DocumentObject,
    text: str,
    level: int = 1,
    *,
    auto_number: bool = True
) -> Paragraph:
    """
    添加政府公文格式的标题

    Args:
        doc: Word文档对象
        text: 标题文本
        level: 标题级别（1-4）
        auto_number: 是否自动添加序号（默认True）

    Returns:
        标题段落对象
    """
    level = min(max(level, 1), 4)

    # 自动添加序号（如果文本中不包含序号格式）
    if auto_number:
        # 检查文本是否已经包含序号格式
        has_number = False
        if level == 1 and re.match(r'^[一二三四五六七八九十]+、', text.strip()):
            has_number = True
        elif level == 2 and re.match(r'^（[一二三四五六七八九十]+）', text.strip()):
            has_number = True
        elif level == 3 and re.match(r'^\d+\.', text.strip()):
            has_number = True
        elif level == 4 and re.match(r'^（\d+）', text.strip()):
            has_number = True

        if not has_number:
            number = _get_heading_number(level)
            text = f"{number} {text}"

    # 创建标题段落
    paragraph = doc.add_heading(text, level=level)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Pt(0)

    # 根据级别设置字体大小
    size_map = {
        1: HEADING_LEVEL_1_SIZE_PT,
        2: HEADING_LEVEL_2_SIZE_PT,
        3: HEADING_LEVEL_3_SIZE_PT,
        4: HEADING_LEVEL_4_SIZE_PT,
    }
    size_pt = size_map.get(level, BODY_SIZE_PT)

    for run in paragraph.runs:
        set_run_font(run, HEADING_FONT, size_pt)

    return paragraph


def format_government_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                set_paragraph_format(
                    paragraph,
                    alignment=WD_ALIGN_PARAGRAPH.CENTER,
                    first_line_indent_chars=0,
                    line_spacing_pt=20,
                    space_before_pt=0,
                    space_after_pt=0,
                )
                for run in paragraph.runs:
                    set_run_font(run, BODY_FONT, 10.5, bold=(row_idx == 0))


def add_government_table(
    doc: DocumentObject,
    rows: Iterable[Iterable[object]],
    *,
    header: bool = True,
):
    data = [[str(cell) for cell in row] for row in rows]
    if not data:
        return None
    table = doc.add_table(rows=len(data), cols=len(data[0]))
    for row_idx, row_data in enumerate(data):
        for col_idx, value in enumerate(row_data):
            _set_table_cell_text(table.cell(row_idx, col_idx), value, bold=header and row_idx == 0)
    format_government_table(table)
    return table


def resolve_report_image_path(src: str, base_dir: str | Path | None = None) -> Path | None:
    """
    Resolve an HTML image src to a local file path.

    Supports absolute paths, paths relative to a report.html directory, file://
    URLs, local /api/reports/{report_id}/assets/... URLs, and assets/charts/...
    relative paths. Remote HTTP images are intentionally not downloaded here.
    """
    if not src:
        return None

    # 支持 assets/charts/ 和 assets/images/ 相对路径
    if src.startswith("assets/charts/") or src.startswith("assets/images/"):
        if base_dir:
            candidate = Path(base_dir) / src
            if candidate.exists() and candidate.is_file():
                return candidate.resolve()
        return None

    parsed = urlparse(src)
    if parsed.scheme in {"http", "https"}:
        return None
    if parsed.scheme == "file":
        candidate = Path(unquote(parsed.path))
    elif src.startswith("/api/reports/"):
        parts = src.split("/")
        # /api/reports/{report_id}/assets/{asset_path...}
        if len(parts) >= 6 and parts[4] == "assets":
            report_id = unquote(parts[3])
            asset_path = Path(*[unquote(part) for part in parts[5:]])
            candidate = DEFAULT_REFERENCE_DOCX.parents[1] / "reports" / report_id / "assets" / asset_path
        else:
            return None
    else:
        candidate = Path(unquote(src))
        if not candidate.is_absolute() and base_dir is not None:
            candidate = Path(base_dir) / candidate

    try:
        candidate = candidate.resolve()
    except OSError:
        return None
    return candidate if candidate.exists() and candidate.is_file() else None


def add_government_image(
    doc: DocumentObject,
    image_path: str | Path,
    *,
    caption: str | None = None,
    width_inches: float = 5.8,
):
    """Embed an image in a DOCX report with centered caption."""
    path = Path(image_path)
    if not path.exists() or not path.is_file():
        raise FileNotFoundError(f"Image not found: {path}")

    paragraph = doc.add_paragraph()
    set_image_paragraph_format(paragraph)
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(width_inches))

    if caption:
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(
            caption_paragraph,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line_indent_chars=0,
            line_spacing_pt=20,
            space_before_pt=0,
            space_after_pt=6,
        )
        caption_run = caption_paragraph.add_run(caption)
        set_run_font(caption_run, BODY_FONT, 10.5, color="666666")

    return paragraph


def sanitize_report_text(text: str | None) -> str:
    """Normalize text copied from HTML reports for stable DOCX/PDF rendering."""
    if not text:
        return ""
    normalized = str(text)
    normalized = normalized.replace("\xa0", " ")
    normalized = re.sub(r"[\u200b-\u200f\ufeff]", "", normalized)
    # Emoji and other non-BMP pictographs often render as tofu in LibreOffice
    # PDF previews. Formal DOCX exports do not need them.
    normalized = re.sub(r"[\U00010000-\U0010ffff]", "", normalized)
    normalized = re.sub(r"[\u2600-\u27bf]", "", normalized)
    normalized = re.sub(r"\s+", " ", normalized).strip()
    return normalized


def _add_centered_meta_paragraph(
    doc: DocumentObject,
    text: str,
    *,
    size_pt: float = 12,
    color: str | None = "666666",
):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(
        paragraph,
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        first_line_indent_chars=0,
        line_spacing_pt=20,
        space_before_pt=0,
        space_after_pt=4,
    )
    run = paragraph.add_run(text)
    set_run_font(run, BODY_FONT, size_pt, color=color)
    return paragraph


def _element_text(element) -> str:
    return sanitize_report_text(element.get_text(" ", strip=True))


def _is_inside_table(element) -> bool:
    return element.find_parent("table") is not None


def _is_inside_image_container(element) -> bool:
    parent = element.find_parent(class_="img-container")
    return parent is not None and parent is not element


def _caption_text(element) -> str:
    text = _element_text(element)
    return re.sub(r"^[▲△▴▵\s]+", "", text).strip()


def _handle_html_table(doc: DocumentObject, table_element) -> bool:
    rows = []
    for row_element in table_element.find_all("tr"):
        cells = row_element.find_all(["th", "td"], recursive=False)
        if not cells:
            continue
        row = [_element_text(cell) for cell in cells]
        rows.append(row)

    if not rows:
        return False

    max_cols = max(len(row) for row in rows)
    normalized_rows = [row + [""] * (max_cols - len(row)) for row in rows]
    add_government_table(doc, normalized_rows, header=True)
    return True


def _handle_html_image_container(
    doc: DocumentObject,
    container,
    *,
    base_dir: Path,
    width_inches: float,
    image_stats: dict,
) -> bool:
    image = container.find("img")
    if not image:
        return False

    src = image.get("src", "")
    image_path = resolve_report_image_path(src, base_dir)
    caption_element = container.find(class_="caption")
    caption = _caption_text(caption_element) if caption_element else sanitize_report_text(image.get("alt"))

    if image_path:
        add_government_image(doc, image_path, caption=caption or None, width_inches=width_inches)
        image_stats["embedded"] += 1
    else:
        image_stats["missing"].append(src)
        if caption:
            add_government_paragraph(doc, f"[图片缺失：{caption}]")
    return True


def _handle_html_node(
    doc: DocumentObject,
    node,
    *,
    base_dir: Path,
    width_inches: float,
    image_stats: dict,
    seen_text: set[str],
) -> None:
    from bs4 import NavigableString, Tag

    if isinstance(node, NavigableString):
        return
    if not isinstance(node, Tag):
        return
    if node.name in SKIP_HTML_TAGS or _is_inside_table(node) or _is_inside_image_container(node):
        return

    classes = set(node.get("class") or [])

    if node.name == "h1":
        return
    if node.name == "h2":
        text = _element_text(node)
        if text:
            add_government_heading(doc, text, level=1)
        return
    if node.name == "h3":
        text = _element_text(node)
        if text:
            add_government_heading(doc, text, level=2)
        return
    if node.name == "table":
        _handle_html_table(doc, node)
        return
    if "img-container" in classes:
        _handle_html_image_container(
            doc,
            node,
            base_dir=base_dir,
            width_inches=width_inches,
            image_stats=image_stats,
        )
        return
    if node.name == "img":
        image_path = resolve_report_image_path(node.get("src", ""), base_dir)
        caption = sanitize_report_text(node.get("alt"))
        if image_path:
            add_government_image(doc, image_path, caption=caption or None, width_inches=width_inches)
            image_stats["embedded"] += 1
        else:
            image_stats["missing"].append(node.get("src", ""))
        return
    if node.name == "p":
        text = _element_text(node)
        if text and text not in seen_text:
            add_government_paragraph(doc, text)
            seen_text.add(text)
        return
    if node.name in {"ul", "ol"}:
        for index, item in enumerate(node.find_all("li", recursive=False), start=1):
            text = _element_text(item)
            if not text or text in seen_text:
                continue
            prefix = f"{index}. " if node.name == "ol" else ""
            add_government_paragraph(doc, prefix + text)
            seen_text.add(text)
        return
    if node.name == "li":
        text = _element_text(node)
        if text and text not in seen_text:
            add_government_paragraph(doc, text)
            seen_text.add(text)
        return
    if "caption" in classes:
        return

    for child in node.children:
        _handle_html_node(
            doc,
            child,
            base_dir=base_dir,
            width_inches=width_inches,
            image_stats=image_stats,
            seen_text=seen_text,
        )


def convert_html_report_to_government_docx(
    html_path: str | Path,
    output_path: str | Path | None = None,
    *,
    width_inches: float = 5.8,
) -> dict:
    """
    Convert a local HTML report package to a government-style DOCX.

    The HTML file is treated as the source of truth. Local images referenced by
    relative paths, file URLs, or /api/reports/{report_id}/assets/... URLs are
    embedded into the Word document instead of left as placeholders.
    """
    from bs4 import BeautifulSoup

    source = Path(html_path)
    if not source.exists() or not source.is_file():
        raise FileNotFoundError(f"HTML report not found: {source}")

    output = Path(output_path) if output_path is not None else source.with_suffix(".docx")
    output.parent.mkdir(parents=True, exist_ok=True)

    soup = BeautifulSoup(source.read_text(encoding="utf-8", errors="replace"), "html.parser")
    doc = Document()
    apply_government_report_style(doc)

    header = soup.select_one(".header")
    title_text = ""
    if header:
        title_element = header.find("h1")
        title_text = _element_text(title_element) if title_element else ""
    if not title_text and soup.title:
        title_text = sanitize_report_text(soup.title.get_text(" ", strip=True))
    if title_text:
        add_government_title(doc, title_text)

    if header:
        subtitle = header.select_one(".subtitle")
        subtitle_text = _element_text(subtitle) if subtitle else ""
        if subtitle_text:
            _add_centered_meta_paragraph(doc, subtitle_text)

        meta_items = [_element_text(item) for item in header.select(".meta-item")]
        meta_items = [item for item in meta_items if item]
        if meta_items:
            _add_centered_meta_paragraph(doc, "；".join(meta_items), size_pt=10.5)

    image_stats = {"embedded": 0, "missing": []}
    seen_text: set[str] = set()
    roots = soup.select(".container > .section")
    if not roots and soup.body:
        roots = [soup.body]

    for root in roots:
        _handle_html_node(
            doc,
            root,
            base_dir=source.parent,
            width_inches=width_inches,
            image_stats=image_stats,
            seen_text=seen_text,
        )

    footer_items = []
    for footer in soup.select(".footer p"):
        text = _element_text(footer)
        if text:
            footer_items.append(text)
    if footer_items:
        for text in footer_items:
            _add_centered_meta_paragraph(doc, text, size_pt=9, color="888888")

    apply_government_report_style(doc)
    doc.save(str(output))
    return {
        "converted": True,
        "html_path": str(source),
        "path": str(output),
        "embedded_images": image_stats["embedded"],
        "missing_images": image_stats["missing"],
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
    }


def infer_report_html_path_for_docx(docx_path: str | Path) -> Path | None:
    """Infer the source report.html for a generated DOCX path."""
    path = Path(docx_path)
    candidates = [
        path.parent / "report.html",
        path.parent / path.stem / "report.html",
    ]
    if path.name == "report.docx":
        candidates.append(path.with_name("report.html"))
    for candidate in candidates:
        if candidate.exists() and candidate.is_file():
            return candidate
    return None


def replace_image_placeholders_in_docx(
    docx_path: str | Path,
    *,
    base_dir: str | Path | None = None,
    width_inches: float = 5.8,
) -> dict:
    """
    Replace paragraphs like "[图片: assets/images/a.png]" with embedded images.

    This is a defensive tool-layer cleanup for generated DOCX files. It lets
    agents generate a document naively while still producing a real Word file
    with images embedded when the referenced local image exists.
    """
    path = Path(docx_path)
    if not path.exists() or path.suffix.lower() != ".docx":
        return {"replaced": 0, "missing": [], "path": str(path)}

    image_base_dir = Path(base_dir) if base_dir is not None else path.parent
    doc = Document(str(path))
    pattern = re.compile(r"^\s*\[图片[:：]\s*(?P<src>[^\]]+?)\]\s*$")
    replaced = 0
    missing: list[str] = []

    for paragraph in doc.paragraphs:
        match = pattern.match(paragraph.text.strip())
        if not match:
            continue

        src = match.group("src").strip()
        image_path = resolve_report_image_path(src, image_base_dir)
        if not image_path:
            missing.append(src)
            continue

        paragraph.clear()
        set_image_paragraph_format(paragraph)
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))
        replaced += 1

    if replaced:
        doc.save(str(path))

    return {"replaced": replaced, "missing": missing, "path": str(path)}


def create_government_reference_docx(output_path: str | Path = DEFAULT_REFERENCE_DOCX) -> Path:
    """Create the default Quarto/Pandoc reference DOCX."""
    global _heading_counters, _last_heading_level

    # 重置计数器
    _heading_counters = {1: 0, 2: 0, 3: 0, 4: 0}
    _last_heading_level = 0

    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    apply_government_report_style(doc)
    add_government_title(doc, "政府报告标题")

    # 添加各级标题示例（自动添加序号）
    add_government_heading(doc, "一级标题", level=1, auto_number=True)
    add_government_paragraph(doc, "正文使用仿宋三号，首行缩进两个字符，固定行距二十八磅。")

    add_government_heading(doc, "二级标题", level=2, auto_number=True)
    add_government_paragraph(doc, "此段用于为 Pandoc/Quarto 提供默认正文样式。")

    add_government_heading(doc, "三级标题", level=3, auto_number=True)
    add_government_paragraph(doc, "三级标题内容示例。")

    add_government_heading(doc, "四级标题", level=4, auto_number=True)
    add_government_table(doc, [["指标", "数值"], ["示例", "100"]])

    apply_government_report_style(doc)
    doc.save(str(output))
    return output


def ensure_government_reference_docx(
    output_path: str | Path = DEFAULT_REFERENCE_DOCX,
) -> Path:
    """Return the default reference DOCX, refreshing it to current style rules."""
    output = Path(output_path)
    create_government_reference_docx(output)
    return output
