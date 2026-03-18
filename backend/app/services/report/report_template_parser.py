"""
报告模板解析器

解析DOCX模板，提取结构、占位符和数据需求
"""

from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
from docx.table import Table
import structlog

logger = structlog.get_logger()


class ReportTemplateParser:
    """
    报告模板解析器

    功能：
    - 解析DOCX文档结构
    - 识别段落和表格
    - 提取占位符
    - 分析数据需求
    """

    # 常见占位符模式
    PLACEHOLDER_PATTERNS = [
        ("{", "}"),          # {variable}
        ("{{", "}}"),        # {{variable}}
        ("${", "}"),         # ${variable}
        ("#", "#"),          # #variable#
        ("[", "]"),          # [variable]
        ("%", "%"),          # %variable%
    ]

    def __init__(self):
        self.logger = logger

    def parse_template(self, template_path: str) -> Dict[str, Any]:
        """
        解析模板文件

        Args:
            template_path: 模板文件路径

        Returns:
            解析结果
        """
        path = Path(template_path)
        if not path.exists():
            raise FileNotFoundError(f"模板文件不存在: {template_path}")

        if not path.suffix.lower() == ".docx":
            raise ValueError("只支持DOCX格式文件")

        doc = Document(str(path))

        result = {
            "template_path": template_path,
            "structure": self._extract_structure(doc),
            "placeholders": self._extract_placeholders(doc),
            "tables": self._extract_tables(doc),
            "metadata": {
                "file_name": path.name,
                "file_size": path.stat().st_size,
            }
        }

        self.logger.info(
            "template_parsed",
            path=template_path,
            paragraphs=result["structure"]["paragraph_count"],
            tables=result["structure"]["table_count"],
            placeholders=len(result["placeholders"])
        )

        return result

    def _extract_structure(self, doc: Document) -> Dict[str, Any]:
        """提取文档结构"""
        paragraphs = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                paragraphs.append({
                    "index": i,
                    "text": para.text.strip()[:200],
                    "style": para.style.name if para.style else "Normal"
                })

        return {
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables),
            "paragraphs": paragraphs[:10],
        }

    def _extract_placeholders(self, doc: Document) -> List[Dict[str, Any]]:
        """提取占位符"""
        placeholders = []

        # 从段落中提取
        for para_idx, para in enumerate(doc.paragraphs):
            text = para.text
            found = self._find_placeholders_in_text(
                text,
                "paragraph",
                paragraph_index=para_idx
            )
            placeholders.extend(found)

        # 从表格中提取
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text
                    found = self._find_placeholders_in_text(
                        text,
                        "table",
                        table_index=table_idx,
                        row=row_idx,
                        column=cell_idx
                    )
                    placeholders.extend(found)

        return placeholders

    def _find_placeholders_in_text(
        self,
        text: str,
        location_type: str,
        **location_info
    ) -> List[Dict[str, Any]]:
        """在文本中查找占位符"""
        placeholders = []

        for start_pattern, end_pattern in self.PLACEHOLDER_PATTERNS:
            if start_pattern in text:
                start = text.find(start_pattern)
                while start != -1:
                    end = text.find(end_pattern, start + len(start_pattern))
                    if end == -1:
                        break

                    placeholder = text[start:end + len(end_pattern)]
                    placeholders.append({
                        "type": location_type,
                        "placeholder": placeholder,
                        "context": text[max(0, start-30):min(len(text), end+len(end_pattern)+30)],
                        **location_info
                    })

                    # 查找下一个
                    start = text.find(start_pattern, end + len(end_pattern))

        return placeholders

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """提取表格信息"""
        tables = []

        for i, table in enumerate(doc.tables):
            tables.append({
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "has_headers": self._has_table_headers(table),
                "headers": self._extract_table_headers(table)
            })

        return tables

    def _has_table_headers(self, table: Table) -> bool:
        """判断表格是否有表头"""
        if not table.rows:
            return False

        first_row = table.rows[0]
        for cell in first_row.cells:
            if cell.paragraphs:
                for run in cell.paragraphs[0].runs:
                    if run.bold:
                        return True

        return False

    def _extract_table_headers(self, table: Table) -> List[str]:
        """提取表格表头"""
        if not table.rows:
            return []

        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip())

        return headers

    def analyze_data_requirements(
        self,
        parsed_template: Dict[str, Any],
        user_description: Optional[str] = None
    ) -> List[Dict[str, Any]]:
        """
        分析数据需求

        Args:
            parsed_template: 解析后的模板
            user_description: 用户描述

        Returns:
            数据需求列表
        """
        requirements = []

        # 从占位符推断
        placeholder_types = {
            "时间": ["time", "date", "日期", "时间", "start_time", "end_time"],
            "城市": ["city", "城市", "地区"],
            "站点": ["station", "站点"],
            "污染物": ["pollutant", "污染物", "pm25", "pm10", "o3", "no2", "so2", "co"],
        }

        for ph in parsed_template.get("placeholders", []):
            placeholder_lower = ph["placeholder"].lower()
            for req_type, keywords in placeholder_types.items():
                if any(kw in placeholder_lower for kw in keywords):
                    requirements.append({
                        "type": req_type,
                        "source": "placeholder",
                        "placeholder": ph["placeholder"]
                    })
                    break

        # 从表格推断
        for table in parsed_template.get("tables", []):
            if table["rows"] > 1:
                requirements.append({
                    "type": "表格数据",
                    "source": "table",
                    "rows": table["rows"],
                    "columns": table["columns"],
                    "headers": table.get("headers", [])
                })

        # 去重
        seen = set()
        unique_requirements = []
        for req in requirements:
            key = (req["type"], str(req.get("placeholder", "")))
            if key not in seen:
                seen.add(key)
                unique_requirements.append(req)

        return unique_requirements
