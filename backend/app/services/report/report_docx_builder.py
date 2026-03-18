"""
报告DOCX构建器

使用python-docx构建DOCX报告文档
"""

from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import structlog

logger = structlog.get_logger()


class ReportDocxBuilder:
    """
    报告DOCX构建器

    功能：
    - 创建新文档或基于模板创建
    - 替换占位符
    - 插入数据表格
    - 添加元信息
    """

    def __init__(self):
        self.logger = logger
        self.doc = None

    def create_document(
        self,
        template_path: Optional[str] = None,
        title: Optional[str] = None
    ) -> Document:
        """
        创建文档

        Args:
            template_path: 模板路径（可选）
            title: 文档标题（可选，仅用于新文档）

        Returns:
            Document对象
        """
        if template_path and Path(template_path).exists():
            self.doc = Document(str(template_path))
            self.logger.info("loaded_template", path=template_path)
        else:
            self.doc = Document()
            if title:
                self._add_title(title)

        return self.doc

    def _add_title(self, title: str):
        """添加文档标题"""
        heading = self.doc.add_heading(title, 0)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 设置字体
        for run in heading.runs:
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(18)

    def replace_placeholders(self, replacements: Dict[str, str]):
        """
        替换文档中的占位符

        Args:
            replacements: 替换规则 {placeholder: replacement_value}
        """
        if not self.doc:
            raise ValueError("请先创建文档")

        # 替换段落中的占位符
        for para in self.doc.paragraphs:
            self._replace_in_paragraph(para, replacements)

        # 替换表格中的占位符
        for table in self.doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_in_paragraph(para, replacements)

        self.logger.info(
            "placeholders_replaced",
            count=len(replacements)
        )

    def _replace_in_paragraph(
        self,
        paragraph,
        replacements: Dict[str, str]
    ):
        """在段落中替换占位符"""
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                # 清空段落
                paragraph.clear()

                # 添加替换后的文本
                run = paragraph.add_run(value)

                # 保留原有格式（如果可能）
                # 注意：python-docx的格式处理较复杂，这里做简化处理

    def add_table(
        self,
        data: List[List[str]],
        title: Optional[str] = None,
        style: str = 'Light Grid Accent 1'
    ):
        """
        添加表格

        Args:
            data: 表格数据（二维列表）
            title: 表格标题（可选）
            style: 表格样式
        """
        if not self.doc:
            raise ValueError("请先创建文档")

        if not data:
            return

        # 添加标题
        if title:
            heading = self.doc.add_heading(title, level=2)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 创建表格
        table = self.doc.add_table(rows=len(data), cols=len(data[0]))
        table.style = style

        # 填充数据
        for i, row_data in enumerate(data):
            row = table.rows[i]
            for j, cell_value in enumerate(row_data):
                cell = row.cells[j]
                cell.text = str(cell_value)

                # 表头加粗
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True

        self.logger.info(
            "table_added",
            rows=len(data),
            columns=len(data[0])
        )

    def add_paragraph(
        self,
        text: str,
        style: Optional[str] = None,
        bold: bool = False,
        alignment: Optional[str] = None
    ):
        """
        添加段落

        Args:
            text: 文本内容
            style: 段落样式（可选）
            bold: 是否加粗
            alignment: 对齐方式（left/center/right）
        """
        if not self.doc:
            raise ValueError("请先创建文档")

        para = self.doc.add_paragraph(text)

        if style:
            para.style = style

        if bold:
            for run in para.runs:
                run.bold = True

        if alignment:
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            }
            para.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def add_heading(
        self,
        text: str,
        level: int = 1,
        alignment: Optional[str] = None
    ):
        """
        添加标题

        Args:
            text: 标题文本
            level: 标题级别（1-3）
            alignment: 对齐方式
        """
        if not self.doc:
            raise ValueError("请先创建文档")

        heading = self.doc.add_heading(text, level=level)

        if alignment:
            alignment_map = {
                "left": WD_PARAGRAPH_ALIGNMENT.LEFT,
                "center": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "right": WD_PARAGRAPH_ALIGNMENT.RIGHT,
            }
            heading.alignment = alignment_map.get(alignment, WD_PARAGRAPH_ALIGNMENT.LEFT)

    def add_generation_info(
        self,
        time_range: Optional[Dict[str, str]] = None
    ):
        """
        添加生成信息

        Args:
            time_range: 数据时间范围
        """
        if not self.doc:
            raise ValueError("请先创建文档")

        self.doc.add_paragraph()  # 空行

        info = self.doc.add_paragraph()
        info.add_run("---\n").bold = True
        info.add_run(f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

        if time_range:
            if time_range.get("start_time") and time_range.get("end_time"):
                info.add_run(
                    f"数据时间范围: {time_range['start_time']} 至 {time_range['end_time']}\n"
                )

        info.add_run("生成工具: 大气污染溯源分析系统\n")
        info.add_run("---")

    def save(self, output_path: str):
        """
        保存文档

        Args:
            output_path: 输出文件路径
        """
        if not self.doc:
            raise ValueError("请先创建文档")

        output_file = Path(output_path)
        output_file.parent.mkdir(parents=True, exist_ok=True)

        self.doc.save(str(output_file))

        self.logger.info(
            "document_saved",
            path=output_path,
            size=output_file.stat().st_size if output_file.exists() else 0
        )

    def build_report(
        self,
        template_path: Optional[str] = None,
        time_range: Optional[Dict[str, str]] = None,
        data_ids: Optional[List[str]] = None,
        data_matches: Optional[Dict[str, Any]] = None,
        tables: Optional[List[Dict[str, Any]]] = None,
        output_path: Optional[str] = None,
        title: Optional[str] = None
    ) -> str:
        """
        构建完整报告

        Args:
            template_path: 模板路径
            time_range: 时间范围
            data_ids: 数据ID列表
            data_matches: 数据匹配规则
            tables: 表格数据列表
            output_path: 输出路径
            title: 报告标题

        Returns:
            输出文件路径
        """
        # 创建文档
        self.create_document(template_path, title)

        # 替换占位符
        if data_matches:
            replacements = self._prepare_replacements(time_range, data_matches)
            self.replace_placeholders(replacements)

        # 添加表格
        if tables:
            for table_info in tables:
                self.add_table(
                    data=table_info.get("data", []),
                    title=table_info.get("title"),
                    style=table_info.get("style", "Light Grid Accent 1")
                )

        # 添加生成信息
        self.add_generation_info(time_range)

        # 保存文档
        if output_path:
            self.save(output_path)
            return output_path

        return ""

    def _prepare_replacements(
        self,
        time_range: Optional[Dict[str, str]],
        data_matches: Dict[str, Any]
    ) -> Dict[str, str]:
        """准备替换规则"""
        replacements = {}

        # 添加时间范围替换
        if time_range:
            replacements.update({
                "{start_time}": time_range.get("start_time", ""),
                "{end_time}": time_range.get("end_time", ""),
                "{{start_time}}": time_range.get("start_time", ""),
                "{{end_time}}": time_range.get("end_time", ""),
                "${start_time}": time_range.get("start_time", ""),
                "${end_time}": time_range.get("end_time", ""),
            })

        # 添加数据匹配
        for key, value in data_matches.items():
            # 支持多种占位符格式
            placeholders = [
                f"{{{key}}}",
                f"{{{{{key}}}}}",
                f"${{{key}}}",
            ]
            for ph in placeholders:
                replacements[ph] = str(value)

        return replacements
