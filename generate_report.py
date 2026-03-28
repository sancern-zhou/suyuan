from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('美以伊朗冲突时间脉络报告', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加段落
doc.add_paragraph('报告生成时间：2026年3月28日')
doc.add_paragraph('')

# 摘要部分
doc.add_heading('一、摘要', level=1)
doc.add_paragraph('本报告基于网络搜索结果，梳理了2026年美以与伊朗冲突的关键事件、时间脉络、背景及影响。冲突始于2026年2月28日，美以联合对伊朗发动大规模突袭，导致伊朗最高领袖身亡，随后伊朗发动反击，冲突持续近一个月。')
doc.add_paragraph('')

# 背景部分
doc.add_heading('二、冲突背景', level=1)
doc.add_paragraph('1. 历史恩怨：伊朗与美国自1979年伊斯兰革命后关系恶化，1980年断交并实施经济制裁。2001年后，美国将伊朗列为“支持恐怖主义国家”，矛盾持续加深。')
doc.add_paragraph('2. 近期紧张：2017年后，美国对伊政策趋强硬，2025年6月空袭伊朗核设施，矛盾急剧升级。')
doc.add_paragraph('3. 直接导火索：2026年2月28日，美以未经安理会授权，袭击并杀害伊朗最高领导人，蓄意挑起战争。')
doc.add_paragraph('')

# 时间脉络部分
doc.add_heading('三、时间脉络', level=1)
events = [
    ('2026年2月28日', '美以联合对伊朗发动大规模突袭，出动约200架战斗机和“战斧”巡航导弹，袭击伊朗境内约30个目标，包括总统府。伊朗最高领袖哈梅内伊在袭击中身亡。'),
    ('2026年3月1日', '伊朗民众在德黑兰示威，高呼反美反以口号。伊朗伊斯兰革命卫队发起代号“真实承诺-4”的反击行动，用导弹和无人机打击美国在巴林、卡塔尔、阿联酋的军事基地及以色列目标。伊朗首次使用“法塔赫-2”高超音速导弹袭击美军基地。'),
    ('2026年3月上旬', '冲突持续，伊朗每日发射弹道导弹高达百枚，但随后因导弹发射架被毁、库存下降、生产线停滞，火力输出大幅缩水。'),
    ('2026年3月26日', '伊朗武装部队发言人评估，美军至少被伊朗打死800人。冲突已持续近一个月。'),
    ('2026年3月28日', '冲突持续29天，美以未能实现“速决剧本”，双方陷入战略博弈。')
]

for date, description in events:
    p = doc.add_paragraph()
    p.add_run(f'{date}：').bold = True
    p.add_run(description)

doc.add_paragraph('')

# 影响部分
doc.add_heading('四、影响分析', level=1)
doc.add_paragraph('1. 地区格局：冲突加剧中东紧张局势，影响地区力量平衡，可能引发更大范围的冲突。')
doc.add_paragraph('2. 全球能源安全：中东是全球石油供应关键地区，冲突可能影响石油生产和运输，推高全球油价。')
doc.add_paragraph('3. 军事战略：美以未能速决，暴露现代战争中工业供应链的重要性，伊朗导弹库存快速消耗显示工业体系的关键作用。')
doc.add_paragraph('4. 国际反应：联合国安理会未授权此次袭击，国际社会对冲突升级表示关切。')
doc.add_paragraph('')

# 结论部分
doc.add_heading('五、结论', level=1)
doc.add_paragraph('美以伊朗冲突始于2026年2月28日，持续近一个月，造成重大人员伤亡和地区不稳定。冲突背景复杂，涉及历史恩怨和近期紧张局势。影响深远，涉及地区格局、全球能源安全和军事战略。未来冲突走向取决于双方战略调整和国际社会介入。')
doc.add_paragraph('')
doc.add_paragraph('报告结束。')

# 保存文档
file_path = os.path.join(os.path.expanduser('~'), '美以伊朗冲突时间脉络报告.docx')
doc.save(file_path)
print(f'报告已保存至：{file_path}')