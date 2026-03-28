from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

# 创建文档
doc = Document()

# 标题
title = doc.add_heading('美以伊朗冲突时间脉络报告（全面版）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加段落
doc.add_paragraph('报告生成时间：2026年3月28日')
doc.add_paragraph('')

# 摘要部分
doc.add_heading('一、摘要', level=1)
doc.add_paragraph('本报告基于网络搜索结果，全面梳理了2026年美以与伊朗冲突的详细背景、时间脉络、国际反应、军事细节及经济影响。冲突始于2026年2月28日，美以联合对伊朗发动代号“史诗怒火”/“咆哮的狮子”的大规模军事打击，导致伊朗最高领导人身亡，随后伊朗发动“真实承诺-4”反击行动，冲突持续近一个月，对地区格局和全球能源安全产生深远影响。')
doc.add_paragraph('')

# 详细历史背景
doc.add_heading('二、详细历史背景', level=1)
doc.add_paragraph('1. 二战期间：伊朗遭美英苏三国占领，这是美国干涉伊朗的滥觞。')
doc.add_paragraph('2. 巴列维王朝时期：伊朗巴列维国王执政期间，其内政外交均被美国控制，经济遭掠夺，文化被渗透。')
doc.add_paragraph('3. 1979年伊斯兰革命：伊朗学生占领美国大使馆并扣留52名人质长达444天，表达对美国支持巴列维王朝并长期干涉伊朗内政的强烈不满。')
doc.add_paragraph('4. 1980年断交与制裁：1980年4月7日，美国宣布同伊朗断交，并正式对伊朗实施经济制裁，包括石油禁运、驱逐在美生活的伊朗人、冻结伊朗政府资产和投资等，美伊成为死敌。')
doc.add_paragraph('5. 2001年后矛盾加深：“9·11”事件后，美国政府将伊朗列为“支持恐怖主义国家”，并公开指责伊朗试图发展大规模杀伤性武器，是威胁世界和平的“邪恶轴心”之一。')
doc.add_paragraph('6. 2017年后政策趋硬：美国政府对伊政策更趋强硬，外交上打造地区反伊朗联盟，经济上极限施压，军事上加强威慑，炸死伊朗革命卫队苏莱曼尼将军，与以色列联手打击伊朗在地区的重要盟友。')
doc.add_paragraph('7. 2025年6月冲突升级：美国空袭伊朗三处核设施，美伊矛盾急剧升级。')
doc.add_paragraph('8. 2026年2月28日直接冲突：在伊美谈判取得进展背景下，美国和以色列未经安理会授权，袭击并杀害伊朗最高领导人，蓄意挑起对伊朗的战争。')
doc.add_paragraph('')

# 时间脉络
doc.add_heading('三、时间脉络（详细节点）', level=1)
doc.add_paragraph('**2026年2月28日**')
doc.add_paragraph('• 9时50分（伊朗时间）：以色列出动约200架F-35I隐身战机，在美军电子战飞机与卫星情报系统支援下，对伊朗德黑兰、伊斯法罕、库姆、布什尔等多省份核心目标发动首轮空袭。')
doc.add_paragraph('• 同时，美军从“福特”号、“林肯”号双航母战斗群发射“战斧”巡航导弹，出动B-2隐形轰炸机投放GBU-57巨型钻地弹，同步打击伊朗核设施、军事指挥中枢与导弹阵地。')
doc.add_paragraph('• 打击目标覆盖三大类约30处关键设施：政权中枢（德黑兰总统府、最高领袖官邸周边、情报总部）、核设施（福尔多地下核设施、纳坦兹铀浓缩工厂、布什尔核电站）、军事基础设施（革命卫队总部、弹道导弹发射基地、防空系统指挥中心）。')
doc.add_paragraph('• 伊朗最高领袖哈梅内伊在袭击中身亡。')
doc.add_paragraph('')
doc.add_paragraph('**2026年3月1日**')
doc.add_paragraph('• 伊朗民众在德黑兰的恩格拉布（革命）广场聚集举行示威活动，高呼反美和反以色列口号。')
doc.add_paragraph('• 伊朗伊斯兰革命卫队发起代号“真实承诺-4”的反击行动，用导弹和无人机打击美国在巴林、卡塔尔、阿联酋的军事基地及以色列目标。')
doc.add_paragraph('• 伊朗首次使用“法塔赫-2”高超音速导弹袭击美军基地。')
doc.add_paragraph('')
doc.add_paragraph('**2026年3月上旬**')
doc.add_paragraph('• 冲突持续，伊朗每日发射弹道导弹高达百枚，但随后因导弹发射架被毁、库存下降、生产线停滞，火力输出大幅缩水。')
doc.add_paragraph('• 美军迄今已袭击超过1万个伊朗军事目标。')
doc.add_paragraph('')
doc.add_paragraph('**2026年3月26日**')
doc.add_paragraph('• 伊朗武装部队发言人评估，美军至少被伊朗打死800人。')
doc.add_paragraph('• 伊朗已正式回应美国提出的包含15点内容的停火协议，回应已通过斡旋的国家正式发出。')
doc.add_paragraph('• 伊朗伊斯兰革命卫队发动“真实承诺-4”行动第82波攻势，打击美军位于巴林和科威特境内的军事基地。')
doc.add_paragraph('')
doc.add_paragraph('**2026年3月28日**')
doc.add_paragraph('• 冲突持续29天，美以未能实现“速决剧本”，双方陷入战略博弈。')
doc.add_paragraph('• 伊朗陆军司令警告，对敌人而言，地面战“更加危险且代价更大”。')
doc.add_paragraph('')

# 国际反应
doc.add_heading('四、国际反应', level=1)
doc.add_paragraph('1. 联合国安理会：美以联合军事打击系未经联合国安理会授权、无正当防卫依据的非法使用武力行为，严重侵犯伊朗主权与领土完整，已呈现出“侵略行为”的特征。')
doc.add_paragraph('2. 国际社会：国际社会对冲突升级表示关切，认为该行动是对二战后形成的国际秩序的野蛮冲击。')
doc.add_paragraph('3. 海湾国家：长期来看受益于以伊两个战略对手的互斗，缺少止战的意愿。')
doc.add_paragraph('4. 中国等国家：关注冲突对地区稳定和全球能源安全的影响。')
doc.add_paragraph('')

# 军事细节
doc.add_heading('五、军事细节', level=1)
doc.add_paragraph('• 美以联合行动代号：“史诗怒火”（美军）与“咆哮的狮子”（以军）。')
doc.add_paragraph('• 以色列出动：约200架F-35I隐身战机。')
doc.add_paragraph('• 美军出动：“福特”号、“林肯”号双航母战斗群，B-2隐形轰炸机，发射“战斧”巡航导弹、投放GBU-57巨型钻地弹。')
doc.add_paragraph('• 伊朗反击：代号“真实承诺-4”，使用导弹和无人机打击目标，首次使用“法塔赫-2”高超音速导弹。')
doc.add_paragraph('• 伊朗警告：若敌方试图对伊朗岛屿或本土发起地面行动，伊朗将在曼德海峡开辟新战线。')
doc.add_paragraph('')

# 经济影响
doc.add_heading('六、经济影响', level=1)
doc.add_paragraph('1. 全球能源市场：冲突导致全球能源市场剧烈震荡，中东是全球石油供应关键地区，冲突可能影响石油生产和运输，推高全球油价。')
doc.add_paragraph('2. 霍尔木兹海峡：即使解封，日通行量可能较冲突前下降3-4成，类似两伊战争期间袭船期间的下降幅度。')
doc.add_paragraph('3. 美国武器转运：随着对伊朗战事消耗大量弹药，美国防部正考虑将原本提供给乌克兰的武器转运至中东。')
doc.add_paragraph('4. 资产波动：大类资产影响显而易见，近期资产波动巨大，相关性显著提升。')
doc.add_paragraph('')

# 结论
doc.add_heading('七、结论', level=1)
doc.add_paragraph('美以伊朗冲突始于2026年2月28日，持续近一个月，造成重大人员伤亡和地区不稳定。冲突背景复杂，涉及历史恩怨和近期紧张局势。国际反应强烈，军事细节丰富，经济影响深远。未来冲突走向取决于双方战略调整和国际社会介入，协议可能最早在4月中下旬签订。')
doc.add_paragraph('')
doc.add_paragraph('报告结束。')

# 保存文档
file_path = os.path.join(os.path.expanduser('~'), '美以伊朗冲突时间脉络报告（全面版）.docx')
doc.save(file_path)
print(f'报告已保存至：{file_path}')