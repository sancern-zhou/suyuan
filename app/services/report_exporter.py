"""
报告导出服务

使用前端截图方案：
- 图表截图由前端ECharts的getDataURL()完成
- 后端只负责组装HTML并转换为PDF/Word格式
"""

from typing import Dict, Any, List, Optional
import structlog
import re
from datetime import datetime

logger = structlog.get_logger()


class EchartsPlaceholderHandler:
    """
    ECharts占位符处理器 - 将占位符转换为实际内容

    【废弃】此类仅为处理遗留内容保留，新代码应使用URL直接输出方案：
    - 旧方案：LLM生成 [ECHARTS_PLACEHOLDER:xxx] 占位符 → 后端替换为图片
    - 新方案：LLM直接生成 ![标题](/api/image/xxx) Markdown URL → 前端直接渲染
    """

    def __init__(self):
        self.placeholder_pattern = re.compile(r'\[ECHARTS_PLACEHOLDER:([^\]]+)\]')

    def process_section_content(
        self,
        content: str,
        charts: List[Dict[str, Any]]
    ) -> str:
        """
        处理章节内容中的ECharts占位符

        Args:
            content: 原始章节内容
            charts: 图表列表（用于查找对应的图表信息）

        Returns:
            处理后的内容（占位符被替换为实际内容）
        """
        if not content:
            return content

        # 建立图表ID索引
        chart_by_id = {chart.get("id"): chart for chart in charts}

        def replace_placeholder(match):
            chart_id = match.group(1)
            chart = chart_by_id.get(chart_id)

            if chart:
                title = chart.get("title", "图表")
                chart_type = chart.get("type", "chart")

                # 检查是否有图片数据
                image_data = self._extract_chart_image_data(chart)

                if image_data and image_data.startswith("data:image"):
                    # 有图片数据，替换为实际图片
                    return f'![{title}]({image_data})'
                else:
                    # 没有图片数据，替换为提示信息
                    return f'**{title}** ({chart_type}，交互式图表)'

            # 占位符对应的图表不存在
            return f'**[图表 {chart_id} 已移除]**'

        # 替换所有占位符
        processed_content = self.placeholder_pattern.sub(replace_placeholder, content)

        return processed_content

    def _extract_chart_image_data(self, chart: Dict[str, Any]) -> Optional[str]:
        """从图表对象中提取图片数据"""
        # 优先级顺序：从最可靠到最不可靠
        image_sources = [
            ("unified_image_data", "统一数据字段"),
            ("preview_image", "前端截图"),
            ("data", "Executor数据"),
            ("payload.data", "Payload数据"),
            ("payload.image", "Payload图片"),
            ("image", "直接图片字段")
        ]

        for field_path, source_name in image_sources:
            if "." in field_path:
                # 处理嵌套字段（如payload.data）
                parts = field_path.split(".")
                value = chart
                for part in parts:
                    if isinstance(value, dict) and part in value:
                        value = value[part]
                    else:
                        value = None
                        break
            else:
                # 处理直接字段
                value = chart.get(field_path)

            # 检查是否是有效的图片数据
            if value and isinstance(value, str) and value.startswith("data:image"):
                logger.debug(
                    "chart_image_extracted",
                    chart_id=chart.get("id"),
                    source=source_name,
                    field=field_path
                )
                return value

        return None


class ChartNumberingSystem:
    """
    图表编号系统 - 基于引用而非位置
    """

    def __init__(self):
        self.used_numbers = set()
        self.next_number = 1

    def assign_number(
        self,
        chart: Dict[str, Any],
        fallback_title: str = ""
    ) -> int:
        """
        为图表分配编号

        Args:
            chart: 图表对象
            fallback_title: fallback标题

        Returns:
            分配的编号
        """
        # 优先使用文本引用编号
        if "text_reference_number" in chart:
            number = chart["text_reference_number"]
            if number not in self.used_numbers:
                self.used_numbers.add(number)
                return number

        # 找到下一个可用编号
        while self.next_number in self.used_numbers:
            self.next_number += 1

        self.used_numbers.add(self.next_number)
        return self.next_number

    def reset(self):
        """重置编号系统"""
        self.used_numbers.clear()
        self.next_number = 1


class ReportExporter:
    """报告导出服务"""

    def __init__(self):
        self.placeholder_handler = EchartsPlaceholderHandler()
        self._used_descriptions = set()  # 用于容错分配

    def _normalize_chart_data(self, chart: Dict[str, Any]) -> Dict[str, Any]:
        """
        标准化图表数据 - 从多个可能的字段读取图片数据

        Args:
            chart: 原始图表数据

        Returns:
            标准化后的图表数据，包含unified_image_data字段
        """
        normalized = chart.copy()

        # 优先级顺序：从最可靠到最不可靠
        image_sources = [
            ("preview_image", "前端截图"),
            ("data", "Executor数据"),
            ("payload.data", "Payload数据"),
            ("payload.image", "Payload图片"),
            ("image", "直接图片字段")
        ]

        image_data = None
        source_used = None

        for field_path, source_name in image_sources:
            if "." in field_path:
                # 处理嵌套字段（如payload.data）
                parts = field_path.split(".")
                value = normalized
                for part in parts:
                    if isinstance(value, dict) and part in value:
                        value = value[part]
                    else:
                        value = None
                        break
            else:
                # 处理直接字段
                value = normalized.get(field_path)

            # 检查是否是有效的图片数据
            if value and isinstance(value, str) and value.startswith("data:image"):
                image_data = value
                source_used = source_name
                break

        # 添加统一字段
        normalized["unified_image_data"] = image_data
        normalized["_image_source"] = source_used  # 用于调试

        return normalized

    async def generate(
        self,
        format: str,
        report_content: Optional[Dict[str, Any]],
        charts: List[Dict[str, Any]]
    ) -> bytes:
        """
        生成报告
        
        Args:
            format: 导出格式 (pdf/docx/html)
            report_content: 报告内容（来自报告专家）
            charts: 图表列表（包含前端截图的base64）
            
        Returns:
            bytes: 生成的文件内容
        """
        logger.info(
            "generating_report",
            format=format,
            chart_count=len(charts),
            has_report_content=report_content is not None
        )
        
        # 1. 组装HTML
        html_content = self._compose_html(report_content, charts)
        
        # 2. 转换格式
        if format == "pdf":
            return await self._html_to_pdf(html_content)
        elif format == "docx":
            return await self._html_to_docx(html_content, charts)
        else:
            return html_content.encode("utf-8")
    
    def _compose_html(
        self,
        report_content: Optional[Dict[str, Any]],
        charts: List[Dict[str, Any]]
    ) -> str:
        """组装HTML报告"""
        
        # 解析报告内容
        title = "大气污染溯源分析报告"
        summary = ""
        detail_html = ""
        conclusions = []
        recommendations = []
        confidence = 85
        
        if report_content:
            title = report_content.get("title", title)
            summary = report_content.get("summary", "")
            confidence_val = report_content.get("confidence", 0.85)
            confidence = int(confidence_val * 100) if confidence_val <= 1 else int(confidence_val)
            conclusions = report_content.get("conclusions", [])
            recommendations = report_content.get("recommendations", [])
            
            # 从sections获取Markdown内容
            sections = report_content.get("sections", [])
            if sections and isinstance(sections[0], dict):
                md_content = sections[0].get("markdown_content", "")
                if md_content:
                    detail_html = self._markdown_to_html(md_content)
        
        # 提取各章节Markdown（如果存在）
        weather_md = ""
        component_md = ""
        conclusion_md = ""
        detail_md = ""
        
        if report_content:
            sections = report_content.get("sections", [])
            if isinstance(sections, list):
                for section in sections:
                    if not isinstance(section, dict):
                        continue
                    section_type = section.get("type") or section.get("name") or ""
                    content = section.get("markdown_content", "")
                    if not content:
                        continue
                    
                    # 解析预设标识
                    parsed_content = self._parse_section_markers(content, section_type)
                    
                    lower_type = str(section_type).lower()
                    if "weather" in lower_type or "气象" in str(section_type):
                        weather_md = parsed_content
                        logger.info("weather_section_extracted", content_length=len(parsed_content), has_markers=any(m in content for m in ["[WEATHER_SECTION_START]", "[WEATHER_SECTION_END]"]))
                    elif "component" in lower_type or "组分" in str(section_type):
                        component_md = parsed_content
                        logger.info("component_section_extracted", content_length=len(parsed_content), has_markers=any(m in content for m in ["[COMPONENT_SECTION_START]", "[COMPONENT_SECTION_END]"]))
                    elif "conclusion" in lower_type or "结论" in str(section_type) or "建议" in str(section_type):
                        conclusion_md = parsed_content
                        logger.info("conclusion_section_extracted", content_length=len(parsed_content), has_markers=any(m in content for m in ["[CONCLUSION_SECTION_START]", "[CONCLUSION_SECTION_END]"]))
                    elif "detail" in lower_type or "详细" in str(section_type):
                        detail_md = parsed_content
                    elif not section_type:
                        # 没有type字段的，作为详细分析
                        detail_md = parsed_content
        
        # 【修复】在生成HTML之前，先从文本中移除图表解析部分（因为图表解析会在图表旁边显示）
        # 这样可以避免文本和图表分离的问题
        weather_md_clean = self._remove_chart_sections(weather_md) if weather_md else ""
        component_md_clean = self._remove_chart_sections(component_md) if component_md else ""

        # 【关键修复】处理ECharts占位符
        # 标准化图表数据
        normalized_charts = [self._normalize_chart_data(chart) for chart in charts]

        # 处理各章节的占位符
        if weather_md:
            weather_md_processed = self.placeholder_handler.process_section_content(
                weather_md, normalized_charts
            )
            weather_md_clean = self._remove_chart_sections(weather_md_processed)
            weather_detail_html = self._markdown_to_html(weather_md_clean)

        if component_md:
            component_md_processed = self.placeholder_handler.process_section_content(
                component_md, normalized_charts
            )
            component_md_clean = self._remove_chart_sections(component_md_processed)
            component_detail_html = self._markdown_to_html(component_md_clean)

        conclusion_detail_html = self._markdown_to_html(conclusion_md) if conclusion_md else ""

        # 生成图表HTML（按专家归类，并匹配章节内容中的图表引用）
        # 注意：这里传入原始的weather_md和component_md，用于提取图表解析文本
        charts_html = self._generate_grouped_charts_html(
            normalized_charts,
            weather_content=weather_md,
            component_content=component_md,
            other_content=""
        )
        summary_html = self._markdown_to_html(summary) if summary else ""
        
        # 生成结论HTML（如果从sections中提取到了结论内容，优先使用）
        conclusion_section = ""
        if conclusion_detail_html:
            # 使用LLM生成的结论章节
            conclusion_section = f"""
            <div class="section">
                <h2>结论与建议</h2>
                <div class="detail-content">{conclusion_detail_html}</div>
            </div>
            """
        else:
            # 降级：使用提取的conclusions和recommendations
            conclusions_html = ""
            if conclusions:
                items = "".join(f"<li>{c}</li>" for c in conclusions)
                conclusions_html = f"""
                <div class="conclusions">
                    <h3>主要结论</h3>
                    <ul>{items}</ul>
                </div>
                """
            
            recommendations_html = ""
            if recommendations:
                items = "".join(f"<li>{r}</li>" for r in recommendations)
                recommendations_html = f"""
                <div class="recommendations">
                    <h3>控制建议</h3>
                    <ul>{items}</ul>
                </div>
                """
            
            if conclusions_html or recommendations_html:
                conclusion_section = f"""
                <div class="section">
                    <h2>结论与建议</h2>
                    {conclusions_html}
                    {recommendations_html}
                </div>
                """
        
        # 生成详细分析HTML
        detail_section = ""
        if detail_html or detail_md:
            detail_content = self._markdown_to_html(detail_md) if detail_md else detail_html
            if detail_content:
                detail_section = f"""
                <div class="section">
                    <h2>详细分析</h2>
                    <div class="detail-content">{detail_content}</div>
                </div>
                """
        
        # 气象分析章节（文字+图表）
        weather_section = ""
        if weather_detail_html or "weather" in charts_html:
            weather_section = f"""
            <div class="section">
                <h2>气象分析</h2>
                {f'<div class="detail-content">{weather_detail_html}</div>' if weather_detail_html else ''}
                {charts_html.get("weather", "")}
            </div>
            """
        
        # 组分分析章节（文字+图表）
        component_section = ""
        if component_detail_html or "component" in charts_html:
            component_section = f"""
            <div class="section">
                <h2>组分分析</h2>
                {f'<div class="detail-content">{component_detail_html}</div>' if component_detail_html else ''}
                {charts_html.get("component", "")}
            </div>
            """
        
        # 其他图表
        other_section = ""
        if "other" in charts_html:
            other_section = f"""
            <div class="section">
                <h2>其他分析图表</h2>
                {charts_html.get("other", "")}
            </div>
            """
        
        # 组装完整HTML（顺序：结论 -> 气象 -> 组分 -> 其他 -> 详细）
        generated_at = datetime.now().strftime("%Y-%m-%d %H:%M")
        
        html = f'''<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>{title}</title>
    <style>
        @page {{
            size: A4;
            margin-top: 2.5cm;
            margin-bottom: 2.5cm;
            margin-left: 3cm;
            margin-right: 3cm;
        }}
        * {{
            box-sizing: border-box;
        }}
        body {{
            font-family: "FangSong", "SimSun", "Songti SC", "Microsoft YaHei", serif;
            font-size: 14pt;
            line-height: 1.75;
            color: #333;
            max-width: 210mm;
            margin: 0 auto;
            padding: 0;
            background: white;
        }}
        .report-header {{
            text-align: center;
            border-bottom: 2px solid #1976d2;
            padding-bottom: 20px;
            margin-bottom: 32px;
        }}
        .report-header h1 {{
            color: #1976d2;
            margin: 0 0 12px 0;
            font-size: 22pt;
            letter-spacing: 1px;
            font-family: "SimHei", "Microsoft YaHei", sans-serif;
        }}
        .report-meta {{
            color: #666;
            font-size: 11pt;
            margin: 0;
            letter-spacing: 0.5px;
        }}
        .section {{
            margin: 30px 0;
            page-break-inside: avoid;
        }}
        .page-break {{
            page-break-after: always;
            margin: 28px 0;
            border-top: 1px dashed #e0e0e0;
        }}
        .section h2 {{
            color: #1976d2;
            border-left: 4px solid #1976d2;
            padding-left: 12px;
            margin: 0 0 18px 0;
            font-size: 16pt;
            font-family: "SimHei", "Microsoft YaHei", sans-serif;
        }}
        .section h3 {{
            color: #333;
            margin: 18px 0 10px 0;
            font-size: 14pt;
            font-family: "SimHei", "Microsoft YaHei", sans-serif;
        }}
        .detail-content, .conclusions, .recommendations {{
            text-align: justify;
            text-indent: 2em;
        }}
        p {{
            margin: 10px 0;
            text-align: justify;
            text-indent: 2em;
        }}
        ul, ol {{
            padding-left: 26px;
            margin: 10px 0;
        }}
        li {{
            margin: 6px 0;
            line-height: 1.75;
        }}
        .charts-section {{
            margin: 34px 0;
        }}
        .charts-section h2 {{
            color: #1976d2;
            border-left: 4px solid #1976d2;
            padding-left: 12px;
            margin: 0 0 22px 0;
            font-size: 16pt;
            font-family: "SimHei", "Microsoft YaHei", sans-serif;
        }}
        .charts-grid {{
            display: flex;
            flex-wrap: wrap;
            gap: 30px;
        }}
        .chart-item {{
            page-break-inside: avoid;
            break-inside: avoid;
            margin-bottom: 30px;
        }}
        .chart-item.full-width {{
            width: 100%;
        }}
        .chart-item img {{
            width: 100%;
            height: auto;
            border: 1px solid #ddd;
            border-radius: 4px;
            background: #f7f7f7;
        }}
        .chart-caption {{
            text-align: center;
            font-size: 11pt;
            color: #555;
            margin: 10px 0 0 0;
        }}
        .chart-description {{
            text-align: justify;
            font-size: 12pt;
            color: #333;
            margin: 8px 0 0 0;
            padding: 0 10px;
            line-height: 1.6;
        }}
        .chart-state-note {{
            text-align: center;
            font-size: 10pt;
            color: #777;
            margin-top: 6px;
        }}
        .chart-group {{
            page-break-inside: avoid;
            break-inside: avoid;
            margin-bottom: 30px;
            border: 1px solid #e0e0e0;
            border-radius: 4px;
            padding: 15px;
            background: #fafafa;
        }}
        .chart-group-title {{
            text-align: center;
            font-size: 12pt;
            color: #1976d2;
            margin: 0 0 12px 0;
            font-weight: bold;
            font-family: "SimHei", "Microsoft YaHei", sans-serif;
        }}
        .conclusions, .recommendations {{
            margin: 20px 0;
        }}
        .conclusions ul, .recommendations ul {{
            padding-left: 26px;
            margin: 10px 0;
        }}
        .conclusions li, .recommendations li {{
            margin: 6px 0;
        }}
        .report-footer {{
            margin-top: 44px;
            padding-top: 20px;
            border-top: 1px solid #ddd;
            text-align: center;
        }}
        .report-footer p {{
            font-size: 11pt;
            color: #777;
            margin: 6px 0;
            letter-spacing: 0.5px;
        }}
        table {{
            width: 100%;
            border-collapse: collapse;
            margin: 18px 0;
        }}
        th, td {{
            border: 1px solid #ddd;
            padding: 10px 14px;
            text-align: left;
        }}
        th {{
            background: #f5f5f5;
        }}
    </style>
</head>
<body>
    <div class="report-header">
        <h1>{title}</h1>
        <p class="report-meta">
            生成时间：{generated_at} | 分析置信度：{confidence}%
        </p>
    </div>
    
    {f'<div class="section"><h2>执行摘要</h2><div class="detail-content">{summary_html}</div></div>' if summary_html else ''}
    
    {conclusion_section}
    
    {weather_section}
    
    {component_section}
    
    {other_section}
    
    {detail_section}
    
    <div class="report-footer">
        <p>本报告由大气污染溯源分析系统自动生成</p>
        <p>报告仅供参考，具体决策请结合实际情况</p>
    </div>
</body>
</html>'''
        
        return html
    
    def _generate_grouped_charts_html(
        self,
        charts: List[Dict[str, Any]],
        weather_content: str = None,
        component_content: str = None,
        other_content: str = None
    ) -> Dict[str, str]:
        """
        生成分组后的图表HTML：气象 / 组分 / 其他
        返回 {"weather": html, "component": html, "other": html}（可能缺省）
        
        Args:
            charts: 图表列表
            weather_content: 气象分析章节的Markdown内容（用于匹配图表引用）
            component_content: 组分分析章节的Markdown内容
            other_content: 其他分析章节的Markdown内容
        """
        from collections import defaultdict
        
        grouped: Dict[str, List[Dict[str, Any]]] = defaultdict(list)
        for chart in charts or []:
            meta = chart.get("meta", {}) or {}
            
            # 方法0：优先从source_data_id推断（最可靠）
            source_data_id = meta.get("source_data_id") or meta.get("data_id") or ""
            source_data_id_lower = str(source_data_id).lower()
            
            # 方法1：从meta中获取expert信息
            expert_tag = str(
                meta.get("expert")
                or meta.get("generator")
                or meta.get("scenario")
                or meta.get("data_source")
                or chart.get("expert")  # 也检查chart的直接字段
                or ""
            ).lower()
            
            # 方法2：从图表标题和类型推断（如果meta中没有expert信息）
            chart_title = str(chart.get("title", "") or "").lower()
            chart_type = str(chart.get("type", "") or "").lower()
            
            # 判断逻辑：优先使用source_data_id，然后使用meta中的expert，最后根据标题和类型推断
            is_weather = False
            is_component = False
            
            # 优先从source_data_id判断
            if source_data_id:
                if any(keyword in source_data_id_lower for keyword in [
                    "meteorology", "trajectory", "气象", "轨迹", "风速", "风向", "边界层"
                ]):
                    is_weather = True
                elif any(keyword in source_data_id_lower for keyword in [
                    "vocs", "pmf", "obm", "component", "组分", "源解析", "贡献", "air_quality", "regional", "quick_obm"
                ]):
                    is_component = True
            
            # 如果source_data_id无法判断，使用expert_tag和标题/类型
            if not is_weather and not is_component:
                is_weather = (
                    "weather" in expert_tag 
                    or "气象" in str(meta.get("expert") or "")
                    or "meteorological_trajectory" in expert_tag
                    or "气象" in chart_title
                    or "meteorology" in chart_type
                    or "trajectory" in chart_type
                    or "轨迹" in chart_title
                    or "风速" in chart_title
                    or "风向" in chart_title
                    or "边界层" in chart_title
                )
                
                is_component = (
                    "component" in expert_tag 
                    or "组分" in str(meta.get("expert") or "")
                    or "quick_obm" in expert_tag
                    or "组分" in chart_title
                    or "vocs" in chart_type
                    or "pmf" in chart_type
                    or "obm" in chart_type
                    or "vocs" in chart_title
                    or "pmf" in chart_title
                    or "源解析" in chart_title
                    or "贡献" in chart_title
                )
            
            if is_weather:
                grouped["weather"].append(chart)
                logger.info("chart_grouped_to_weather", chart_title=chart.get("title"), expert_tag=expert_tag, has_meta=bool(meta))
            elif is_component:
                grouped["component"].append(chart)
                logger.info("chart_grouped_to_component", chart_title=chart.get("title"), expert_tag=expert_tag, has_meta=bool(meta))
            else:
                grouped["other"].append(chart)
                logger.info("chart_grouped_to_other", chart_title=chart.get("title"), expert_tag=expert_tag, chart_type=chart_type, has_meta=bool(meta), meta_keys=list(meta.keys()) if meta else [])
        
        html_map: Dict[str, str] = {}
        # 为每个分组生成HTML，并传入对应的章节内容以匹配图表引用
        section_contents = {
            "weather": weather_content or "",
            "component": component_content or "",
            "other": other_content or ""
        }
        
        for key, items in grouped.items():
            if not items:
                continue
            section_content = section_contents.get(key, "")
            logger.info(
                "building_charts_grid",
                category=key,
                chart_count=len(items),
                section_content_length=len(section_content) if section_content else 0,
                section_content_preview=section_content[:200] if section_content else "",
                has_chart_markers=bool(re.search(r'\[CHART_\d+_START\]', section_content or ""))
            )
            html_map[key] = self._build_charts_grid(items, section_content=section_content)
        return html_map
    
    def _parse_section_markers(self, content: str, section_type: str) -> str:
        """
        解析预设标识，提取章节内容
        
        Args:
            content: 包含预设标识的Markdown内容
            section_type: 章节类型（weather/component/conclusion）
            
        Returns:
            解析后的内容（去除标识，保留章节内容）
        """
        # 根据章节类型确定标识
        if section_type == "weather" or "weather" in str(section_type).lower() or "气象" in str(section_type):
            start_marker = "[WEATHER_SECTION_START]"
            end_marker = "[WEATHER_SECTION_END]"
        elif section_type == "component" or "component" in str(section_type).lower() or "组分" in str(section_type):
            start_marker = "[COMPONENT_SECTION_START]"
            end_marker = "[COMPONENT_SECTION_END]"
        elif section_type == "conclusion" or "conclusion" in str(section_type).lower() or "结论" in str(section_type) or "建议" in str(section_type):
            start_marker = "[CONCLUSION_SECTION_START]"
            end_marker = "[CONCLUSION_SECTION_END]"
        else:
            # 如果没有匹配，尝试查找所有可能的标识
            if "[WEATHER_SECTION_START]" in content:
                start_marker = "[WEATHER_SECTION_START]"
                end_marker = "[WEATHER_SECTION_END]"
            elif "[COMPONENT_SECTION_START]" in content:
                start_marker = "[COMPONENT_SECTION_START]"
                end_marker = "[COMPONENT_SECTION_END]"
            elif "[CONCLUSION_SECTION_START]" in content:
                start_marker = "[CONCLUSION_SECTION_START]"
                end_marker = "[CONCLUSION_SECTION_END]"
            else:
                # 没有标识，直接返回原内容
                return content
        
        # 提取标识之间的内容
        start_idx = content.find(start_marker)
        end_idx = content.find(end_marker)
        
        if start_idx >= 0:
            if end_idx > start_idx:
                # 有完整的开始和结束标识
                section_content = content[start_idx + len(start_marker):end_idx].strip()
            else:
                # 只有开始标识，提取到内容末尾
                section_content = content[start_idx + len(start_marker):].strip()
            
            # 【修复】保留图表标识 [CHART_X_START] 和 [CHART_X_END]，用于后续提取图表解析文本
            # 只移除章节级别的标识（如果还存在）
            section_content = re.sub(r'\[WEATHER_SECTION_START\]', '', section_content)
            section_content = re.sub(r'\[WEATHER_SECTION_END\]', '', section_content)
            section_content = re.sub(r'\[COMPONENT_SECTION_START\]', '', section_content)
            section_content = re.sub(r'\[COMPONENT_SECTION_END\]', '', section_content)
            section_content = re.sub(r'\[CONCLUSION_SECTION_START\]', '', section_content)
            section_content = re.sub(r'\[CONCLUSION_SECTION_END\]', '', section_content)
            
            return section_content.strip()
        else:
            # 如果没有找到开始标识，返回原内容（但也要去除可能存在的章节标识）
            # 【修复】保留图表标识，用于后续提取图表解析文本
            cleaned_content = content
            cleaned_content = re.sub(r'\[WEATHER_SECTION_START\]', '', cleaned_content)
            cleaned_content = re.sub(r'\[WEATHER_SECTION_END\]', '', cleaned_content)
            cleaned_content = re.sub(r'\[COMPONENT_SECTION_START\]', '', cleaned_content)
            cleaned_content = re.sub(r'\[COMPONENT_SECTION_END\]', '', cleaned_content)
            cleaned_content = re.sub(r'\[CONCLUSION_SECTION_START\]', '', cleaned_content)
            cleaned_content = re.sub(r'\[CONCLUSION_SECTION_END\]', '', cleaned_content)
            # 注意：不在这里移除图表标识，保留用于提取引用
            return cleaned_content.strip()
    
    def _remove_chart_sections(self, content: str) -> str:
        """
        从章节内容中移除图表解析部分（[CHART_X_START]... [CHART_X_END]）
        因为这些内容会在图表旁边显示，不需要在文本部分重复显示
        
        Args:
            content: 包含图表解析的Markdown内容
            
        Returns:
            移除图表解析后的内容
        """
        if not content:
            return content
        
        # 移除所有图表解析部分
        cleaned = re.sub(r'\[CHART_\d+_START\].*?\[CHART_\d+_END\]', '', content, flags=re.DOTALL)
        
        # 清理多余的空行
        cleaned = re.sub(r'\n\s*\n\s*\n+', '\n\n', cleaned)
        
        return cleaned.strip()
    
    def _extract_chart_references(self, section_content: str) -> List[Dict]:
        """从章节内容中提取图表引用和解析文本"""
        references = []
        
        if not section_content:
            logger.warning("extract_chart_references_empty_content")
            return references
        
        # 方法1：从预设标识中提取 [CHART_X_START]... [CHART_X_END]
        chart_pattern = r'\[CHART_(\d+)_START\](.*?)\[CHART_\d+_END\]'
        chart_matches = re.findall(chart_pattern, section_content, re.DOTALL)
        
        logger.info(
            "extract_chart_references_pattern_match",
            pattern=chart_pattern,
            matches_count=len(chart_matches),
            section_content_preview=section_content[:500] if section_content else "",
            section_content_has_markers=bool(re.search(r'\[CHART_\d+_START\]', section_content or ""))
        )
        
        for chart_num, chart_content in chart_matches:
            # 提取图表标题和解析文本
            lines = chart_content.strip().split('\n')
            title = ""
            description = ""
            
            for line in lines:
                line = line.strip()
                if not line:
                    continue
                # 匹配"图X：标题"格式
                title_match = re.match(r'图\s*\d+[:：]\s*(.+)', line)
                if title_match:
                    title = title_match.group(1).strip()
                elif title and not description:
                    # 第一行非标题内容作为解析文本
                    description = line
                elif description:
                    # 后续内容追加到解析文本
                    description += " " + line
            
            ref_data = {
                "number": int(chart_num),
                "title": title,
                "description": description.strip()
            }
            references.append(ref_data)
            
            # 记录每个引用的提取详情
            logger.info(
                "chart_reference_extracted_detail",
                chart_number=int(chart_num),
                extracted_title=title,
                extracted_description_length=len(description.strip()),
                extracted_description_preview=description.strip()[:100] if description.strip() else "",
                raw_content_preview=chart_content[:200] if chart_content else ""
            )
        
        # 方法2：如果没有预设标识，使用传统方式匹配"图X：标题"
        if not references:
            pattern = r'图\s*(\d+)[:：]\s*([^\n]+)'
            matches = re.findall(pattern, section_content)
            
            for match in matches:
                ref_num = int(match[0])
                ref_title = match[1].strip()
                references.append({
                    "number": ref_num,
                    "title": ref_title,
                    "description": ""
                })
        
        return references
    
    def _match_charts_to_references(
        self,
        charts: List[Dict],
        references: List[Dict]
    ) -> List[Dict]:
        """根据文本引用匹配实际图表"""
        matched = []
        used_charts = set()
        used_refs = set()
        
        # 记录匹配过程的详细信息
        match_details = []
        
        # 先按引用顺序匹配
        for ref in references:
            ref_num = ref["number"]
            ref_title = ref["title"]
            best_match = None
            best_match_idx = -1
            match_candidates = []
            
            for idx, chart in enumerate(charts):
                if idx in used_charts:
                    continue
                
                chart_title = chart.get("title", "")
                if not chart_title:
                    continue
                
                # 方法1：精确匹配或包含匹配
                is_contain_match = ref_title in chart_title or chart_title in ref_title
                
                # 方法2：模糊匹配（移除标点、空格，比较核心内容）
                ref_clean = re.sub(r'[^\w\u4e00-\u9fa5]', '', ref_title.lower())
                chart_clean = re.sub(r'[^\w\u4e00-\u9fa5]', '', chart_title.lower())
                
                # 检查关键词重叠
                ref_keywords = set([k for k in ref_clean.split() if len(k) > 1])
                chart_keywords = set([k for k in chart_clean.split() if len(k) > 1])
                keyword_overlap = ref_keywords & chart_keywords
                overlap_ratio = len(keyword_overlap) / len(ref_keywords | chart_keywords) if (ref_keywords | chart_keywords) else 0
                is_fuzzy_match = overlap_ratio > 0.3  # 30%以上关键词重叠
                
                # 方法3：检查是否包含核心关键词（如"轨迹"、"企业"等）
                core_keywords = ["轨迹", "企业", "站点", "分析", "分布", "风向", "风速"]
                has_core_keyword = any(kw in ref_title and kw in chart_title for kw in core_keywords)
                
                is_match = is_contain_match or is_fuzzy_match or has_core_keyword
                
                if is_match:
                    match_candidates.append({
                        "idx": idx,
                        "chart_title": chart_title,
                        "match_type": "contain" if is_contain_match else ("fuzzy" if is_fuzzy_match else "keyword"),
                        "overlap_ratio": overlap_ratio,
                        "keyword_overlap": list(keyword_overlap)
                    })
                    # 选择匹配度最高的
                    if best_match is None:
                        best_match = chart
                        best_match_idx = idx
                    elif is_contain_match and not (ref_title in best_match.get("title", "") or best_match.get("title", "") in ref_title):
                        # 精确匹配优先
                        best_match = chart
                        best_match_idx = idx
                    elif overlap_ratio > 0.5:  # 关键词重叠度高的优先
                        best_match = chart
                        best_match_idx = idx
            
            if best_match:
                matched.append({
                    **best_match,
                    "text_reference_number": ref_num,
                    "text_reference_title": ref_title
                })
                used_charts.add(best_match_idx)
                used_refs.add(ref_num)
                match_details.append({
                    "ref_number": ref_num,
                    "ref_title": ref_title,
                    "matched": True,
                    "chart_title": best_match.get("title"),
                    "match_type": match_candidates[0]["match_type"] if match_candidates else "unknown",
                    "candidates_count": len(match_candidates)
                })
            else:
                match_details.append({
                    "ref_number": ref_num,
                    "ref_title": ref_title,
                    "matched": False,
                    "candidates_count": len(match_candidates),
                    "available_charts": [c.get("title", "") for c in charts if charts.index(c) not in used_charts]
                })
        
        # 添加未匹配的图表（按原始顺序）
        unmatched_charts = []
        for idx, chart in enumerate(charts):
            if idx not in used_charts:
                matched.append(chart)
                unmatched_charts.append({
                    "idx": idx,
                    "title": chart.get("title", "")
                })
        
        # 按引用编号排序（如果有引用编号）
        matched_with_refs = [m for m in matched if "text_reference_number" in m]
        matched_without_refs = [m for m in matched if "text_reference_number" not in m]

        # 有引用的图表按引用编号排序
        matched_with_refs.sort(key=lambda x: x.get("text_reference_number", 999))

        # 合并排序结果
        matched = matched_with_refs + matched_without_refs
        
        # 记录详细的匹配结果
        logger.info(
            "chart_matching_details",
            total_references=len(references),
            total_charts=len(charts),
            matched_count=len([c for c in matched if "text_reference_number" in c]),
            unmatched_refs=[d for d in match_details if not d["matched"]],
            unmatched_charts=unmatched_charts,
            match_details=match_details
        )
        
        return matched
    
    def _build_charts_grid(
        self,
        charts: List[Dict[str, Any]],
        section_content: str = None
    ) -> str:
        """统一的图表网格构建，默认单列全宽，扩大间距避免拥挤，支持图表分组"""
        if not charts:
            return ""

        # 【关键修复】先标准化所有图表数据，使用统一数据通道
        normalized_charts = [self._normalize_chart_data(chart) for chart in charts]

        # 初始化已使用的描述编号集合（用于容错分配）
        self._used_descriptions = set()

        # 如果提供了章节内容，尝试匹配图表引用
        matched_charts = normalized_charts
        chart_descriptions = {}  # 存储图表编号对应的解析文本
        
        # 记录输入的图表信息（用于诊断）
        logger.info(
            "build_charts_grid_input",
            category=section_content[:50] if section_content else "unknown",
            charts_count=len(charts),
            charts_titles=[c.get("title", "") for c in charts],
            has_section_content=bool(section_content),
            section_content_length=len(section_content) if section_content else 0
        )
        
        if section_content:
            references = self._extract_chart_references(section_content)
            if references:
                # 构建描述映射
                for ref in references:
                    chart_descriptions[ref["number"]] = ref.get("description", "")
                
                logger.info(
                    "chart_references_extracted",
                    references_count=len(references),
                    references=[{"number": r["number"], "title": r["title"], "has_description": bool(r.get("description")), "description_preview": r.get("description", "")[:100]} for r in references],
                    descriptions_map={k: v[:50] + "..." if len(v) > 50 else v for k, v in chart_descriptions.items()},
                    chart_titles_from_content=[r["title"] for r in references],
                    chart_titles_from_charts=[c.get("title", "") for c in charts]
                )
                
                matched_charts = self._match_charts_to_references(charts, references)
                logger.info(
                    "charts_matched_to_references",
                    references_count=len(references),
                    matched_count=len([c for c in matched_charts if "text_reference_number" in c]),
                    unmatched_count=len([c for c in matched_charts if "text_reference_number" not in c]),
                    matched_charts=[{"title": c.get("title"), "ref_number": c.get("text_reference_number"), "ref_title": c.get("text_reference_title")} for c in matched_charts if "text_reference_number" in c],
                    unmatched_charts=[{"title": c.get("title")} for c in matched_charts if "text_reference_number" not in c]
                )
            else:
                logger.warning(
                    "no_chart_references_found",
                    section_content_preview=section_content[:500] if section_content else "",
                    section_content_has_chart_markers=bool(re.search(r'\[CHART_\d+_START\]', section_content or ""))
                )
        
        # 按 chart_group_id 分组图表
        from collections import defaultdict
        grouped_charts = defaultdict(list)
        ungrouped_charts = []
        
        for chart in matched_charts:
            meta = chart.get("meta", {}) or {}
            chart_group_id = meta.get("chart_group_id")
            if chart_group_id:
                grouped_charts[chart_group_id].append(chart)
            else:
                ungrouped_charts.append(chart)
        
        items = []
        chart_index = 0
        
        # 处理分组图表（合并为一组，对应一个解析内容）
        for group_id, group_charts in grouped_charts.items():
            chart_index += 1
            group_title = f"图表组 {chart_index}"
            
            # 尝试从章节内容中查找该组的解析文本
            # 对于分组图表，尝试查找组内所有图表的解析文本，使用第一个非空的描述
            group_description = ""
            if group_charts and section_content:
                # 方法1：优先使用第一个图表的引用编号
                first_chart = group_charts[0]
                if "text_reference_number" in first_chart:
                    ref_num = first_chart["text_reference_number"]
                    group_description = chart_descriptions.get(ref_num, "")
                
                # 方法2：如果第一个图表没有描述，尝试查找组内其他图表的描述
                if not group_description:
                    for chart in group_charts:
                        if "text_reference_number" in chart:
                            ref_num = chart["text_reference_number"]
                            desc = chart_descriptions.get(ref_num, "")
                            if desc:
                                group_description = desc
                                break
                
                # 诊断日志
                logger.debug(
                    "group_chart_description",
                    group_id=group_id,
                    chart_count=len(group_charts),
                    first_chart_title=first_chart.get("title"),
                    first_chart_has_ref_number="text_reference_number" in first_chart,
                    first_chart_ref_number=first_chart.get("text_reference_number"),
                    chart_descriptions_keys=list(chart_descriptions.keys()),
                    has_description=bool(group_description),
                    description_length=len(group_description) if group_description else 0,
                    description_preview=group_description[:100] if group_description else ""
                )
            
            # 生成组内所有图表的HTML
            group_items = []
            for i, chart in enumerate(group_charts):
                # 【修复】使用统一数据通道读取图片数据
                unified_image_data = chart.get("unified_image_data", "")
                title = chart.get("text_reference_title") or chart.get("title", f"图表 {i + 1}")
                chart_type = chart.get("type", "chart")
                user_state = chart.get("user_state")

                size_class = "full-width"

                if unified_image_data and unified_image_data.startswith("data:image"):
                    img_html = f'<img src="{unified_image_data}" alt="{title}">'
                else:
                    img_html = f'''
                    <div style="width:100%;height:220px;background:#f5f5f5;display:flex;align-items:center;justify-content:center;border-radius:4px;">
                        <span style="color:#999;">{title} ({chart_type})</span>
                    </div>
                    '''
                
                state_note = ""
                if user_state:
                    notes = []
                    data_zoom = user_state.get("dataZoom", [])
                    if data_zoom and any(dz.get("start", 0) > 0 or dz.get("end", 100) < 100 for dz in data_zoom):
                        notes.append("已调整时间范围")
                    legend_selected = user_state.get("legendSelected", {})
                    if legend_selected and any(v is False for v in legend_selected.values()):
                        notes.append("已隐藏部分指标")
                    if notes:
                        state_note = f'<p class="chart-state-note">({", ".join(notes)})</p>'
                
                # 组内图表使用子标题
                sub_title = f"{title}" if i == 0 else f"{title}（续）"
                group_items.append(f'''
                <div class="chart-item {size_class}">
                    {img_html}
                    <p class="chart-caption">{sub_title}</p>
                    {state_note}
                </div>
                ''')
            
            # 为整个组添加统一的描述和标题
            description_html = ""
            if group_description:
                description_html = f'<p class="chart-description">{group_description}</p>'
            
            # 使用组内第一个图表的引用编号作为组的编号
            group_display_index = chart_index
            if group_charts and "text_reference_number" in group_charts[0]:
                group_display_index = group_charts[0]["text_reference_number"]
            
            items.append(f'''
            <div class="chart-group">
                <p class="chart-group-title">图 {group_display_index}：{group_charts[0].get("text_reference_title") or group_charts[0].get("title", group_title)}</p>
                {description_html}
                <div class="charts-grid">
                    {"".join(group_items)}
                </div>
            </div>
            ''')
        
        # 处理未分组的图表
        for i, chart in enumerate(ungrouped_charts):
            # 【修复】使用统一数据通道读取图片数据
            unified_image_data = chart.get("unified_image_data", "")
            # 优先使用文本中的引用标题，否则使用图表标题
            title = chart.get("text_reference_title") or chart.get("title", f"图表 {chart_index + 1}")
            chart_type = chart.get("type", "chart")
            user_state = chart.get("user_state")
            
            # 优先使用文本中的引用编号
            if "text_reference_number" in chart:
                display_index = chart["text_reference_number"]
            else:
                order_index = chart.get("order", chart_index)
                display_index = (order_index if order_index is not None else chart_index) + 1
            
            # 获取图表解析文本
            description = chart_descriptions.get(display_index, "")
            
            # 诊断日志：记录描述文本获取过程（改为INFO级别，确保能看到）
            logger.info(
                "chart_description_lookup",
                chart_title=chart.get("title"),
                chart_index=i,
                has_text_reference_number="text_reference_number" in chart,
                text_reference_number=chart.get("text_reference_number"),
                display_index=display_index,
                chart_descriptions_keys=list(chart_descriptions.keys()),
                description_found=bool(description),
                description_length=len(description) if description else 0,
                description_preview=description[:100] if description else "",
                # 添加对比信息，便于诊断
                available_descriptions_count=len(chart_descriptions),
                chart_descriptions_numbers=list(chart_descriptions.keys())
            )
            
            # 如果匹配失败但描述存在，尝试按顺序分配
            if not description and chart_descriptions:
                # 尝试使用图表在列表中的位置作为索引
                sequential_index = i + 1
                if sequential_index in chart_descriptions:
                    description = chart_descriptions[sequential_index]
                    logger.info(
                        "chart_description_fallback_sequential",
                        chart_title=chart.get("title"),
                        sequential_index=sequential_index,
                        description_found=bool(description)
                    )
                else:
                    # 如果还是没有，尝试查找第一个未使用的描述
                    if not hasattr(self, '_used_descriptions'):
                        self._used_descriptions = set()
                    for num, desc in sorted(chart_descriptions.items()):
                        if desc and num not in self._used_descriptions:
                            description = desc
                            self._used_descriptions.add(num)
                            logger.info(
                                "chart_description_fallback_first_available",
                                chart_title=chart.get("title"),
                                used_description_number=num,
                                description_found=bool(description)
                            )
                            break
            
            size_class = "full-width"  # 统一全宽，避免图片挤在一起

            if unified_image_data and unified_image_data.startswith("data:image"):
                img_html = f'<img src="{unified_image_data}" alt="{title}">'
            else:
                img_html = f'''
                <div style="width:100%;height:220px;background:#f5f5f5;display:flex;align-items:center;justify-content:center;border-radius:4px;">
                    <span style="color:#999;">{title} ({chart_type})</span>
                </div>
                '''
            
            state_note = ""
            if user_state:
                notes = []
                data_zoom = user_state.get("dataZoom", [])
                if data_zoom and any(dz.get("start", 0) > 0 or dz.get("end", 100) < 100 for dz in data_zoom):
                    notes.append("已调整时间范围")
                legend_selected = user_state.get("legendSelected", {})
                if legend_selected and any(v is False for v in legend_selected.values()):
                    notes.append("已隐藏部分指标")
                if notes:
                    state_note = f'<p class="chart-state-note">({", ".join(notes)})</p>'
            
            # 如果有解析文本，添加描述
            description_html = ""
            if description:
                description_html = f'<p class="chart-description">{description}</p>'
            
            item_html = f'''
            <div class="chart-item {size_class}">
                {img_html}
                <p class="chart-caption">图 {display_index}：{title}</p>
                {description_html}
                {state_note}
            </div>
            '''
            items.append(item_html)
            chart_index += 1
        
        return f'''
        <div class="charts-section">
            <div class="charts-grid">
                {"".join(items)}
            </div>
        </div>
        '''
    
    def _markdown_to_html(self, md_text: str) -> str:
        """Markdown转HTML"""
        # 【修复】在转换为HTML之前，移除所有标识，避免在最终报告中显示这些标识
        # 这些标识已经用于提取图表引用，现在可以安全移除
        cleaned_md = re.sub(r'\[CHART_\d+_START\]', '', md_text)
        cleaned_md = re.sub(r'\[CHART_\d+_END\]', '', cleaned_md)
        # 同时移除章节级别的标识（作为保险，虽然_parse_section_markers应该已经移除了）
        cleaned_md = re.sub(r'\[WEATHER_SECTION_START\]', '', cleaned_md)
        cleaned_md = re.sub(r'\[WEATHER_SECTION_END\]', '', cleaned_md)
        cleaned_md = re.sub(r'\[COMPONENT_SECTION_START\]', '', cleaned_md)
        cleaned_md = re.sub(r'\[COMPONENT_SECTION_END\]', '', cleaned_md)
        cleaned_md = re.sub(r'\[CONCLUSION_SECTION_START\]', '', cleaned_md)
        cleaned_md = re.sub(r'\[CONCLUSION_SECTION_END\]', '', cleaned_md)
        
        try:
            import markdown
            return markdown.markdown(cleaned_md, extensions=['tables', 'fenced_code'])
        except ImportError:
            # 简单处理
            lines = cleaned_md.split('\n')
            html_lines = []
            for line in lines:
                if line.startswith('### '):
                    html_lines.append(f'<h3>{line[4:]}</h3>')
                elif line.startswith('## '):
                    html_lines.append(f'<h4>{line[3:]}</h4>')
                elif line.startswith('# '):
                    html_lines.append(f'<h3>{line[2:]}</h3>')
                elif line.startswith('- '):
                    html_lines.append(f'<li>{line[2:]}</li>')
                elif line.strip():
                    html_lines.append(f'<p>{line}</p>')
            return '\n'.join(html_lines)
    
    async def _html_to_pdf(self, html_content: str) -> bytes:
        """HTML转PDF - 多种方案fallback"""
        
        # 方案1: 尝试使用pdfkit (需要wkhtmltopdf)
        try:
            import pdfkit
            
            options = {
                'encoding': 'UTF-8',
                'page-size': 'A4',
                'margin-top': '15mm',
                'margin-right': '15mm',
                'margin-bottom': '15mm',
                'margin-left': '15mm',
                'enable-local-file-access': None
            }
            pdf_bytes = pdfkit.from_string(html_content, False, options=options)
            logger.info("pdf_generated_with_pdfkit", size=len(pdf_bytes))
            return pdf_bytes
        except Exception as e:
            logger.warning("pdfkit_failed", error=str(e))
        
        # 方案2: 尝试使用weasyprint
        try:
            from weasyprint import HTML
            
            pdf_bytes = HTML(string=html_content, base_url=".").write_pdf()
            logger.info("pdf_generated_with_weasyprint", size=len(pdf_bytes))
            return pdf_bytes
        except Exception as e:
            logger.warning("weasyprint_failed", error=str(e))
        
        # 方案3: 返回HTML，让前端提示用户使用浏览器打印
        logger.warning("pdf_fallback_to_html", message="PDF库不可用，返回HTML格式")
        raise ImportError(
            "PDF导出暂不可用。\n"
            "请选择HTML格式导出，然后在浏览器中使用Ctrl+P打印为PDF。\n"
            "或安装wkhtmltopdf: https://wkhtmltopdf.org/downloads.html"
        )
    
    async def _html_to_docx(
        self,
        html_content: str,
        charts: List[Dict[str, Any]]
    ) -> bytes:
        """HTML转Word文档"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt, Cm
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            import io
            import base64
            import re
            
            doc = Document()
            
            # 设置页面边距
            sections = doc.sections
            for section in sections:
                section.top_margin = Cm(2)
                section.bottom_margin = Cm(2)
                section.left_margin = Cm(2.5)
                section.right_margin = Cm(2.5)
            
            # 从HTML提取内容
            title_match = re.search(r'<h1[^>]*>([^<]+)</h1>', html_content)
            title = title_match.group(1) if title_match else "大气污染溯源分析报告"
            
            meta_match = re.search(r'<p class="report-meta"[^>]*>([^<]+)</p>', html_content)
            meta = meta_match.group(1) if meta_match else ""
            
            # 添加标题
            heading = doc.add_heading(title, 0)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加元数据
            if meta:
                meta_para = doc.add_paragraph(meta)
                meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            doc.add_paragraph()
            
            # 添加图表
            if charts:
                doc.add_heading("分析图表", 1)
                
                for i, chart in enumerate(charts):
                    preview_image = chart.get("preview_image", "")
                    chart_title = chart.get("title", f"图表 {i + 1}")
                    
                    if preview_image and preview_image.startswith("data:image"):
                        try:
                            # 提取base64数据
                            _, data = preview_image.split(",", 1)
                            img_bytes = base64.b64decode(data)
                            
                            img_stream = io.BytesIO(img_bytes)
                            doc.add_picture(img_stream, width=Inches(5.5))
                            
                            # 添加图片说明
                            caption = doc.add_paragraph(f"图 {i + 1}：{chart_title}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                        except Exception as e:
                            logger.warning(f"Failed to add chart image: {e}")
                            doc.add_paragraph(f"[图表: {chart_title}]")
                    else:
                        doc.add_paragraph(f"[图表: {chart_title}]")
                    
                    doc.add_paragraph()
            
            # 提取结论
            conclusions_match = re.search(
                r'<div class="conclusions">\s*<h3>主要结论</h3>\s*<ul>(.*?)</ul>',
                html_content,
                re.DOTALL
            )
            if conclusions_match:
                doc.add_heading("主要结论", 1)
                items = re.findall(r'<li>([^<]+)</li>', conclusions_match.group(1))
                for item in items:
                    doc.add_paragraph(item.strip(), style='List Bullet')
            
            # 提取建议
            recommendations_match = re.search(
                r'<div class="recommendations">\s*<h3>控制建议</h3>\s*<ul>(.*?)</ul>',
                html_content,
                re.DOTALL
            )
            if recommendations_match:
                doc.add_heading("控制建议", 1)
                items = re.findall(r'<li>([^<]+)</li>', recommendations_match.group(1))
                for item in items:
                    doc.add_paragraph(item.strip(), style='List Bullet')
            
            # 添加页脚
            doc.add_paragraph()
            footer = doc.add_paragraph("本报告由大气污染溯源分析系统自动生成")
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 保存到内存
            output = io.BytesIO()
            doc.save(output)
            result = output.getvalue()
            
            logger.info("docx_generated", size=len(result))
            return result
            
        except ImportError:
            logger.warning("python-docx_not_installed")
            raise ImportError("Word导出需要安装python-docx: pip install python-docx")
